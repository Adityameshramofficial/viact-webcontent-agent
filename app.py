"""viact.ai Content Intelligence Suite — Multi-Agent Content Pipeline"""
import os
import sys
import json
import html as _html_lib

import streamlit as st


def _t(text: str) -> str:
    """Sanitize dynamic text for HTML embedding — escapes HTML chars, strips newlines."""
    return _html_lib.escape(str(text)).replace('\n', ' ').replace('\r', '')


def _html(s: str) -> str:
    """Collapse a multi-line HTML string to a single line so CommonMark never treats
    indented lines as code blocks and blank lines never break HTML block mode."""
    return " ".join(part.strip() for part in s.splitlines() if part.strip())

# ── Load secrets (localhost: .env  |  Streamlit Cloud: st.secrets) ───────────
# override=True ensures an updated .env is picked up without restarting the process
try:
    from dotenv import load_dotenv
    load_dotenv(override=True)
except Exception:
    pass

# Streamlit Cloud fallback: if a key wasn't set by .env (cloud has no .env),
# pull from st.secrets.  Locally, .env wins because load_dotenv already set it.
_SECRET_KEYS = [
    "GROQ_API_KEY", "TAVILY_API_KEY", "FIRECRAWL_API_KEY",
    "SHEET_ID", "GOOGLE_CLIENT_ID", "GOOGLE_CLIENT_SECRET",
    "GCP_SERVICE_ACCOUNT",
]
for _k in _SECRET_KEYS:
    try:
        if not os.environ.get(_k):
            os.environ[_k] = str(st.secrets[_k])
    except Exception:
        pass

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))

try:
    from generate_images import generate_image as _gen_img, extract_dims as _edims
except Exception:
    def _gen_img(*a, **kw): return None  # type: ignore
    def _edims(p, dw=1200, dh=630): return dw, dh  # type: ignore

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="viact.ai Webpage Agent",
    page_icon="🏗️",
    layout="wide",
)

# ── viAct Design System — Industrial Intelligence Dashboard ──────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Oxanium:wght@300;400;600;700&family=Jost:wght@300;400;500;600;700;800&family=Caveat:wght@500;600&display=swap');

/* ── viAct Design Token System ── */
:root {
  --bg-dark:     #0a0a0f;
  --bg-dark-2:   #0d0d16;
  --bg-card:     #12121c;
  --bg-card-2:   #0e0e1a;
  --orange:      #ff6a3d;
  --orange-soft: rgba(255,106,61,0.10);
  --orange-bdr:  rgba(255,106,61,0.22);
  --t-white:  #ffffff;
  --t-g1:     #E9ECF1;
  --t-g2:     #C9D0D9;
  --t-g3:     #A8B0BE;
  --t-g4:     #818181;
  --bdr-dark: rgba(255,255,255,0.07);
  --shadow-sm: 0 2px 12px rgba(0,0,0,0.08);
  --shadow-md: 0 8px 32px rgba(0,0,0,0.18);
  --shadow-or: 0 8px 32px rgba(255,106,61,0.14);
}

/* ── Global reset ── */
html, body, .stApp { font-family: 'Jost', sans-serif !important; background: var(--bg-dark) !important; }
.stMainBlockContainer, .main .block-container { max-width: 1020px !important; padding: 0 2rem 6rem 2rem !important; }
#MainMenu, footer, header { visibility: hidden !important; }

/* ── Typography ── */
h1, h2, h3, h4 { font-family: 'Jost', sans-serif !important; color: var(--t-g1) !important; }
p, li, span, label, div { color: var(--t-g2); }
.stMarkdown p { color: var(--t-g2); }

/* ── Sidebar ── */
[data-testid="stSidebar"] { background-color: var(--bg-dark-2) !important; border-right: 1px solid rgba(255,255,255,0.05); }

/* ── Input labels — Oxanium precision style ── */
.stTextInput label, .stTextArea label, .stFileUploader label, .stSelectbox label {
    font-family: 'Oxanium', sans-serif !important;
    font-size: 0.62rem !important; font-weight: 600 !important;
    letter-spacing: 2px !important; text-transform: uppercase !important;
    color: #3a3a4a !important;
}

/* ── Text inputs ── */
.stTextInput input {
    background: var(--bg-card) !important; border: 1.5px solid rgba(255,255,255,0.06) !important;
    border-radius: 8px !important; color: var(--t-g1) !important;
    font-family: 'Jost', sans-serif !important; font-size: 0.95rem !important;
    padding: 0.65rem 0.9rem !important;
}
.stTextInput input:focus { border-color: var(--orange) !important; box-shadow: 0 0 0 3px var(--orange-soft) !important; outline: none !important; }
.stTextInput input::placeholder { color: #252535 !important; }

/* ── Text areas ── */
div[data-baseweb="textarea"] > div { background: var(--bg-card) !important; border: 1.5px solid rgba(255,255,255,0.06) !important; border-radius: 8px !important; }
.stTextArea textarea { background: transparent !important; border: none !important; color: #aab0be !important; font-family: 'Jost', sans-serif !important; font-size: 0.85rem !important; padding: 0.65rem 0.9rem !important; line-height: 1.65 !important; }
.stTextArea textarea:focus { outline: none !important; }
div[data-baseweb="textarea"] > div:focus-within { border-color: var(--orange) !important; box-shadow: 0 0 0 3px var(--orange-soft) !important; }
.stTextArea textarea::placeholder { color: #1e1e2e !important; }

/* ── Selectbox ── */
div[data-baseweb="select"] > div { background: var(--bg-card) !important; border: 1.5px solid rgba(255,255,255,0.06) !important; border-radius: 8px !important; color: var(--t-g1) !important; }

/* ── File uploader ── */
[data-testid="stFileUploaderDropzone"] { background: var(--bg-card) !important; border: 1.5px dashed rgba(255,255,255,0.07) !important; border-radius: 8px !important; padding: 0.75rem 1.1rem !important; box-sizing: border-box !important; }
[data-testid="stFileUploaderDropzone"] > div { display: flex !important; flex-direction: row !important; align-items: center !important; justify-content: space-between !important; gap: 1rem !important; flex-wrap: nowrap !important; }
[data-testid="stFileUploaderDropzone"]:hover { border-color: var(--orange-bdr) !important; }
[data-testid="stFileUploaderDropzone"] span, [data-testid="stFileUploaderDropzone"] small, [data-testid="stFileUploaderDropzone"] p { color: #2e2e3e !important; font-family: 'Jost', sans-serif !important; }
[data-testid="stFileUploaderDropzone"] button, [data-testid="stFileUploader"] section button { border-radius: 6px !important; border: 1px solid rgba(255,255,255,0.08) !important; color: #444 !important; background: var(--bg-dark-2) !important; font-size: 0.72rem !important; letter-spacing: 0.5px !important; padding: 0.35rem 0.9rem !important; flex-shrink: 0 !important; white-space: nowrap !important; }
[data-testid="stFileUploaderDropzone"] button:hover, [data-testid="stFileUploader"] section button:hover { border-color: var(--orange) !important; color: var(--orange) !important; background: var(--bg-dark-2) !important; box-shadow: none !important; }

/* ── Primary button — solid orange ── */
button[kind="primary"], button[data-testid="baseButton-primary"] {
    font-family: 'Oxanium', sans-serif !important; font-weight: 600 !important;
    font-size: 0.7rem !important; letter-spacing: 2px !important;
    text-transform: uppercase !important; border-radius: 6px !important;
    background: var(--orange) !important; color: #fff !important;
    border: 2px solid var(--orange) !important; transition: all 0.18s ease !important;
    cursor: pointer !important;
}
button[kind="primary"]:hover, button[data-testid="baseButton-primary"]:hover {
    background: #e55a2e !important; border-color: #e55a2e !important;
    box-shadow: var(--shadow-or) !important; transform: translateY(-1px) !important;
}
button[kind="primary"]:focus-visible, button[data-testid="baseButton-primary"]:focus-visible {
    outline: 2px solid var(--orange) !important; outline-offset: 3px !important;
}

/* ── Secondary button — ghost ── */
button[kind="secondary"], button[data-testid="baseButton-secondary"] {
    font-family: 'Oxanium', sans-serif !important; font-weight: 600 !important;
    font-size: 0.65rem !important; letter-spacing: 1.5px !important;
    text-transform: uppercase !important; border-radius: 20px !important;
    background: transparent !important; color: var(--orange) !important;
    border: 1.5px solid var(--orange-bdr) !important; transition: all 0.18s ease !important;
    cursor: pointer !important;
}
button[kind="secondary"]:hover, button[data-testid="baseButton-secondary"]:hover { background: var(--orange-soft) !important; border-color: var(--orange) !important; }

/* ── Download button ── */
[data-testid="stDownloadButton"] button { border-radius: 5px !important; border-color: rgba(255,255,255,0.07) !important; color: #444 !important; font-size: 0.65rem !important; padding: 0.3rem 0.8rem !important; cursor: pointer !important; }
[data-testid="stDownloadButton"] button:hover { border-color: var(--orange) !important; color: var(--orange) !important; background: transparent !important; }

/* ── Tabs — Oxanium nav bar style ── */
.stTabs [data-baseweb="tab-list"] {
    background: var(--bg-dark-2) !important;
    border: 1px solid rgba(255,255,255,0.05) !important;
    border-radius: 10px !important; gap: 4px !important;
    padding: 5px !important; margin-bottom: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important; color: #2e2e3e !important;
    font-family: 'Oxanium', sans-serif !important; font-weight: 600 !important;
    font-size: 0.72rem !important; letter-spacing: 1px !important;
    text-transform: uppercase !important;
    padding: 0.65rem 1.2rem !important; border: none !important;
    border-radius: 7px !important; transition: all 0.15s ease !important;
    cursor: pointer !important;
}
.stTabs [data-baseweb="tab"]:hover { color: var(--t-g3) !important; background: rgba(255,255,255,0.03) !important; }
.stTabs [aria-selected="true"] {
    color: #fff !important; background: var(--orange) !important;
    box-shadow: 0 2px 16px rgba(255,106,61,0.35) !important;
}
.stTabs [data-baseweb="tab-panel"] {
    background: var(--bg-dark-2) !important;
    border: 1px solid rgba(255,255,255,0.05) !important;
    border-top: none !important; border-radius: 0 0 10px 10px !important;
    padding: 1.8rem !important;
}

/* ── Expanders ── */
details[data-testid="stExpander"] { border: 1px solid rgba(255,255,255,0.05) !important; border-radius: 8px !important; background: var(--bg-card-2) !important; overflow: hidden !important; }
details[data-testid="stExpander"] summary { font-family: 'Oxanium', sans-serif !important; color: #444 !important; font-size: 0.72rem !important; font-weight: 600 !important; letter-spacing: 1px !important; text-transform: uppercase !important; padding: 0.7rem 1rem !important; list-style: none !important; cursor: pointer !important; }
details[data-testid="stExpander"] summary:hover { color: var(--orange) !important; }
details[data-testid="stExpander"] summary::-webkit-details-marker { display: none; }
details[data-testid="stExpander"] > div { background: var(--bg-dark) !important; padding: 1rem !important; }

/* ── Code blocks ── */
pre, .stCode { background: #05050a !important; border: 1px solid rgba(255,255,255,0.05) !important; border-radius: 8px !important; }
code { background: var(--bg-card) !important; color: var(--orange) !important; border: 1px solid rgba(255,255,255,0.06) !important; border-radius: 4px !important; padding: 2px 6px !important; font-size: 0.8rem !important; }
pre code { color: var(--t-g2) !important; }

/* ── Alerts ── */
.stAlert { border-radius: 8px !important; font-family: 'Jost', sans-serif !important; }
div[data-testid="stAlert"] { border-radius: 8px !important; border-left-width: 3px !important; }
.stSuccess { border-left-color: var(--orange) !important; }

/* ── Divider ── */
hr { border-color: rgba(255,255,255,0.05) !important; }

/* ── Radio ── */
div[data-testid="stRadio"] label { color: var(--t-g2) !important; }

/* ── Spinner ── */
.stSpinner > div > div { border-top-color: var(--orange) !important; }

/* ── Checkbox ── */
div[data-testid="stCheckbox"] label { color: var(--t-g2) !important; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 3px; height: 3px; }
::-webkit-scrollbar-track { background: var(--bg-dark); }
::-webkit-scrollbar-thumb { background: rgba(255,255,255,0.08); border-radius: 4px; }
::-webkit-scrollbar-thumb:hover { background: var(--orange); }

/* ══════════════════════ CUSTOM COMPONENT CLASSES ══════════════════════ */

/* Eyebrow label — Oxanium precision tags */
.eyebrow {
    display: inline-flex; align-items: center; gap: 10px;
    font-family: 'Oxanium', sans-serif; font-weight: 300;
    font-size: 10px; letter-spacing: 0.2em; text-transform: uppercase;
    color: var(--t-g3); margin-bottom: 10px;
}
.eyebrow::before { content: ''; display: block; width: 16px; height: 2px; background: var(--orange); }

/* Glass Card */
.glass-card {
    background: rgba(18,18,28,0.92); backdrop-filter: blur(12px);
    border: 1px solid var(--orange-bdr); border-radius: 12px;
    padding: 24px; position: relative; overflow: hidden;
    box-shadow: 0 4px 24px rgba(0,0,0,0.4); transition: all 0.2s ease;
    margin-bottom: 16px;
}
.glass-card:hover { border-color: rgba(255,106,61,0.35); box-shadow: 0 8px 32px rgba(255,106,61,0.09); }

/* Dashboard Stat Card */
.stat-card {
    background: var(--bg-card); border: 1px solid rgba(255,255,255,0.05);
    border-radius: 10px; padding: 18px 20px; position: relative; overflow: hidden;
    transition: border-color 0.2s ease;
}
.stat-card::before {
    content: ''; position: absolute; top: 0; left: 0; right: 0; height: 2px;
    background: var(--orange); opacity: 0.7;
}
.stat-card:hover { border-color: var(--orange-bdr); }
.stat-card-label { font-family: 'Oxanium', sans-serif; font-size: 0.6rem; font-weight: 600; letter-spacing: 2px; text-transform: uppercase; color: #3a3a4a; margin-bottom: 8px; }
.stat-card-value { font-family: 'Oxanium', sans-serif; font-size: 2rem; font-weight: 700; color: var(--orange); line-height: 1; margin-bottom: 4px; }
.stat-card-sub { font-family: 'Jost', sans-serif; font-size: 0.72rem; color: #2e2e3e; }

/* Metrics */
.metric-title { color: #3a3a4a; font-family: 'Oxanium', sans-serif; font-size: 0.62rem; font-weight: 600; text-transform: uppercase; letter-spacing: 2px; margin-bottom: 6px; }
.metric-value { color: var(--orange); font-family: 'Oxanium', sans-serif; font-size: 2.2rem; font-weight: 700; line-height: 1.1; margin-bottom: 6px; }

/* Status Badges */
.badge-confirmed { background: #0d4429; color: #3fb950; padding: 3px 10px; border-radius: 20px; font-family: 'Oxanium', sans-serif; font-size: 0.62rem; font-weight: 600; letter-spacing: 1px; border: 1px solid #238636; display: inline-block; }
.badge-high { background: var(--orange-soft); color: var(--orange); padding: 3px 10px; border-radius: 20px; font-family: 'Oxanium', sans-serif; font-size: 0.62rem; font-weight: 600; letter-spacing: 1px; border: 1px solid var(--orange-bdr); display: inline-block; }
.badge-medium { background: rgba(210,153,34,0.1); color: #d6a126; padding: 3px 10px; border-radius: 20px; font-family: 'Oxanium', sans-serif; font-size: 0.62rem; font-weight: 600; letter-spacing: 1px; border: 1px solid rgba(210,153,34,0.3); display: inline-block; }
.badge-low { background: rgba(200,60,60,0.1); color: #f85149; padding: 3px 10px; border-radius: 20px; font-family: 'Oxanium', sans-serif; font-size: 0.62rem; font-weight: 600; letter-spacing: 1px; border: 1px solid rgba(200,60,60,0.3); display: inline-block; }

/* Step indicators */
.step-active { background: var(--orange); color: #fff; padding: 10px 18px; border-radius: 8px; text-align: center; font-weight: 700; font-size: 0.88rem; }
.step-done { background: rgba(63,185,80,0.08); color: #3fb950; padding: 10px 18px; border-radius: 8px; text-align: center; font-weight: 600; font-size: 0.88rem; border: 1px solid rgba(63,185,80,0.2); }
.step-idle { background: var(--bg-card); color: #252535; padding: 10px 18px; border-radius: 8px; text-align: center; font-weight: 600; font-size: 0.88rem; border: 1px solid rgba(255,255,255,0.04); }
.step-wrap { text-align: center; }
.step-num-active { background: var(--orange); color: #fff; width: 32px; height: 32px; border-radius: 50%; display: inline-flex; align-items: center; justify-content: center; font-family: 'Oxanium', sans-serif; font-weight: 700; font-size: 0.9rem; margin-bottom: 6px; }
.step-num-done { background: rgba(63,185,80,0.1); color: #3fb950; width: 32px; height: 32px; border-radius: 50%; display: inline-flex; align-items: center; justify-content: center; font-family: 'Oxanium', sans-serif; font-weight: 700; font-size: 0.9rem; margin-bottom: 6px; border: 1px solid rgba(63,185,80,0.25); }
.step-num-idle { background: var(--bg-card); color: #252535; width: 32px; height: 32px; border-radius: 50%; display: inline-flex; align-items: center; justify-content: center; font-family: 'Oxanium', sans-serif; font-weight: 700; font-size: 0.9rem; margin-bottom: 6px; border: 1px solid rgba(255,255,255,0.05); }
.step-title-active { color: var(--t-g1); font-family: 'Oxanium', sans-serif; font-weight: 600; font-size: 0.82rem; letter-spacing: 0.5px; }
.step-title-done { color: #3fb950; font-family: 'Oxanium', sans-serif; font-weight: 600; font-size: 0.82rem; letter-spacing: 0.5px; }
.step-title-idle { color: #252535; font-family: 'Oxanium', sans-serif; font-weight: 600; font-size: 0.82rem; letter-spacing: 0.5px; }
.step-sub { color: #2e2e3e; font-size: 0.72rem; margin-top: 3px; }
.step-connector { color: rgba(255,255,255,0.06); font-size: 1.2rem; align-self: flex-start; padding-top: 14px; }

/* Pipeline */
.pipeline-box { background: var(--bg-card-2); border: 1px solid var(--orange-bdr); border-radius: 10px; padding: 18px 16px; flex: 1; min-width: 0; }
.pipeline-arrow { color: var(--orange); font-size: 1.5rem; align-self: center; flex-shrink: 0; padding: 0 6px; opacity: 0.6; }
.pipeline-tag { background: var(--orange-soft); color: var(--orange); border: 1px solid var(--orange-bdr); border-radius: 4px; padding: 2px 8px; font-family: 'Oxanium', sans-serif; font-size: 0.62rem; font-weight: 600; text-transform: uppercase; letter-spacing: 1.5px; display: inline-block; margin-bottom: 8px; }
.pipeline-output { background: rgba(63,185,80,0.04); border: 1px solid rgba(63,185,80,0.15); border-radius: 4px; padding: 5px 10px; margin-top: 10px; font-size: 0.78rem; color: #3fb950; }

/* Output chips */
.output-chip { background: var(--bg-card); border: 1px solid rgba(255,255,255,0.05); border-radius: 8px; padding: 12px 14px; }
.output-chip-icon { font-size: 1.3rem; margin-bottom: 5px; }
.output-chip-title { color: var(--t-g1); font-weight: 700; font-size: 0.85rem; margin-bottom: 3px; }
.output-chip-desc { color: #3a3a4a; font-size: 0.76rem; line-height: 1.4; }

/* Log box — terminal style */
.log-box { background: #02020a; border: 1px solid rgba(255,255,255,0.05); border-radius: 8px; padding: 14px 16px; font-family: 'Courier New', monospace; font-size: 0.78rem; color: #58a6ff; line-height: 1.6; max-height: 260px; overflow-y: auto; }

/* Caption */
small, .caption { color: #2e2e3e !important; font-size: 0.82rem !important; }

/* Status pulse dot */
@keyframes pulse-dot {
    0%,100% { opacity: 1; }
    50% { opacity: 0.3; }
}
.status-dot { width: 6px; height: 6px; border-radius: 50%; background: #3fb950; display: inline-block; animation: pulse-dot 2s ease-in-out infinite; box-shadow: 0 0 6px rgba(63,185,80,0.5); }
.status-dot-or { background: var(--orange); box-shadow: 0 0 6px rgba(255,106,61,0.5); }

/* Section header — eyebrow + title pattern */
.section-eyebrow { font-family: 'Oxanium', sans-serif; font-size: 0.6rem; font-weight: 600; letter-spacing: 3px; text-transform: uppercase; color: var(--orange); margin-bottom: 4px; display: flex; align-items: center; gap: 8px; }
.section-eyebrow::before { content: ''; display: block; width: 14px; height: 2px; background: var(--orange); }
.section-title { font-family: 'Jost', sans-serif; font-size: 1.05rem; font-weight: 700; color: var(--t-g1); margin: 0 0 14px 0; }

/* Caveat accent */
.caveat-accent { font-family: 'Caveat', cursive; font-size: 1.1rem; color: var(--orange); font-weight: 500; }

/* ── Sidebar ── */
[data-testid="stSidebar"] { background: var(--bg-dark-2) !important; border-right: 1px solid rgba(255,255,255,0.04) !important; }
</style>
""", unsafe_allow_html=True)

# ── Dashboard Header ─────────────────────────────────────────────────────────
st.markdown(_html("""
<div style="padding:1.4rem 0 1rem 0; display:flex; align-items:center; gap:14px; border-bottom:1px solid rgba(255,255,255,0.04); margin-bottom:1.2rem;">
  <div style="width:36px; height:36px; background:rgba(255,106,61,0.1); border:1px solid rgba(255,106,61,0.2); border-radius:8px; display:flex; align-items:center; justify-content:center; flex-shrink:0;">
    <svg width="18" height="18" viewBox="0 0 24 24" fill="none">
      <rect x="2" y="3" width="8" height="8" rx="1.5" fill="#ff6a3d" opacity="0.9"/>
      <rect x="14" y="3" width="8" height="8" rx="1.5" fill="#ff6a3d" opacity="0.4"/>
      <rect x="2" y="13" width="8" height="8" rx="1.5" fill="#ff6a3d" opacity="0.4"/>
      <rect x="14" y="13" width="8" height="8" rx="1.5" fill="#ff6a3d" opacity="0.7"/>
    </svg>
  </div>
  <div style="flex:1;">
    <div style="font-family:'Oxanium',sans-serif; font-size:0.48rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#ff6a3d; margin-bottom:3px;">viAct · AI Content Platform</div>
    <div style="font-family:'Oxanium',sans-serif; font-size:1.1rem; font-weight:700; color:#fff; letter-spacing:-0.3px; line-height:1;">Content Intelligence Suite</div>
  </div>
  <div style="display:flex; align-items:center; gap:16px; padding:8px 16px; background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:7px;">
    <div style="text-align:center;">
      <div style="font-family:'Oxanium',sans-serif; font-size:0.48rem; letter-spacing:2px; text-transform:uppercase; color:#2a2a3a;">Agents</div>
      <div style="font-family:'Oxanium',sans-serif; font-size:1rem; font-weight:700; color:#ff6a3d; line-height:1.2;">08</div>
    </div>
    <div style="width:1px; height:20px; background:rgba(255,255,255,0.05);"></div>
    <div style="display:flex; align-items:center; gap:5px;">
      <span class="status-dot"></span>
      <span style="font-family:'Oxanium',sans-serif; font-size:0.52rem; font-weight:600; color:#3fb950; letter-spacing:1px; text-transform:uppercase;">Live</span>
    </div>
  </div>
</div>
"""), unsafe_allow_html=True)

# ── Agent navigation — horizontal nav bar ────────────────────────────────────
_AGENT_NAV = [
    ("market_radar", "01", "Market Radar",    "#ff6a3d"),
    ("industry",     "02", "Industry Pages",  "#3fb950"),
    ("casestudy",    "03", "Case Studies",    "#58a6ff"),
    ("product",      "04", "Product Pages",   "#d6a126"),
    ("va",           "07", "Video Analytics", "#bc8cff"),
    ("solutions",    "08", "Solutions Pages", "#ff6a3d"),
    ("blog",         "09", "Blog Writer",     "#e879f9"),
    ("meta_seo",     "10", "Meta / SEO",      "#38bdf8"),
    ("partner",      "11", "Partner Outreach","#22d3ee"),
]
if "agent_nav" not in st.session_state:
    st.session_state["agent_nav"] = "market_radar"
_sel = st.session_state["agent_nav"]

_nav_cols = st.columns(9, gap="small")
for (_key, _num, _name, _color), _col in zip(_AGENT_NAV, _nav_cols):
    _active = _sel == _key
    with _col:
        _btn_style = (
            f"background:{_color} !important; color:#fff !important; border-color:{_color} !important;"
            if _active else
            "background:#12121c !important; color:#3a3a4a !important; border-color:rgba(255,255,255,0.06) !important;"
        )
        st.markdown(
            f"<style>.nav-btn-{_key} button {{ {_btn_style} font-family:'Oxanium',sans-serif !important; font-size:0.6rem !important; letter-spacing:1.2px !important; border-radius:7px !important; padding:8px 4px !important; line-height:1.3 !important; }}</style>",
            unsafe_allow_html=True,
        )
        if st.button(f"{_num}\n{_name.upper()}", key=f"nav_{_key}", use_container_width=True):
            st.session_state["agent_nav"] = _key
            st.rerun()

st.markdown("<div style='height:16px;'></div>", unsafe_allow_html=True)

# ── Agent content routing ─────────────────────────────────────────────────────
# (old agent overview cards — replaced by nav bar above)
_DEAD_CODE_CARDS_START = """SKIP_START

  <!-- Agent 01 — Market Radar -->
  <div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:10px; padding:16px 18px; position:relative; overflow:hidden; cursor:default; transition:border-color 0.2s;">
    <div style="position:absolute; top:0; left:0; right:0; height:2px; background:#ff6a3d; opacity:0.8;"></div>
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
      <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#ff6a3d;">AGENT 01</div>
      <div style="font-family:'Oxanium',sans-serif; font-size:1rem; font-weight:700; color:#ff6a3d;">25+</div>
    </div>
    <div style="font-family:'Jost',sans-serif; color:#E9ECF1; font-weight:700; font-size:0.9rem; margin-bottom:3px;">Market Radar</div>
    <div style="font-family:'Oxanium',sans-serif; color:#ff6a3d; font-size:0.55rem; font-weight:600; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:10px;">Daily Intel · Topic Gen</div>
    <p style="margin:0 0 10px 0; color:#2e2e3e; font-size:0.75rem; line-height:1.55;">Monitor 25 competitors and auto-suggest daily content topics — Industry, Case Study, VA, Solutions.</p>
    <div style="display:flex; gap:4px; flex-wrap:wrap; margin-bottom:10px;">
      <span style="background:rgba(255,106,61,0.08); color:#ff6a3d; border:1px solid rgba(255,106,61,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px; letter-spacing:0.5px;">Tavily</span>
      <span style="background:rgba(255,106,61,0.08); color:#ff6a3d; border:1px solid rgba(255,106,61,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px; letter-spacing:0.5px;">RSS</span>
      <span style="background:rgba(255,106,61,0.08); color:#ff6a3d; border:1px solid rgba(255,106,61,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px; letter-spacing:0.5px;">Auto Daily</span>
    </div>
    <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; letter-spacing:1.5px; color:#1e1e2e; text-transform:uppercase;">Tab 1 · Sheet + Email</div>
  </div>

  <!-- Agent 02 — Industry Pages -->
  <div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:10px; padding:16px 18px; position:relative; overflow:hidden; cursor:default;">
    <div style="position:absolute; top:0; left:0; right:0; height:2px; background:#3fb950; opacity:0.8;"></div>
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
      <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#3fb950;">AGENT 02</div>
      <div style="font-family:'Oxanium',sans-serif; font-size:1rem; font-weight:700; color:#3fb950;">8</div>
    </div>
    <div style="font-family:'Jost',sans-serif; color:#E9ECF1; font-weight:700; font-size:0.9rem; margin-bottom:3px;">Industry Pages</div>
    <div style="font-family:'Oxanium',sans-serif; color:#3fb950; font-size:0.55rem; font-weight:600; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:10px;">Dynamic Landing Pages</div>
    <p style="margin:0 0 10px 0; color:#2e2e3e; font-size:0.75rem; line-height:1.55;">Industry + .docx upload &mdash; 8-section Wix CMS page: Hero, Metrics, Use Cases, Testimonials, CTA.</p>
    <div style="display:flex; gap:4px; flex-wrap:wrap; margin-bottom:10px;">
      <span style="background:rgba(63,185,80,0.07); color:#3fb950; border:1px solid rgba(63,185,80,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Construction</span>
      <span style="background:rgba(63,185,80,0.07); color:#3fb950; border:1px solid rgba(63,185,80,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Oil &amp; Gas</span>
      <span style="background:rgba(63,185,80,0.07); color:#3fb950; border:1px solid rgba(63,185,80,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">+6 more</span>
    </div>
    <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; letter-spacing:1.5px; color:#1e1e2e; text-transform:uppercase;">Tab 2 · Dynamic Pages Sheet</div>
  </div>

  <!-- Agent 03 — Case Studies -->
  <div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:10px; padding:16px 18px; position:relative; overflow:hidden; cursor:default;">
    <div style="position:absolute; top:0; left:0; right:0; height:2px; background:#58a6ff; opacity:0.8;"></div>
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
      <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#58a6ff;">AGENT 03</div>
      <div style="font-family:'Oxanium',sans-serif; font-size:1rem; font-weight:700; color:#58a6ff;">56</div>
    </div>
    <div style="font-family:'Jost',sans-serif; color:#E9ECF1; font-weight:700; font-size:0.9rem; margin-bottom:3px;">Case Studies</div>
    <div style="font-family:'Oxanium',sans-serif; color:#58a6ff; font-size:0.55rem; font-weight:600; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:10px;">56-Field CMS Schema</div>
    <p style="margin:0 0 10px 0; color:#2e2e3e; font-size:0.75rem; line-height:1.55;">Enter project details &mdash; AI generates: Problem, Solution, Impact, Testimonials, Metrics + 56 Wix fields.</p>
    <div style="display:flex; gap:4px; flex-wrap:wrap; margin-bottom:10px;">
      <span style="background:rgba(88,166,255,0.07); color:#58a6ff; border:1px solid rgba(88,166,255,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Tavily</span>
      <span style="background:rgba(88,166,255,0.07); color:#58a6ff; border:1px solid rgba(88,166,255,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">.docx</span>
      <span style="background:rgba(88,166,255,0.07); color:#58a6ff; border:1px solid rgba(88,166,255,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Image Briefs</span>
    </div>
    <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; letter-spacing:1.5px; color:#1e1e2e; text-transform:uppercase;">Tab 3 · Dynamic Pages Sheet</div>
  </div>

</div>

<div style="display:grid; grid-template-columns:1fr 1fr 1fr; gap:10px; margin-bottom:1.6rem;">

  <!-- Agent 04 — Product Pages -->
  <div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:10px; padding:16px 18px; position:relative; overflow:hidden; cursor:default;">
    <div style="position:absolute; top:0; left:0; right:0; height:2px; background:#d6a126; opacity:0.8;"></div>
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
      <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#d6a126;">AGENT 04</div>
      <div style="font-family:'Oxanium',sans-serif; font-size:1rem; font-weight:700; color:#d6a126;">CMS</div>
    </div>
    <div style="font-family:'Jost',sans-serif; color:#E9ECF1; font-weight:700; font-size:0.9rem; margin-bottom:3px;">Product Pages</div>
    <div style="font-family:'Oxanium',sans-serif; color:#d6a126; font-size:0.55rem; font-weight:600; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:10px;">viAct Product CMS</div>
    <p style="margin:0 0 10px 0; color:#2e2e3e; font-size:0.75rem; line-height:1.55;">CMS content for viAct product pages &mdash; features, specs, use cases, SEO — Wix ready.</p>
    <div style="display:flex; gap:4px; flex-wrap:wrap; margin-bottom:10px;">
      <span style="background:rgba(214,161,38,0.07); color:#d6a126; border:1px solid rgba(214,161,38,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">viAct Products</span>
      <span style="background:rgba(214,161,38,0.07); color:#d6a126; border:1px solid rgba(214,161,38,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Wix CMS</span>
    </div>
    <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; letter-spacing:1.5px; color:#1e1e2e; text-transform:uppercase;">Tab 4 · Dynamic Pages Sheet</div>
  </div>

  <!-- Agent 07 — Video Analytics -->
  <div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:10px; padding:16px 18px; position:relative; overflow:hidden; cursor:default;">
    <div style="position:absolute; top:0; left:0; right:0; height:2px; background:#bc8cff; opacity:0.8;"></div>
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
      <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#bc8cff;">AGENT 07</div>
      <div style="font-family:'Oxanium',sans-serif; font-size:1rem; font-weight:700; color:#bc8cff;">27</div>
    </div>
    <div style="font-family:'Jost',sans-serif; color:#E9ECF1; font-weight:700; font-size:0.9rem; margin-bottom:3px;">Video Analytics</div>
    <div style="font-family:'Oxanium',sans-serif; color:#bc8cff; font-size:0.55rem; font-weight:600; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:10px;">27 Detection Pages</div>
    <p style="margin:0 0 10px 0; color:#2e2e3e; font-size:0.75rem; line-height:1.55;">Enter a detection type (PPE, Intrusion, Fall…) &mdash; full page: Hero, Challenges, How It Works, SEO.</p>
    <div style="display:flex; gap:4px; flex-wrap:wrap; margin-bottom:10px;">
      <span style="background:rgba(188,140,255,0.07); color:#bc8cff; border:1px solid rgba(188,140,255,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">27 Types</span>
      <span style="background:rgba(188,140,255,0.07); color:#bc8cff; border:1px solid rgba(188,140,255,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Tavily</span>
      <span style="background:rgba(188,140,255,0.07); color:#bc8cff; border:1px solid rgba(188,140,255,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Alt Texts</span>
    </div>
    <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; letter-spacing:1.5px; color:#1e1e2e; text-transform:uppercase;">Tab 5 · Dynamic Pages Sheet</div>
  </div>

  <!-- Agent 08 — Solutions Pages -->
  <div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:10px; padding:16px 18px; position:relative; overflow:hidden; cursor:default;">
    <div style="position:absolute; top:0; left:0; right:0; height:2px; background:#ff6a3d; opacity:0.5;"></div>
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
      <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#ff6a3d;">AGENT 08</div>
      <div style="font-family:'Oxanium',sans-serif; font-size:1rem; font-weight:700; color:#ff6a3d;">14</div>
    </div>
    <div style="font-family:'Jost',sans-serif; color:#E9ECF1; font-weight:700; font-size:0.9rem; margin-bottom:3px;">Solutions Pages</div>
    <div style="font-family:'Oxanium',sans-serif; color:#ff6a3d; font-size:0.55rem; font-weight:600; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:10px;">14 Solution Verticals</div>
    <p style="margin:0 0 10px 0; color:#2e2e3e; font-size:0.75rem; line-height:1.55;">Full CMS page for 14 viAct solutions — Hero, Dashboard, 5 Features, FAQs, SEO + image prompts.</p>
    <div style="display:flex; gap:4px; flex-wrap:wrap; margin-bottom:10px;">
      <span style="background:rgba(255,106,61,0.08); color:#ff6a3d; border:1px solid rgba(255,106,61,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">Nano Banana</span>
      <span style="background:rgba(255,106,61,0.08); color:#ff6a3d; border:1px solid rgba(255,106,61,0.18); border-radius:20px; font-family:'Oxanium',sans-serif; font-size:0.55rem; font-weight:600; padding:2px 7px;">7 Img Prompts</span>
    </div>
    <div style="font-family:'Oxanium',sans-serif; font-size:0.52rem; letter-spacing:1.5px; color:#1e1e2e; text-transform:uppercase;">Tab 6 · Dynamic Pages Sheet</div>
  </div>

SKIP_END"""

# =============================================================================
# AGENT CONTENT — routed via sidebar navigation
# =============================================================================
if _sel == "market_radar":
    # 3-AGENT MARKET RADAR → WEBPAGE PIPELINE (Tavily + Firecrawl + Groq)
    # =============================================================================

    if "r3_step" not in st.session_state:
        st.session_state["r3_step"] = 0

    step = st.session_state["r3_step"]

    if step == 0:
        st.markdown(_html("""
        <div style="display:flex;gap:8px;margin-bottom:20px;align-items:stretch;">
          <div style="flex:1;background:#12121c;border:1px solid rgba(255,255,255,0.05);border-top:2px solid #ff6a3d;border-radius:8px;padding:14px 16px;">
            <div style="font-family:'Oxanium',sans-serif;font-size:0.55rem;font-weight:600;letter-spacing:2px;text-transform:uppercase;color:#ff6a3d;margin-bottom:8px;">STEP 01</div>
            <div style="color:#E9ECF1;font-family:'Jost',sans-serif;font-weight:700;font-size:0.82rem;margin-bottom:6px;">Run Daily Scan</div>
            <div style="color:#2e2e3e;font-size:0.74rem;line-height:1.55;">Click "Run Daily Intel Scan" to scan 25 competitors, surface trends, and auto-suggest 4 new content topics.</div>
          </div>
          <div style="color:#1e1e2e;font-size:1.1rem;align-self:center;padding:0 4px;">&#8594;</div>
          <div style="flex:1;background:#12121c;border:1px solid rgba(255,255,255,0.05);border-top:2px solid #ff6a3d;border-radius:8px;padding:14px 16px;">
            <div style="font-family:'Oxanium',sans-serif;font-size:0.55rem;font-weight:600;letter-spacing:2px;text-transform:uppercase;color:#ff6a3d;margin-bottom:8px;">STEP 02</div>
            <div style="color:#E9ECF1;font-family:'Jost',sans-serif;font-weight:700;font-size:0.82rem;margin-bottom:6px;">Use a Suggested Topic</div>
            <div style="color:#2e2e3e;font-size:0.74rem;line-height:1.55;">After the scan, suggested topics appear below. Click "Use This" — the topic auto-fills in the target agent tab.</div>
          </div>
          <div style="color:#1e1e2e;font-size:1.1rem;align-self:center;padding:0 4px;">&#8594;</div>
          <div style="flex:1;background:#12121c;border:1px solid rgba(255,255,255,0.05);border-top:2px solid #ff6a3d;border-radius:8px;padding:14px 16px;">
            <div style="font-family:'Oxanium',sans-serif;font-size:0.55rem;font-weight:600;letter-spacing:2px;text-transform:uppercase;color:#ff6a3d;margin-bottom:8px;">STEP 03</div>
            <div style="color:#E9ECF1;font-family:'Jost',sans-serif;font-weight:700;font-size:0.82rem;margin-bottom:6px;">Or Build a Pillar Page</div>
            <div style="color:#2e2e3e;font-size:0.74rem;line-height:1.55;">Use "Run Deep Market Radar" to research a specific SEO gap and generate a full pillar page + 3 supporting blogs.</div>
          </div>
        </div>
        """), unsafe_allow_html=True)

    # ── Step indicator ────────────────────────────────────────────────────────────
    _step_data = [
        ("1", "Competitor Research",   "Scan 8 sites · find content gaps"),
        ("2", "Pick Your Topic",       "Choose gap · add reference material"),
        ("3", "Content Ready",         "Review 10 sections · push to Sheets"),
    ]
    _si_cols = st.columns([1, 0.08, 1, 0.08, 1])
    _col_idxs = [0, 2, 4]
    for _si, ((_num, _title, _sub), _ci) in enumerate(zip(_step_data, _col_idxs)):
        if _si < step:
            _num_cls, _title_cls, _check = "step-num-done", "step-title-done", "&#10003;"
        elif _si == step:
            _num_cls, _title_cls, _check = "step-num-active", "step-title-active", _num
        else:
            _num_cls, _title_cls, _check = "step-num-idle", "step-title-idle", _num
        _si_cols[_ci].markdown(_html(f"""
    <div class="step-wrap">
    <div class="{_num_cls}">{_check}</div>
    <div class="{_title_cls}">{_title}</div>
    <div class="step-sub">{_sub}</div>
    </div>
    """), unsafe_allow_html=True)
        if _ci < 4:
            _si_cols[_ci + 1].markdown(
                f"<div style='text-align:center; color:#2d303a; font-size:1.4rem; padding-top:8px;'>&#8594;</div>",
                unsafe_allow_html=True,
            )

    if step > 0:
        st.write("")
        if st.button("↩ Start Over", key="r3_reset"):
            for k in [k for k in st.session_state if k.startswith("r3_")]:
                del st.session_state[k]
            st.rerun()

    st.write("")

    # =============================================================================
    # STEP 0 — API Key Check + Run Market Radar
    # =============================================================================
    if step == 0:
        # ── Pipeline flow diagram ──────────────────────────────────────────────────
        st.markdown(
            "<p style='font-family:\"Oxanium\",sans-serif; color:#2e2e3e; font-size:0.6rem; font-weight:600; text-transform:uppercase; letter-spacing:3px; margin-bottom:12px; display:flex; align-items:center; gap:8px;'><span style='display:inline-block;width:14px;height:2px;background:#ff6a3d;'></span>HOW IT WORKS &mdash; 3 AI AGENTS IN SEQUENCE</p>",
            unsafe_allow_html=True,
        )
        st.markdown(_html("""
    <div style="display:flex; gap:0; align-items:stretch; margin-bottom:20px;">

    <div class="pipeline-box" style="border-color:rgba(255,106,61,0.35);">
    <div class="pipeline-tag">Agent 1 &middot; Tavily Search</div>
    <div style="color:#e6edf3; font-weight:700; font-size:1rem; margin-bottom:6px;">&#128269; Find the Gap</div>
    <div style="color:#8b949e; font-size:0.82rem; line-height:1.6;">
    &#8226; Searches <strong style="color:#c9d1d9;">8 competitor websites</strong> for safety topics<br>
    &#8226; AI extracts 10&ndash;15 topic names from results<br>
    &#8226; Checks each topic on <code>viact.ai</code> &#8212; live<br>
    &#8226; Only topics with <strong style="color:#3fb950;">0 viAct pages</strong> pass through
    </div>
    <div class="pipeline-output">&#8594; Output: 3 Confirmed Content Gaps</div>
    </div>

    <div class="pipeline-arrow">&#9654;</div>

    <div class="pipeline-box" style="border-color:rgba(88,166,255,0.3);">
    <div class="pipeline-tag" style="background:rgba(88,166,255,0.1); color:#58a6ff; border-color:rgba(88,166,255,0.3);">Agent 2 &middot; Firecrawl</div>
    <div style="color:#e6edf3; font-weight:700; font-size:1rem; margin-bottom:6px;">&#128196; Read Competitor Pages</div>
    <div style="color:#8b949e; font-size:0.82rem; line-height:1.6;">
    &#8226; Downloads full competitor page content<br>
    &#8226; Bypasses anti-bot protection automatically<br>
    &#8226; Converts messy HTML to clean text<br>
    &#8226; Blocked pages marked <strong style="color:#f85149;">[ACCESS DENIED]</strong> &#8212; never faked
    </div>
    <div class="pipeline-output" style="border-color:rgba(88,166,255,0.2); color:#58a6ff;">&#8594; Output: Real Competitor Markdown</div>
    </div>

    <div class="pipeline-arrow">&#9654;</div>

    <div class="pipeline-box" style="border-color:rgba(63,185,80,0.3);">
    <div class="pipeline-tag" style="background:rgba(63,185,80,0.08); color:#3fb950; border-color:rgba(63,185,80,0.3);">Agent 3 &middot; Llama 3.3 70B</div>
    <div style="color:#e6edf3; font-weight:700; font-size:1rem; margin-bottom:6px;">&#9999; Write the Webpage</div>
    <div style="color:#8b949e; font-size:0.82rem; line-height:1.6;">
    &#8226; Uses <strong style="color:#c9d1d9;">only</strong> the scraped content &#8212; zero hallucination<br>
    &#8226; Writes headline, body, FAQs, meta tags<br>
    &#8226; Adds Singapore &amp; UAE regulatory context<br>
    &#8226; Generates schema markup + image prompts
    </div>
    <div class="pipeline-output" style="border-color:rgba(63,185,80,0.2); color:#3fb950;">&#8594; Output: 10-Section Publish-Ready Page</div>
    </div>

    </div>
    """), unsafe_allow_html=True)

        # ── What you'll get ────────────────────────────────────────────────────────
        st.markdown(
            "<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:12px;'>WHAT YOU GET &mdash; 10 SECTIONS, READY TO PUBLISH</p>",
            unsafe_allow_html=True,
        )
        _outputs = [
            ("&#128196;", "Full Webpage Body",    "Complete Markdown page — paste into any CMS"),
            ("&#127919;", "H1 Headline",          "Problem-focused, no marketing fluff"),
            ("&#128269;", "SEO Meta Tags",        "Title &lt;60 chars, description &lt;155 chars"),
            ("&#10067;",  "5 Schema FAQs",        "40-60 word answers for Google rich results"),
            ("&#127991;", "JSON-LD Schema",       "Paste-ready &lt;script&gt; for structured data"),
            ("&#127760;", "GEO Package",          "Opening 200 words for ChatGPT/Perplexity citations"),
            ("&#128247;", "2 Image Prompts",      "Realistic APAC photography briefs, not CGI"),
            ("&#128279;", "Internal Links",       "Verified viact.ai URLs only — never invented"),
            ("&#128202;", "Decision Logic",       "Email paragraph with proof for Gary/Surendra"),
            ("&#128220;", "Keyword Set",          "Primary + LSI keywords + heading map"),
        ]
        _out_cols = st.columns(5)
        for _oi, (_icon, _title, _desc) in enumerate(_outputs):
            _out_cols[_oi % 5].markdown(_html(f"""
    <div class="output-chip">
    <div class="output-chip-icon">{_icon}</div>
    <div class="output-chip-title">{_title}</div>
    <div class="output-chip-desc">{_desc}</div>
    </div>
    """), unsafe_allow_html=True)

        st.write("")

        # ── Competitor grid (static — no API calls) ────────────────────────────────
        from research_competitors import get_all_competitors
        _all_competitors = get_all_competitors()

        st.markdown(_html(f"""
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
    <p style="margin:0; color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px;">
    COMPETITOR WEBSITES AGENT 1 WILL SCAN &nbsp;<span style="color:#ff6a3d;">({len(_all_competitors)} sites)</span>
    </p>
    <span style="color:#8b949e; font-size:0.76rem;">All searched simultaneously via Tavily</span>
    </div>
    """), unsafe_allow_html=True)
        _comp_cols = st.columns(4)
        for _ci, _comp in enumerate(_all_competitors):
            _domain = _comp["url"].split("//")[-1].split("/")[0]
            _comp_cols[_ci % 4].markdown(
                _html(f"""
    <div class="glass-card" style="padding:12px 14px;">
    <div style="color:#e6edf3; font-weight:700; font-size:0.88rem; margin-bottom:3px;">{_t(_comp['name'])}</div>
    <div style="color:#8b949e; font-size:0.74rem; font-family:monospace;">{_t(_domain)}</div>
    </div>
    """), unsafe_allow_html=True)

        st.write("")

        # ── API Key Status ─────────────────────────────────────────────────────────
        st.markdown(_html("""
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:10px;">
    <p style="margin:0; color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px;">API CONNECTIONS</p>
    <span style="color:#8b949e; font-size:0.76rem;">All 3 agents need their own API key to work</span>
    </div>
    """), unsafe_allow_html=True)

        key_configs = [
            ("TAVILY_API_KEY",    "Tavily",    "&#128269;", "Searches competitor websites &amp; confirms gaps on viact.ai",  "Agent 1 — Required", True),
            ("FIRECRAWL_API_KEY", "Firecrawl", "&#128293;", "Downloads full competitor page content bypassing anti-bot",     "Agent 2 — Optional", False),
            ("GROQ_API_KEY",      "Groq",      "&#129302;", "Runs Llama 3.3 70B to extract topics and write page content",   "Agent 1 &amp; 3 — Required", True),
        ]

        key_cols = st.columns(3)
        all_required_present = True
        for col, (key_name, label, icon, desc, agent_label, required) in zip(key_cols, key_configs):
            val = os.getenv(key_name, "")
            present = bool(val)
            if required and not present:
                all_required_present = False
            status_color = "#3fb950" if present else ("#f85149" if required else "#d6a126")
            status_text  = "Connected" if present else ("Missing" if required else "Optional")
            masked = f"{val[:8]}…" if present else "Not set"
            col.markdown(_html(f"""
    <div class="glass-card" style="padding:16px; border-color:rgba({'63,185,80' if present else ('248,81,73' if required else '210,153,34')},0.25);">
    <div style="display:flex; justify-content:space-between; align-items:flex-start; margin-bottom:8px;">
    <div style="display:flex; align-items:center; gap:8px;">
    <span style="font-size:1.3rem;">{icon}</span>
    <span style="color:#e6edf3; font-weight:700; font-size:0.98rem;">{_t(label)}</span>
    </div>
    <span style="background:rgba({'63,185,80' if present else ('248,81,73' if required else '210,153,34')},0.1); color:{status_color}; border:1px solid rgba({'63,185,80' if present else ('248,81,73' if required else '210,153,34')},0.3); border-radius:20px; padding:2px 8px; font-size:0.72rem; font-weight:700;">{status_text}</span>
    </div>
    <div style="color:#8b949e; font-size:0.78rem; line-height:1.5; margin-bottom:6px;">{desc}</div>
    <div style="display:flex; justify-content:space-between; align-items:center;">
    <span style="color:#ff6a3d; font-size:0.72rem; font-weight:600;">{agent_label}</span>
    <span style="color:{status_color}; font-size:0.74rem; font-family:monospace;">{_t(masked)}</span>
    </div>
    </div>
    """), unsafe_allow_html=True)

        if not all_required_present:
            st.markdown(_html("""
    <div style="background:rgba(248,81,73,0.06); border:1px solid rgba(248,81,73,0.3); border-radius:8px; padding:12px 16px; font-size:0.84rem; color:#f85149; margin-bottom:4px;">
    <strong>&#9888; Missing required API keys.</strong> Add them to your <code>.env</code> file or Streamlit Cloud secrets:
    <br><code style="color:#c9d1d9;">GROQ_API_KEY=gsk_...</code> &nbsp;&middot;&nbsp; <code style="color:#c9d1d9;">TAVILY_API_KEY=tvly-...</code>
    </div>
    """), unsafe_allow_html=True)

        st.write("")

        # ── Industry Vertical Selector ─────────────────────────────────────────────
        INDUSTRY_OPTIONS = {
            "Construction Safety":     "construction safety",
            "Oil & Gas Safety":        "oil gas safety",
            "Manufacturing Safety":    "manufacturing safety",
            "Mining Safety":           "mining safety",
            "Facility Management":     "facility management safety",
            "Food & Beverage Safety":  "food beverage safety",
        }
        _ind_label = st.selectbox(
            "🏭  Industry Vertical — which sector should Agent 1 search?",
            options=list(INDUSTRY_OPTIONS.keys()),
            index=0,
            key="r3_industry_label",
            help="Agent 1 will search competitor sites using this industry keyword.",
        )
        st.session_state["r3_industry"] = INDUSTRY_OPTIONS[_ind_label]

        st.markdown(_html("""
    <div style="background:rgba(255,106,61,0.04); border:1px solid rgba(255,106,61,0.15); border-radius:8px; padding:12px 16px; margin-bottom:12px; display:flex; align-items:center; gap:12px;">
    <span style="font-size:1.4rem;">&#9201;</span>
    <div>
    <div style="color:#e6edf3; font-weight:600; font-size:0.88rem;">Takes about 2&ndash;3 minutes</div>
    <div style="color:#8b949e; font-size:0.78rem;">Agent 1 runs ~12 live searches (8 competitors + 10-15 viact.ai checks). No action needed &mdash; just watch the progress below.</div>
    </div>
    </div>
    """), unsafe_allow_html=True)
        _col_scan, _col_radar = st.columns([2, 3])
        with _col_scan:
            run_intel = st.button(
                "📡 Run Daily Intel Scan",
                type="primary",
                use_container_width=True,
                key="r3_intel",
                help="Scan competitor news + trends + generate today's 3 content topics (~30 sec)",
                disabled=not all_required_present,
            )
        with _col_radar:
            run_radar = st.button(
                "🔍 Run Deep Market Radar",
                use_container_width=True,
                key="r3_run",
                help="Deep competitor content gap scan — takes 2-3 minutes",
                disabled=not all_required_present,
            )

        if run_intel:
            from competitor_news_monitor import run_daily_monitor
            from push_to_sheets import push_competitor_intel, push_daily_topics, push_product_launches, push_competitor_site_changes
            with st.spinner("Scanning competitors, trends & topics... ~45 sec"):
                _intel_result = run_daily_monitor(progress_callback=lambda m: st.toast(m))
            push_competitor_intel(_intel_result)
            push_daily_topics(_intel_result)
            push_product_launches(_intel_result)
            push_competitor_site_changes(_intel_result)

            # Save suggested topics to session state for other tabs to pick up
            _dt = _intel_result.get("daily_topics", {})
            if _dt.get("industry_topic"):
                st.session_state["suggested_industry_topic"] = _dt["industry_topic"]
            if _dt.get("case_study_topic"):
                st.session_state["suggested_cs_topic"] = _dt["case_study_topic"]
            if _dt.get("va_topic"):
                st.session_state["suggested_va_topic"] = _dt["va_topic"]
            if _dt.get("solutions_topic"):
                st.session_state["suggested_solutions_topic"] = _dt["solutions_topic"]

            # ── Today's 3 Topics card ─────────────────────────────────────────
            _dt = _intel_result.get("daily_topics", {})
            _ind_t = _dt.get("industry_topic", {})
            _cs_t  = _dt.get("case_study_topic", {})
            _va_t  = _dt.get("va_topic", {})
            st.markdown("### 💡 Today's 3 Content Topics")
            _tc1, _tc2, _tc3 = st.columns(3)
            with _tc1:
                st.markdown(f"**🏭 Industry Page**")
                st.markdown(f"_{_ind_t.get('industry', '')}_")
                st.markdown(f"**{_ind_t.get('topic', '—')}**")
                st.caption(_ind_t.get("why", ""))
            with _tc2:
                st.markdown(f"**📋 Case Study**")
                st.markdown(f"_{_cs_t.get('company_type', '')}, {_cs_t.get('location', '')}_")
                st.markdown(f"**{_cs_t.get('detection_focus', '—')}**")
                st.caption(_cs_t.get("why", ""))
            with _tc3:
                st.markdown(f"**🎯 Video Analytics**")
                st.markdown(f"**{_va_t.get('detection_name', '—')}**")
                st.caption(_va_t.get("why", ""))
            st.success(f"✅ Scan done — {_intel_result['counts']['competitor_news']} competitor news · {_intel_result['counts']['trends']} trends · {_intel_result['counts'].get('product_launches', 0)} launches · {_intel_result['counts'].get('website_changes', 0)} site changes · topics saved to sheet + session")
            st.divider()

            # ── Competitor Product Launches ───────────────────────────────────
            _launches = _intel_result.get("product_launches", [])
            if _launches:
                st.markdown("### 🚀 Competitor Product Launches Detected")
                for _launch in _launches:
                    _lc1, _lc2 = st.columns([4, 1])
                    with _lc1:
                        st.warning(
                            f"**{_launch['competitor']}** — **{_launch['product_name']}**\n\n"
                            f"{_launch.get('snippet', '')[:150]}..."
                        )
                    with _lc2:
                        if st.button("Respond with Page →", key=f"respond_{abs(hash(_launch['url']))}"):
                            st.session_state["product_prefill_url"]  = _launch["url"]
                            st.session_state["product_prefill_name"] = _launch["product_name"]
                            st.rerun()
                st.divider()

            # ── Competitor Website Changes ────────────────────────────────────
            _wchanges = _intel_result.get("website_changes", [])
            if _wchanges:
                st.markdown("### 🔍 Competitor Website Changes")
                for _wc in _wchanges:
                    _wc1, _wc2 = st.columns([4, 1])
                    with _wc1:
                        _wc_badge = "🆕 New Page" if _wc["change_type"] == "new_page" else "✏️ Updated"
                        _wc_title = _wc.get("title") or _wc.get("url", "")
                        st.info(
                            f"{_wc_badge} **{_wc['competitor']}** — "
                            f"[{_wc_title[:80]}]({_wc['url']})\n\n"
                            f"💡 *{_wc['marketing_response']}*"
                        )
                    with _wc2:
                        if st.button("Build Counter-Page →", key=f"counter_{abs(hash(_wc['url']))}"):
                            st.session_state["suggested_industry_topic"] = {
                                "industry": "Construction Safety",
                                "topic": f"viAct vs {_wc['competitor']}: {_wc_title[:60]}",
                                "why": _wc["marketing_response"],
                            }
                            st.rerun()
                st.divider()

            # ── Intel details ─────────────────────────────────────────────────
            _urgency = _intel_result.get("urgency", "medium").upper()
            _urgency_icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(_urgency, "🟡")
            st.markdown(f"### {_urgency_icon} Urgency: **{_urgency}**")
            with st.expander("📋 Executive Summary", expanded=True):
                st.markdown(f"**{_intel_result['executive_summary']}**")
                st.markdown(f"🔥 **Trending:** {_intel_result['trending_topic']}")
                st.markdown(f"🏢 **Top Competitor Move:** {_intel_result['top_competitor_move']}")
                st.markdown(f"⚡ **viAct Action:** {_intel_result['viact_opportunity']}")
            with st.expander(f"🏢 Competitor News ({_intel_result['counts']['competitor_news']})"):
                for _cn in _intel_result.get("competitor_news", []):
                    st.markdown(f"**[{_cn['title']}]({_cn['url']})**")
                    st.caption(f"{_cn['competitor']} — {_cn['snippet'][:120]}")
                    st.divider()
            with st.expander(f"📈 Industry Trends ({_intel_result['counts']['trends']})"):
                for _tr in _intel_result.get("industry_trends", []):
                    st.markdown(f"**[{_tr['title']}]({_tr['url']})**")
                    st.caption(_tr['snippet'][:120])
                    st.divider()
            with st.expander(f"💡 Marketing Opportunities ({_intel_result['counts']['opportunities']})"):
                for _op in _intel_result.get("marketing_opportunities", []):
                    st.markdown(f"**[{_op['title']}]({_op['url']})**")
                    st.caption(_op['snippet'][:120])
                    st.divider()

        if run_radar:
            from agent1_market_explorer import discover_market_gaps

            st.markdown("<hr/>", unsafe_allow_html=True)

            # ── Pre-flight: quick Tavily key validation ────────────────────────────
            _tavily_key = os.getenv("TAVILY_API_KEY", "")
            _tavily_ok = False
            try:
                import requests as _rq
                _r = _rq.post("https://api.tavily.com/search", json={
                    "api_key": _tavily_key, "query": "test", "max_results": 1
                }, timeout=12)
                _tavily_ok = _r.status_code == 200
                if not _tavily_ok:
                    _err_body = _r.json() if _r.headers.get("content-type","").startswith("application/json") else {}
                    _err_code = _r.status_code
                    _err_msg = _err_body.get("message") or _err_body.get("detail") or f"HTTP {_err_code}"
                    if _err_code in (429, 432):
                        st.warning(
                            f"**Tavily monthly limit reached** ({_err_msg}). "
                            "Search will use Google News RSS as fallback — results may vary slightly."
                        )
                        _tavily_ok = True  # allow pipeline to continue; agents handle fallback internally
                    else:
                        st.error(
                            f"**Tavily API key failed** ({_err_msg}).  "
                            f"Key in use: `{_tavily_key[:20]}...`  \n\n"
                            "**Fix:** Get a new free key at [tavily.com](https://tavily.com) "
                            "and update it in `.env` (local) or Streamlit Cloud → Settings → Secrets."
                        )
            except Exception as _te:
                st.error(f"Tavily connection error: {_te}")

            if not _tavily_ok:
                st.stop()

            # ── 3-phase progress panel (reset on every new run) ───────────────────
            st.session_state["r3_progress"] = {"competitors": [], "topics": [], "gaps": [], "logs": []}

            progress_placeholder = st.empty()

            def _render_progress():
                prog = st.session_state["r3_progress"]
                comp_items  = prog["competitors"]   # "Name|N" strings
                topic_items = prog["topics"]         # plain topic name strings
                gap_items   = prog["gaps"]           # "CONFIRMED|name|score" or "SKIP|name|url"
                log_items   = prog.get("logs", [])   # general agent1 messages

                n_comp  = len(comp_items)
                n_topic = len(topic_items)
                n_confirmed = sum(1 for g in gap_items if g.startswith("CONFIRMED"))
                n_gap   = len(gap_items)

                # Phase A rows
                comp_rows = "".join(
                    f"<div style='display:flex; justify-content:space-between; padding:4px 0; border-bottom:1px solid #1f2430;'>"
                    f"<span style='color:#c9d1d9; font-size:0.82rem;'>{_t(item.split('|')[0])}</span>"
                    f"<span style='color:#8b949e; font-size:0.78rem;'>{_t(item.split('|')[1])} snippet(s)</span>"
                    f"</div>"
                    for item in comp_items
                )

                # Phase B rows
                topic_rows = "".join(
                    f"<div style='padding:3px 0; color:#c9d1d9; font-size:0.82rem; border-bottom:1px solid #1f2430;'>&#8250; {_t(t)}</div>"
                    for t in topic_items
                )

                # Phase C rows
                gap_rows = "".join(
                    (
                        f"<div style='padding:4px 0; border-bottom:1px solid #1f2430;'>"
                        f"<span style='background:rgba(255,106,61,0.12); color:#ff6a3d; border:1px solid rgba(255,106,61,0.3); border-radius:4px; padding:1px 7px; font-size:0.74rem; font-weight:700; margin-right:6px;'>&#10003; CONFIRMED</span>"
                        f"<span style='color:#e6edf3; font-size:0.82rem;'>{_t(parts[1] if len(parts) > 1 else '')}</span>"
                        f"<span style='color:#ff6a3d; font-size:0.74rem; margin-left:6px;'>{_t(parts[2] if len(parts) > 2 else '')}</span>"
                        f"</div>"
                        if g.startswith("CONFIRMED") else
                        f"<div style='padding:4px 0; border-bottom:1px solid #1f2430;'>"
                        f"<span style='color:#484f58; font-size:0.78rem; margin-right:6px;'>&#8627; skip</span>"
                        f"<span style='color:#484f58; font-size:0.8rem;'>{_t(parts[1] if len(parts) > 1 else '')}</span>"
                        f"</div>"
                    )
                    for g in gap_items
                    for parts in [g.split("|")]
                )

                html_out = (
                    f"<div style='background:#0d1117; border:1px solid #2d303a; border-radius:10px; overflow:hidden; margin-top:8px;'>"

                    # Phase A header
                    f"<div style='padding:10px 14px; background:rgba(255,106,61,0.06); border-bottom:1px solid #2d303a; display:flex; justify-content:space-between;'>"
                    f"<span style='color:#ff6a3d; font-weight:700; font-size:0.82rem; text-transform:uppercase; letter-spacing:1px;'>&#128225; Competitors Scanned</span>"
                    f"<span style='color:#8b949e; font-size:0.8rem;'>{n_comp} / {len(_all_competitors)}</span>"
                    f"</div>"
                    + (f"<div style='padding:8px 14px;'>{comp_rows}</div>" if comp_rows else "")

                    # Phase B header
                    + f"<div style='padding:10px 14px; background:rgba(88,166,255,0.05); border-top:1px solid #2d303a; border-bottom:1px solid #2d303a; display:flex; justify-content:space-between;'>"
                    f"<span style='color:#58a6ff; font-weight:700; font-size:0.82rem; text-transform:uppercase; letter-spacing:1px;'>&#128269; Topics Extracted</span>"
                    f"<span style='color:#8b949e; font-size:0.8rem;'>{n_topic} topic(s)</span>"
                    f"</div>"
                    + (f"<div style='padding:8px 14px;'>{topic_rows}</div>" if topic_rows else "")

                    # Phase C header
                    + f"<div style='padding:10px 14px; background:rgba(63,185,80,0.05); border-top:1px solid #2d303a; border-bottom:1px solid #2d303a; display:flex; justify-content:space-between;'>"
                    f"<span style='color:#3fb950; font-weight:700; font-size:0.82rem; text-transform:uppercase; letter-spacing:1px;'>&#10003; Gap Confirmation</span>"
                    f"<span style='color:#8b949e; font-size:0.8rem;'>{n_confirmed} confirmed &middot; {n_gap} checked of {n_topic}</span>"
                    f"</div>"
                    + (f"<div style='padding:8px 14px;'>{gap_rows}</div>" if gap_rows else "")

                    # Log panel — shows errors, dedup notices, viact.ai check details
                    + (
                        f"<div style='padding:10px 14px; background:rgba(139,148,158,0.04); border-top:1px solid #2d303a;'>"
                        f"<div style='color:#8b949e; font-size:0.74rem; font-weight:700; text-transform:uppercase; letter-spacing:1px; margin-bottom:6px;'>Agent Log</div>"
                        + "".join(
                            f"<div style='padding:2px 0; color:{'#f85149' if any(w in m.lower() for w in ('failed','error','warn','missing')) else '#484f58'}; font-size:0.76rem; font-family:monospace;'>{_t(m)}</div>"
                            for m in log_items[-20:]
                        )
                        + "</div>"
                        if log_items else ""
                    )

                    + "</div>"
                )
                progress_placeholder.markdown(html_out, unsafe_allow_html=True)

            def _ui_progress(phase: str, message: str):
                if phase in ("competitors", "topics", "gaps"):
                    st.session_state["r3_progress"][phase].append(message)
                else:
                    # agent1 general log — capture errors, dedup notices, etc.
                    st.session_state["r3_progress"]["logs"].append(message)
                _render_progress()

            with st.spinner("Agent 1 running Tavily searches and confirming gaps..."):
                try:
                    radar_results = discover_market_gaps(
                        progress_callback=_ui_progress,
                        industry=st.session_state.get("r3_industry", "construction safety"),
                    )
                    if not radar_results.get("topics"):
                        _logs = st.session_state.get("r3_progress", {}).get("logs", [])
                        _last_logs = "  \n".join(f"`{m}`" for m in _logs[-5:]) if _logs else ""
                        st.warning(
                            "**No confirmed gaps found this run.**  \n"
                            "Possible reasons: all topics are already covered by viact.ai, "
                            "topics were recently generated (12-week dedup), or all competitor "
                            "searches returned low-signal results.  \n\n"
                            + (f"**Last agent messages:**  \n{_last_logs}  \n\n" if _last_logs else "")
                            + "Try a different **Industry Vertical** or click **Run again**."
                        )
                    else:
                        st.session_state["r3_results"] = radar_results
                        st.session_state["r3_step"] = 1
                        st.rerun()
                except Exception as e:
                    st.error(f"Agent 1 failed: {e}")

    # =============================================================================
    # STEP 1 — Confirmed Gap Cards + HITL Topic Selection
    # =============================================================================
    elif step == 1:
        radar = st.session_state["r3_results"]
        topics = radar.get("topics", [])
        scan_ts = radar.get("scan_timestamp", "")
        n_scanned = radar.get("total_competitors_scanned", 0)

        st.markdown(_html(f"""
    <div style="margin-bottom:16px;">
    <div style="display:flex; align-items:center; gap:12px; margin-bottom:6px;">
    <h3 style="margin:0; color:#e6edf3;">&#127919; These Topics Competitors Cover — viAct Does Not</h3>
    <span style="background:rgba(255,106,61,0.1); border:1px solid rgba(255,106,61,0.3); border-radius:20px; padding:3px 12px; font-size:0.78rem; color:#ff6a3d; font-weight:600;">{_t(n_scanned)} sites scanned &middot; {_t(scan_ts)}</span>
    </div>
    <p style="margin:0; color:#8b949e; font-size:0.84rem;">Each gap below was <strong style="color:#3fb950;">live-verified</strong> by searching viact.ai &#8212; only topics with zero existing pages are shown. Pick one to build a webpage for.</p>
    </div>
    """), unsafe_allow_html=True)

        # ── Render topic cards ─────────────────────────────────────────────────────
        for i, topic in enumerate(topics):
            opp = topic.get("opportunity_score", "?")
            opp_class = {"High": "badge-high", "Medium": "badge-medium", "Low": "badge-low"}.get(opp, "badge-medium")
            comp_count = topic.get("competitor_count", 0)
            evidence = topic.get("competitor_evidence", [])

            evidence_rows = "".join(
                "<tr style='border-bottom:1px solid #2d303a;'>"
                f"<td style='padding:6px 10px; color:#ff6a3d; font-weight:600; font-size:0.8rem; white-space:nowrap;'>{_t(e.get('competitor',''))}</td>"
                f"<td style='padding:6px 10px; font-size:0.78rem;'><a href='{_t(e.get('url',''))}' target='_blank' style='color:#58a6ff; text-decoration:none;'>{_t(e.get('url','')[:55])}…</a></td>"
                f"<td style='padding:6px 10px; color:#8b949e; font-size:0.76rem;'>{_t(e.get('snippet','')[:90])}…</td>"
                "</tr>"
                for e in evidence[:4]
            )
            evidence_table = (
                "<table style='width:100%; border-collapse:collapse;'>"
                "<thead><tr style='background:rgba(255,255,255,0.03);'>"
                "<th style='padding:6px 10px; color:#8b949e; font-size:0.78rem; text-align:left;'>Competitor</th>"
                "<th style='padding:6px 10px; color:#8b949e; font-size:0.78rem; text-align:left;'>URL</th>"
                "<th style='padding:6px 10px; color:#8b949e; font-size:0.78rem; text-align:left;'>Snippet</th>"
                "</tr></thead>"
                f"<tbody>{evidence_rows}</tbody></table>"
                if evidence else
                "<div style='color:#484f58; font-size:0.82rem;'>Gap confirmed via Tavily search — no direct snippet evidence collected.</div>"
            )
            plural_s = "s" if comp_count != 1 else ""

            st.markdown(_html(f"""
    <div class="glass-card" style="border-color:rgba(255,106,61,0.2);">
    <div style="display:flex; justify-content:space-between; align-items:flex-start; margin-bottom:10px;">
    <div>
    <div style="color:#8b949e; font-size:0.72rem; font-weight:700; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:4px;">Content Gap {i+1} of {len(topics)}</div>
    <h4 style="margin:0; color:#e6edf3; font-size:1.15rem; font-weight:700;">{_t(topic['topic'])}</h4>
    </div>
    <div style="display:flex; gap:8px; flex-shrink:0; margin-left:16px;">
    <span class="badge-confirmed">&#10003; viAct Has No Page</span>
    <span class="{opp_class}">{_t(opp)} Opportunity</span>
    </div>
    </div>
    <div style="background:rgba(63,185,80,0.05); border:1px solid rgba(63,185,80,0.2); border-radius:6px; padding:8px 12px; margin-bottom:14px; display:flex; align-items:center; gap:10px;">
    <span style="font-size:1rem;">&#128269;</span>
    <span style="font-size:0.8rem; color:#8b949e;">Tavily searched <code>viact.ai</code> for <strong style="color:#c9d1d9;">&ldquo;{_t(topic['topic'])}&rdquo;</strong> &#8594; <strong style="color:#3fb950;">0 solution pages found</strong> &middot; checked {_t(topic.get('confirmed_at', ''))}</span>
    </div>
    <div style="display:grid; grid-template-columns:1fr 2fr; gap:20px;">
    <div>
    <div style="color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1px; margin-bottom:6px;">Why Trending</div>
    <div style="color:#c9d1d9; font-size:0.87rem;">{_t(topic.get('why_trending', ''))}</div>
    </div>
    <div>
    <div style="color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1px; margin-bottom:8px;">Competitor Evidence ({_t(comp_count)} competitor{plural_s})</div>
    {evidence_table}
    </div>
    </div>
    </div>
    """), unsafe_allow_html=True)

        # ── HITL Gate ──────────────────────────────────────────────────────────────
        st.divider()
        st.subheader("Which gap should we build a page for?")

        topic_options = [f"Gap {i+1}: {t['topic']}" for i, t in enumerate(topics)]
        selected_option = st.selectbox(
            "Select a gap to build a webpage for:",
            topic_options,
            key="r3_topic_choice",
        )
        selected_idx = topic_options.index(selected_option)

        st.write("")

        # ── Reference Library — persistent storage in Google Sheets ───────────────
        with st.expander("📚 Reference Library — auto-loaded in every run (Manual + GitHub Actions)", expanded=False):
            st.caption(
                "Add viAct project stats, case studies, MOM/OSHAD data here ONCE. "
                "Every future pipeline run — including Monday's automated run — will load these automatically."
            )
            _ref_svc = None
            _ref_sheet_id = os.getenv("SHEET_ID", "")
            _lib_rows = []
            try:
                from push_to_sheets import get_sheets_service, read_reference_library, add_reference, ensure_reference_tab
                from googleapiclient.discovery import build as _gsa_build
                _ref_svc = get_sheets_service()
                ensure_reference_tab(_ref_svc, _ref_sheet_id)
                # Load raw rows for display
                _raw = _ref_svc.spreadsheets().values().get(
                    spreadsheetId=_ref_sheet_id,
                    range="Reference_Library!A:D"
                ).execute()
                _lib_rows = _raw.get("values", [])[1:]  # skip header
            except Exception:
                pass

            selected_topic_name = topics[selected_idx]["topic"] if topics else ""

            if _lib_rows:
                st.markdown(f"**{len(_lib_rows)} reference(s) stored** — auto-loaded for every Agent 3 run")
                _display = []
                for r in _lib_rows:
                    _type = r[0] if len(r) > 0 else ""
                    _filter = r[1] if len(r) > 1 else "—"
                    _text = r[2][:80] + "..." if len(r) > 2 and len(r[2]) > 80 else (r[2] if len(r) > 2 else "")
                    _date = r[3] if len(r) > 3 else ""
                    _display.append({"Type": _type, "Topic Filter": _filter or "—", "Text Preview": _text, "Added": _date})
                st.dataframe(_display, use_container_width=True, hide_index=True)
            else:
                st.info("No references stored yet. Add your first one below.")

            st.markdown("**Add a new reference:**")
            _col1, _col2 = st.columns([1, 2])
            with _col1:
                _ref_type = st.selectbox("Type", ["global", "topic"], key="ref_lib_type",
                    help="global = always included  |  topic = only when topic name matches filter")
                _topic_filter = st.text_input("Topic Filter (for topic type)", key="ref_lib_filter",
                    placeholder="e.g. fatigue, permit, confined",
                    help="Leave blank for global. Enter keyword that must appear in the topic name.")
            with _col2:
                _ref_text = st.text_area("Reference Text", key="ref_lib_text", height=100,
                    placeholder=(
                        "e.g. Marina Bay Sands: 0 incidents, 18 months, 2400 workers\n"
                        "MOM WSH 2024: 35% of fatalities from falls\n"
                        "viAct: 90% risk reduction, 400+ sites, $2.5M savings"
                    ))
            if st.button("💾 Save to Reference Library", key="ref_lib_save"):
                if _ref_text.strip() and _ref_svc:
                    try:
                        add_reference(_ref_svc, _ref_sheet_id, _ref_type, _topic_filter.strip(), _ref_text.strip())
                        st.success("Saved! This reference will be loaded in all future runs.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"Save failed: {exc}")
                elif not _ref_text.strip():
                    st.warning("Please enter reference text before saving.")
                else:
                    st.warning("Sheets not connected — check SHEET_ID and service account credentials.")

        st.write("")
        st.markdown("**📄 Additional Reference Material** — for this run only (optional)")
        st.caption(
            "Already have a Reference Library above? This is for one-off data specific to this run only. "
            "What to paste: MOM/BCA report stats · viAct project data · regulatory quotes"
        )

        references = st.text_area(
            "Reference material (optional)",
            placeholder=(
                "e.g. MOM WSH Report 2024: falls from height = 35% of fatalities\n"
                "viAct Marina Bay Sands project: 0 incidents across 18 months\n"
                "BCA: construction sector accounts for 28% of workplace fatalities\n"
                "Leave blank — Reference Library entries are loaded automatically"
            ),
            height=100,
            key="r3_refs_input",
        )

        firecrawl_available = bool(os.getenv("FIRECRAWL_API_KEY"))

        if not firecrawl_available:
            st.markdown(_html("""
    <div style="background:rgba(210,153,34,0.06); border:1px solid rgba(210,153,34,0.3); border-radius:8px; padding:10px 14px; font-size:0.82rem; color:#d6a126; margin:10px 0;">
    <strong>&#9888; Firecrawl key not set</strong> &#8212; Agent 2 will skip competitor page scraping. Agent 3 will use only the short Tavily snippets. Content quality will be lower. Add <code>FIRECRAWL_API_KEY</code> for best results.
    </div>
    """), unsafe_allow_html=True)

        st.write("")
        st.markdown(_html("""
    <div style="background:rgba(255,106,61,0.04); border:1px solid rgba(255,106,61,0.15); border-radius:8px; padding:10px 16px; margin-bottom:10px; display:flex; gap:10px; align-items:center;">
    <span style="font-size:1.2rem;">&#9201;</span>
    <div style="color:#8b949e; font-size:0.8rem;">
    <strong style="color:#c9d1d9;">What happens next:</strong> Agent 2 downloads competitor pages (30s) &#8594; Agent 3 writes your full webpage using only that real content (60-90s). Total: ~2 minutes.
    </div>
    </div>
    """), unsafe_allow_html=True)
        if st.button(
            f"⚡  Build Webpage for Gap {selected_idx + 1}: {topics[selected_idx]['topic'][:50]}{'…' if len(topics[selected_idx]['topic']) > 50 else ''}",
            type="primary",
            use_container_width=True,
            key="r3_generate",
        ):
            # Merge: Reference Library (persistent) + manual textarea (this-run only)
            selected_topic = topics[selected_idx]
            _lib_refs = ""
            try:
                from push_to_sheets import get_sheets_service as _gss2, read_reference_library as _rrl
                _svc2 = _gss2()
                _lib_refs = _rrl(_svc2, os.getenv("SHEET_ID", ""), topic=selected_topic["topic"])
            except Exception:
                pass
            manual_refs = references.strip()
            raw_refs = "\n".join(filter(None, [_lib_refs, manual_refs]))

            st.session_state["r3_selected_idx"] = selected_idx
            st.session_state["r3_references"] = raw_refs
            st.session_state["r3_unverified"] = not bool(raw_refs)

            # ── Agent 2: Firecrawl scrape ──────────────────────────────────────────
            competitor_urls = [
                e["url"] for e in selected_topic.get("competitor_evidence", []) if e.get("url")
            ]

            competitor_data: dict = {}
            if firecrawl_available and competitor_urls:
                from agent2_data_extractor import extract_competitor_content

                a2_status = st.empty()
                a2_log: list[str] = []

                def _a2_cb(phase: str, message: str):
                    a2_log.append(message)
                    a2_status.markdown(
                        "<div class='log-box'>" + "<br>".join(a2_log[-8:]) + "</div>",
                        unsafe_allow_html=True,
                    )

                with st.spinner("Agent 2: Firecrawl scraping competitor pages..."):
                    try:
                        competitor_data = extract_competitor_content(
                            competitor_urls, progress_callback=_a2_cb
                        )
                    except Exception as exc:
                        st.warning(f"Agent 2 failed: {exc} — proceeding with Tavily snippets only.")
            else:
                for url in competitor_urls:
                    competitor_data[url] = {
                        "success": False,
                        "markdown": "[ACCESS DENIED]",
                        "word_count": 0,
                        "error": "Firecrawl key not configured",
                    }

            st.session_state["r3_competitor_data"] = competitor_data

            # ── Agent 3: Content generation ────────────────────────────────────────
            viact_pages = radar.get("viact_known_pages", [])

            with st.spinner("Agent 3: Generating structured content suite (Llama 3.3 70B)..."):
                try:
                    from agent3_content_architect import generate_structured_content, build_webpage_body

                    content = generate_structured_content(
                        topic=selected_topic["topic"],
                        competitor_data=competitor_data,
                        viact_pages=viact_pages,
                        references=raw_refs,
                        radar_topic_entry=selected_topic,
                    )

                    if not content.get("webpage_body"):
                        content["webpage_body"] = build_webpage_body(content)

                    st.session_state["r3_content"] = content
                    st.session_state["r3_step"] = 2
                    # Auto-save pillar page to Sheets immediately after generation
                    try:
                        from push_to_sheets import push_webpage_vertical as _pwv_auto
                        _r3_unverified = [u for u, r in competitor_data.items() if not r.get("success")]
                        _pwv_auto(
                            content=content,
                            decision_logic=content.get("decision_logic", ""),
                            input_source=f"3-Agent Radar — {st.session_state.get('r3_results', {}).get('scan_timestamp', '')}",
                            competitor_urls=list(competitor_data.keys()),
                            unverified=_r3_unverified,
                        )
                        st.toast(f"✓ Auto-saved to Sheets: {selected_topic['topic'][:50]}")
                    except Exception as _r3_ae:
                        st.toast(f"⚠ Auto-save failed: {_r3_ae}")
                    st.rerun()
                except Exception as e:
                    st.error(f"Agent 3 failed: {e}")

    # =============================================================================
    # STEP 2 — Content Preview + Push to Sheets
    # =============================================================================
    elif step == 2:
        radar = st.session_state["r3_results"]
        selected_idx = st.session_state["r3_selected_idx"]
        references = st.session_state.get("r3_references", "")
        unverified = st.session_state.get("r3_unverified", True)
        competitor_data = st.session_state.get("r3_competitor_data", {})
        selected_topic = radar["topics"][selected_idx]
        topic_str = selected_topic["topic"]
        content = st.session_state["r3_content"]

        # ── Header ─────────────────────────────────────────────────────────────────
        _unverified_banner = "<div style='margin-top:10px; background:rgba(210,153,34,0.06); border:1px solid rgba(210,153,34,0.25); border-radius:6px; padding:8px 12px; font-size:0.82rem; color:#d6a126;'>&#9888; <strong>[Unverified]</strong> &#8212; No reference material was provided. Statistics use public MOM/BCA data only. Add real project data before publishing.</div>" if unverified else ""
        _mb = "10px" if unverified else "0"
        st.markdown(_html(f"""
    <div class="glass-card" style="border-color:rgba(63,185,80,0.3); background:rgba(22,25,33,0.8);">
    <div style="display:flex; align-items:center; justify-content:space-between; flex-wrap:wrap; gap:12px; margin-bottom:{_mb};">
    <div style="display:flex; align-items:center; gap:12px;">
    <div style="background:rgba(63,185,80,0.15); border:1px solid #238636; border-radius:50%; width:40px; height:40px; display:flex; align-items:center; justify-content:center; font-size:1.2rem; flex-shrink:0;">&#10003;</div>
    <div>
    <div style="color:#3fb950; font-size:0.72rem; font-weight:700; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:3px;">10-Section Webpage Ready to Publish</div>
    <h3 style="margin:0; color:#e6edf3; font-size:1.1rem;">{_t(topic_str)}</h3>
    </div>
    </div>
    <div style="display:flex; gap:8px; flex-wrap:wrap;">
    <span style="background:rgba(88,166,255,0.1); color:#58a6ff; border:1px solid rgba(88,166,255,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#128269; Tavily verified</span>
    <span style="background:rgba(255,106,61,0.1); color:#ff6a3d; border:1px solid rgba(255,106,61,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#128293; Firecrawl scraped</span>
    <span style="background:rgba(63,185,80,0.08); color:#3fb950; border:1px solid rgba(63,185,80,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#129302; Llama 3.3 written</span>
    </div>
    </div>
    {_unverified_banner}
    </div>
    """), unsafe_allow_html=True)

        # ── Push to Sheets ─────────────────────────────────────────────────────────
        _push_col, _info_col = st.columns([1, 3])
        with _push_col:
            if st.button("📊  Push to Google Sheets", type="primary", key="r3_push", use_container_width=True):
                try:
                    from push_to_sheets import push_webpage_vertical, WEBPAGE_VERTICAL_TAB
                    competitor_urls_list = list(competitor_data.keys())
                    push_webpage_vertical(
                        content=content,
                        decision_logic=content.get("decision_logic", ""),
                        input_source=f"3-Agent Radar — {radar.get('scan_timestamp', '')}",
                        competitor_urls=competitor_urls_list,
                        unverified=unverified,
                    )
                    sheet_url = f"https://docs.google.com/spreadsheets/d/{os.getenv('SHEET_ID', '')}"
                    st.success(f"✅ Saved to **'{WEBPAGE_VERTICAL_TAB}'** tab — [Open Sheet ↗]({sheet_url})")
                except Exception as e:
                    st.error(f"Sheets error: {e}")
        with _info_col:
            st.markdown(_html("""
    <div style="background:rgba(22,25,33,0.5); border:1px solid #2d303a; border-radius:8px; padding:10px 14px; font-size:0.8rem; color:#8b949e; margin-top:4px;">
    &#128221; Saves to the <strong style="color:#c9d1d9;">Webpage Content</strong> tab in vertical format — field name in col A, value in col B. Every push stacks below the previous one.
    </div>
    """), unsafe_allow_html=True)

        st.markdown("<hr/>", unsafe_allow_html=True)

        # ── Blog Cluster — 3 supporting posts ──────────────────────────────────────
        st.markdown(_html("""
<div style="margin-bottom:10px;">
  <span style="color:#58a6ff; font-size:0.72rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px;">Step 2 — Blog Cluster (3 supporting posts)</span>
  <p style="margin:4px 0 0; color:#8b949e; font-size:0.82rem;">Generate 3 short blog posts that support the pillar above — one on regulatory compliance, one on cost/ROI, one on how-to. Each gets pushed to the same date tab in Sheets.</p>
</div>
"""), unsafe_allow_html=True)

        _cluster_state = st.session_state.get("r3_cluster_state", "idle")  # idle | running | done | error

        if _cluster_state == "done":
            _blog_results = st.session_state.get("r3_blog_results", [])
            st.success(f"✅ {len(_blog_results)} blog posts pushed to Sheets in today's date tab.")
            for _bi, _br in enumerate(_blog_results, 1):
                st.markdown(f"**Blog {_bi}:** {_br.get('topic', '')}")
            if st.button("🔄 Regenerate Cluster", key="r3_regen_cluster"):
                st.session_state["r3_cluster_state"] = "idle"
                st.rerun()
        else:
            _cl1, _cl2 = st.columns([1, 3])
            with _cl1:
                _gen_cluster = st.button("📝  Generate Blog Cluster", type="primary", key="r3_gen_cluster", use_container_width=True)
            with _cl2:
                st.markdown(_html("""
<div style="background:rgba(22,25,33,0.5); border:1px solid #2d303a; border-radius:8px; padding:10px 14px; font-size:0.8rem; color:#8b949e; margin-top:4px;">
&#128196; Generates <strong style="color:#c9d1d9;">3 shorter blog posts</strong> (800 words each) that internally link back to the pillar. All pushed to your Sheet automatically.
</div>"""), unsafe_allow_html=True)

            if _gen_cluster:
                from agent3_content_architect import generate_cluster_topics, generate_structured_content
                from push_to_sheets import push_webpage_vertical as _pwv, WEBPAGE_VERTICAL_TAB as _WVT
                _prim_kw = content.get("seo_suite", {}).get("primary_keyword", topic_str)
                _blog_results = []
                _cluster_err = None
                with st.spinner("Generating blog cluster topics..."):
                    try:
                        _blog_topics = generate_cluster_topics(topic_str, _prim_kw)
                    except Exception as _e:
                        _blog_topics = []
                        _cluster_err = str(_e)
                if _cluster_err:
                    st.error(f"Cluster topic generation failed: {_cluster_err}")
                else:
                    _prog_cluster = st.empty()
                    for _j, _bt in enumerate(_blog_topics[:3], 1):
                        with _prog_cluster.container():
                            st.info(f"Writing blog {_j}/3: '{_bt}'...")
                        try:
                            _bc = generate_structured_content(
                                topic=_bt,
                                competitor_data=competitor_data,
                                viact_pages=[],
                                references=references,
                                radar_topic_entry={
                                    "topic": _bt,
                                    "competitor_evidence": [],
                                    "confirmed_at": selected_topic.get("confirmed_at", ""),
                                    "opportunity_score": selected_topic.get("opportunity_score", "Medium"),
                                    "competitor_count": selected_topic.get("competitor_count", 0),
                                    "viact_search_query": f"blog for: {topic_str}",
                                    "why_trending": f"Supporting blog for pillar: {topic_str}",
                                    "pillar_topic": topic_str,
                                },
                                content_type="blog",
                            )
                            _pwv(
                                content=_bc,
                                decision_logic=_bc.get("decision_logic", ""),
                                input_source=f"Blog Cluster — {topic_str[:40]} ({_j}/3)",
                                competitor_urls=list(competitor_data.keys()),
                                unverified=unverified,
                                tab_name=_WVT,
                            )
                            _blog_results.append(_bc)
                        except Exception as _be:
                            st.warning(f"Blog {_j} failed: {_be} — skipping")
                    _prog_cluster.empty()
                    st.session_state["r3_blog_results"] = _blog_results
                    st.session_state["r3_cluster_state"] = "done"
                    st.rerun()

        st.markdown("<hr/>", unsafe_allow_html=True)

        # ── Content Preview Tabs ───────────────────────────────────────────────────
        st.markdown("<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:8px;'>PILLAR PAGE PREVIEW — ALL 10 SECTIONS</p>", unsafe_allow_html=True)
        (
            tab_dl, tab_sources, tab_body, tab_seo,
            tab_faqs, tab_schema, tab_geo,
            tab_visual, tab_links, tab_raw
        ) = st.tabs([
            "📋 Decision Logic",
            "🔍 Proof & Sources",
            "📄 Page Body",
            "🔎 SEO Tags",
            "❓ FAQs",
            "🏷️ Schema Markup",
            "🌐 AI Citations",
            "📷 Image Briefs",
            "🔗 Internal Links",
            "🔧 Raw JSON",
        ])

        with tab_dl:
            st.markdown(_html("""
    <div style="background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; display:flex; gap:10px; align-items:flex-start;">
    <span style="font-size:1.2rem; flex-shrink:0;">&#128161;</span>
    <div style="font-size:0.83rem; color:#58a6ff; line-height:1.5;">
    <strong>What is this?</strong> A ready-to-send paragraph for Gary / Surendra explaining WHY we should build this page.
    It cites the exact Tavily search date, competitor names, and URLs &#8212; not just "we think this is a good idea."
    </div>
    </div>
    """), unsafe_allow_html=True)
            st.text_area(
                "Decision Logic",
                content.get("decision_logic", ""),
                height=220,
                key="r3_dl_text",
            )

        with tab_sources:
            st.markdown(_html("""
    <div style="background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:16px; display:flex; gap:10px; align-items:flex-start;">
    <span style="font-size:1.2rem; flex-shrink:0;">&#128269;</span>
    <div style="font-size:0.83rem; color:#58a6ff; line-height:1.5;">
    <strong>What is this?</strong> The proof trail &#8212; exactly what Agent 1 searched, what Agent 2 scraped, and what Agent 3 actually used to write the content. If someone asks "how do you know this is a real gap?" &#8212; show them this tab.
    </div>
    </div>
    """), unsafe_allow_html=True)
            st.markdown("<h4 style='color:#e6edf3;'>Agent 1 — Tavily Gap Confirmation</h4>", unsafe_allow_html=True)
            st.markdown(_html(f"""
    <div class="glass-card" style="padding:16px;">
    <div style="margin-bottom:8px;"><span style="color:#8b949e; font-size:0.8rem;">Search:</span> <code>{_t(selected_topic.get('viact_search_query', ''))}</code> &#8594; <span style="color:#3fb950; font-weight:700;">0 dedicated solution pages</span></div>
    <div style="margin-bottom:8px;"><span style="color:#8b949e; font-size:0.8rem;">Confirmed:</span> <span style="color:#c9d1d9;">{_t(selected_topic.get('confirmed_at', ''))}</span></div>
    <div><span style="color:#8b949e; font-size:0.8rem;">Opportunity:</span> <span style="color:#ff6a3d; font-weight:700;">{_t(selected_topic.get('opportunity_score', '?'))}</span> &nbsp;&middot;&nbsp; <span style="color:#8b949e; font-size:0.8rem;">{_t(selected_topic.get('competitor_count', 0))} competitors covering this topic</span></div>
    </div>
    """), unsafe_allow_html=True)

            evidence = selected_topic.get("competitor_evidence", [])
            if evidence:
                st.markdown("<div style='color:#e6edf3; font-weight:600; margin:12px 0 8px;'>Competitor snippets:</div>", unsafe_allow_html=True)
                for e in evidence:
                    with st.expander(f"{e.get('competitor', '')} — {e.get('url', '')}"):
                        st.write(e.get("snippet", ""))

            st.markdown("<hr/><h4 style='color:#e6edf3;'>Agent 2 — Firecrawl Scrape Results</h4>", unsafe_allow_html=True)
            if competitor_data:
                for url, result in competitor_data.items():
                    if result.get("success"):
                        st.markdown(f"<div style='color:#3fb950; font-size:0.84rem; margin-bottom:4px;'>&#10003; <code>{_t(url[:65])}</code> &mdash; {_t(result.get('word_count', 0))} words scraped</div>", unsafe_allow_html=True)
                    else:
                        st.markdown(f"<div style='color:#f85149; font-size:0.84rem; margin-bottom:4px;'>&#x1F6AB; <code>{_t(url[:65])}</code> &mdash; ACCESS DENIED ({_t(result.get('error', 'unknown'))})</div>", unsafe_allow_html=True)
            else:
                st.markdown("<div style='color:#484f58; font-size:0.84rem;'>No competitor URLs were scraped.</div>", unsafe_allow_html=True)

            used = content.get("data_sources_used", [])
            denied = content.get("access_denied_urls", [])
            if used:
                st.markdown(f"<div style='margin-top:12px; color:#8b949e; font-size:0.82rem;'>Agent 3 used: {_t(', '.join(used[:4]))}</div>", unsafe_allow_html=True)
            if denied:
                st.markdown(f"<div style='color:#f85149; font-size:0.82rem;'>Agent 3 skipped (ACCESS DENIED): {_t(', '.join(denied[:4]))}</div>", unsafe_allow_html=True)

        with tab_body:
            st.markdown(_html("""
    <div style="background:rgba(63,185,80,0.05); border:1px solid rgba(63,185,80,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; display:flex; gap:10px; align-items:flex-start;">
    <span style="font-size:1.2rem; flex-shrink:0;">&#128196;</span>
    <div style="font-size:0.83rem; color:#3fb950; line-height:1.5;">
    <strong>What is this?</strong> The full webpage in Markdown format. Copy the text below and paste it into your CMS (WordPress, Webflow, Notion, etc.). Use "Preview" to see how it will look rendered.
    </div>
    </div>
    """), unsafe_allow_html=True)
            body = content.get("webpage_body", "")
            st.text_area("Webpage Body (Markdown — paste into CMS)", body, height=500, key="r3_body_text")
            with st.expander("👁️ Preview — see how the page renders"):
                st.markdown(body)

        with tab_seo:
            st.markdown(_html("""
    <div style="background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; display:flex; gap:10px; align-items:flex-start;">
    <span style="font-size:1.2rem; flex-shrink:0;">&#128269;</span>
    <div style="font-size:0.83rem; color:#58a6ff; line-height:1.5;">
    <strong>What is this?</strong> Everything your web developer needs for SEO. Paste the meta title and description into your CMS &lt;head&gt;. Use the keywords for on-page copy. Heading map tells developers the H1/H2 structure.
    </div>
    </div>
    """), unsafe_allow_html=True)
            seo = content.get("seo_suite", {})
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("<div style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1px; margin-bottom:8px;'>Meta</div>", unsafe_allow_html=True)
                st.text_input("Meta Title", seo.get("meta_title", ""), key="r3_meta_title")
                st.caption(f"{len(seo.get('meta_title', ''))} chars — target 50-60")
                st.text_area("Meta Description", seo.get("meta_description", ""), height=90, key="r3_meta_desc")
                st.caption(f"{len(seo.get('meta_description', ''))} chars — target 140-155")
                st.text_input("Canonical Slug", seo.get("canonical_url_slug", ""), key="r3_slug")
            with c2:
                st.markdown("<div style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1px; margin-bottom:12px;'>Keywords</div>", unsafe_allow_html=True)
                st.markdown(f"<div style='margin-bottom:8px;'><span style='color:#8b949e; font-size:0.82rem;'>Primary: </span><code>{_t(seo.get('primary_keyword', ''))}</code></div>", unsafe_allow_html=True)
                st.markdown("<div style='margin-bottom:8px; color:#8b949e; font-size:0.82rem;'>Secondary: " + " &nbsp;&middot;&nbsp; ".join(f"<code>{_t(k)}</code>" for k in seo.get("secondary_keywords", [])) + "</div>", unsafe_allow_html=True)
                st.markdown("<div style='margin-bottom:12px; color:#8b949e; font-size:0.82rem;'>LSI: " + " &nbsp;&middot;&nbsp; ".join(f"<code>{_t(k)}</code>" for k in seo.get("lsi_keywords", [])) + "</div>", unsafe_allow_html=True)
                st.markdown("<div style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1px; margin-bottom:8px;'>Heading Map</div>", unsafe_allow_html=True)
                for h in seo.get("heading_map", []):
                    depth = 1 if h.startswith("H1") else (2 if h.startswith("H2") else 3)
                    indent = (depth - 1) * 14
                    st.markdown(f"<div style='margin-left:{indent}px; color:#c9d1d9; font-size:0.84rem; margin-bottom:4px;'>{'&#9472;' * (depth-1)} {_t(h)}</div>", unsafe_allow_html=True)

        with tab_faqs:
            st.markdown(_html("""
    <div style="background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; display:flex; gap:10px; align-items:flex-start;">
    <span style="font-size:1.2rem; flex-shrink:0;">&#10067;</span>
    <div style="font-size:0.83rem; color:#58a6ff; line-height:1.5;">
    <strong>What is this?</strong> Two types: <strong>Schema FAQs</strong> (5 short answers &#8212; become the JSON-LD markup that Google shows as rich results in search) and <strong>Extended FAQs</strong> (2 longer answers for the actual webpage). Copy both to your CMS.
    </div>
    </div>
    """), unsafe_allow_html=True)
            st.markdown("<div style='color:#e6edf3; font-weight:700; margin-bottom:12px;'>Schema FAQs <span style='color:#8b949e; font-size:0.82rem; font-weight:400;'>— 5 items · 40-60 word answers · used in JSON-LD</span></div>", unsafe_allow_html=True)
            for i, faq in enumerate(content.get("schema_faqs", []), 1):
                with st.expander(f"Q{i}: {faq.get('question', '')}"):
                    st.write(faq.get("answer", ""))
                    st.caption(f"{len(faq.get('answer','').split())} words")
            st.markdown("<hr/><div style='color:#e6edf3; font-weight:700; margin:12px 0;'>Extended FAQs <span style='color:#8b949e; font-size:0.82rem; font-weight:400;'>— 2 items · 80-120 word answers · on-page only</span></div>", unsafe_allow_html=True)
            for i, faq in enumerate(content.get("extended_faqs", []), 1):
                with st.expander(f"Extended Q{i}: {faq.get('question', '')}"):
                    st.write(faq.get("answer", ""))
                    st.caption(f"{len(faq.get('answer','').split())} words")

        with tab_schema:
            st.markdown(_html("""
    <div style="background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; display:flex; gap:10px; align-items:flex-start;">
    <span style="font-size:1.2rem; flex-shrink:0;">&#127991;</span>
    <div style="font-size:0.83rem; color:#58a6ff; line-height:1.5;">
    <strong>What is this?</strong> Structured data code for Google. Give this to your developer and tell them: "Add this inside a &lt;script type=&rdquo;application/ld+json&rdquo;&gt; tag in the page &lt;head&gt;." It makes Google show your FAQs directly in search results.
    </div>
    </div>
    """), unsafe_allow_html=True)
            st.code(content.get("schema_json_ld", ""), language="json")

        with tab_geo:
            geo = content.get("geo_package", {})
            st.markdown(_html("""
    <div style="background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; display:flex; gap:10px; align-items:flex-start;">
    <span style="font-size:1.2rem; flex-shrink:0;">&#127760;</span>
    <div style="font-size:0.83rem; color:#58a6ff; line-height:1.5;">
    <strong>What is this?</strong> The opening 200 words are written specifically so that AI tools like ChatGPT, Perplexity, and Google AI Overviews will quote viAct when someone searches this topic. Use this exact text as the first paragraph of your webpage.
    </div>
    </div>
    """), unsafe_allow_html=True)
            st.markdown("<div style='color:#e6edf3; font-weight:600; margin-bottom:8px;'>Opening 200 Words — AI-citation optimized:</div>", unsafe_allow_html=True)
            st.text_area("Opening 200 words", geo.get("opening_200_words", ""), height=180, key="r3_geo")
            st.markdown("<div style='color:#e6edf3; font-weight:600; margin:14px 0 8px;'>Citation Framing Tips:</div>", unsafe_allow_html=True)
            for tip in geo.get("citation_framing_tips", []):
                st.markdown(f"<div style='background:rgba(22,25,33,0.6); border:1px solid #2d303a; border-radius:6px; padding:8px 12px; margin-bottom:6px; color:#c9d1d9; font-size:0.85rem;'>&#8594; {_t(tip)}</div>", unsafe_allow_html=True)

        with tab_visual:
            prompts = content.get("nano_banana_prompts", content.get("visual_strategy", []))
            st.markdown("<div style='color:#8b949e; font-size:0.82rem; margin-bottom:12px;'>Realistic photography, APAC context, human-centered — not CGI. For use with Nano Banana 2.</div>", unsafe_allow_html=True)
            for i, v in enumerate(prompts, 1):
                with st.expander(f"Image {i} — {v.get('placement', '')}"):
                    st.text_area(f"Prompt {i}", v.get("prompt", ""), height=140, key=f"r3_vis_{i}")
                    st.markdown(f"<div style='color:#8b949e; font-size:0.82rem; margin-top:6px;'><strong style='color:#e6edf3;'>Alt text:</strong> {_t(v.get('alt_text', ''))}</div>", unsafe_allow_html=True)

        with tab_links:
            st.markdown("<div style='color:#e6edf3; font-weight:600; margin-bottom:10px;'>Internal Links — verified viAct.ai URLs only:</div>", unsafe_allow_html=True)
            for link in content.get("internal_links", []):
                url_val = link.get("url", "")
                st.markdown(
                    f"<div style='background:rgba(22,25,33,0.6); border:1px solid #2d303a; border-radius:6px; padding:10px 14px; margin-bottom:8px;'>"
                    f"<div style='color:#ff6a3d; font-weight:600; font-size:0.88rem;'>{_t(link.get('anchor_text', ''))}</div>"
                    f"<div style='margin:4px 0;'><a href='{_t(url_val)}' target='_blank' style='color:#58a6ff; font-size:0.82rem;'>{_t(url_val)}</a></div>"
                    f"<div style='color:#8b949e; font-size:0.8rem;'>{_t(link.get('context', ''))}</div>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

        with tab_raw:
            st.json(content)


elif _sel == "industry":
    # ── Session state ─────────────────────────────────────────────────────────
    if "ip_step" not in st.session_state:
        st.session_state["ip_step"] = 0

    ip_step = st.session_state["ip_step"]

    if ip_step == 0:
        st.markdown(_html("""
        <div style="display:flex;gap:8px;margin-bottom:20px;align-items:stretch;">
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #3fb950;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">🏭</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 1 — Industry &amp; File</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Select an industry from the dropdown (Construction, Oil &amp; Gas, Mining…). Optionally upload an approved .docx reference file with real stats or client data.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #3fb950;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">⚙️</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 2 — Generate</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Click "Generate Industry Page". The AI scrapes viAct + competitor pages and builds a full 8-section Wix CMS landing page — Hero, Metrics, Use Cases, Testimonials, CTA.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #3fb950;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">📊</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 3 — Review &amp; Save</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Review content across tabs — CMS Fields, SEO, FAQs, Image Briefs. When ready, click "Save to Google Sheets".</div>
          </div>
        </div>
        """), unsafe_allow_html=True)

    if ip_step > 0:
        st.write("")
        if st.button("↩ Start Over", key="ip_reset"):
            for _k in [_k for _k in st.session_state if _k.startswith("ip_")]:
                del st.session_state[_k]
            st.rerun()

    st.write("")

    # ── Suggested topic banner (from Daily Intel Scan) ────────────────────────
    _sug_ind = st.session_state.get("suggested_industry_topic")
    if _sug_ind and ip_step == 0:
        _si_col, _si_btn = st.columns([5, 1])
        with _si_col:
            st.info(f"💡 **Today's suggested topic:** {_sug_ind.get('topic', '')} — _{_sug_ind.get('industry', '')}_\n\n_{_sug_ind.get('why', '')}_")
        with _si_btn:
            st.write("")
            if st.button("Use This →", key="ip_use_suggested", use_container_width=True):
                st.session_state["ip_industry_select"] = "Custom (type below) →"
                st.session_state["ip_custom_industry_text"] = _sug_ind.get("industry", "")
                st.rerun()

    # =========================================================================
    # INDUSTRY STEP 0 — Select Industry & Generate
    # =========================================================================
    if ip_step == 0:
        st.markdown(
            "<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:12px;'>GENERATE A FULL INDUSTRY VERTICAL LANDING PAGE</p>",
            unsafe_allow_html=True,
        )
        st.markdown(_html(
            '<div class="glass-card" style="margin-bottom:18px;">'
            '<div style="color:#8b949e; font-size:0.82rem; line-height:1.7;">'
            "Select your industry. The system will:<br>"
            "&nbsp;1. Scrape viAct existing industry page (tone reference via Firecrawl)<br>"
            "&nbsp;2. Scrape 2-3 competitor industry pages (Firecrawl)<br>"
            "&nbsp;3. Generate a complete 8-section landing page: SEO, FAQs, 11 image prompts (Llama 3.3 70B)"
            "</div></div>"
        ), unsafe_allow_html=True)

        import sys as _sys
        _sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))
        from agent3_content_architect import INDUSTRY_VIACT_URLS, INDUSTRY_COMPETITOR_URLS

        _PRESET_INDUSTRIES = list(INDUSTRY_VIACT_URLS.keys()) + [
            "Automotive & EV Safety",
            "Maritime Safety",
            "Logistics & Warehousing Safety",
            "Smart Cities & Public Sector",
            "Pharmaceutical & Chemical Safety",
            "Custom (type below) →",
        ]

        _ind_col, _ref_col = st.columns([1, 1], gap="medium")
        with _ind_col:
            ip_industry_choice = st.selectbox("Industry", _PRESET_INDUSTRIES, key="ip_industry_select")
            if ip_industry_choice == "Custom (type below) →":
                ip_industry = st.text_input(
                    "Custom industry name",
                    placeholder="e.g. Pharmaceutical Safety, Smart Cities Safety...",
                    key="ip_custom_industry_text",
                )
            else:
                ip_industry = ip_industry_choice
        with _ref_col:
            ip_yt = st.text_input(
                "YouTube Video URL (optional — embedded in hero)",
                placeholder="https://www.youtube.com/watch?v=...",
                key="ip_yt_url",
            )

        # ── .docx reference uploader ──────────────────────────────────────────
        _uploaded_doc = st.file_uploader(
            "📄 Upload approved content .docx (auto-fills reference field below)",
            type=["docx"],
            key="ip_doc_upload",
            help="Upload finalized copy — the AI uses it as ground-truth reference for metrics, use cases, and testimonials.",
        )
        if _uploaded_doc is not None:
            if st.session_state.get("ip_last_doc_name") != _uploaded_doc.name:
                try:
                    import io
                    import docx as _docx_lib
                    _doc_obj = _docx_lib.Document(io.BytesIO(_uploaded_doc.read()))
                    _doc_text = "\n".join(p.text for p in _doc_obj.paragraphs if p.text.strip())
                    st.session_state["ip_refs_text"] = _doc_text[:6000]
                    st.session_state["ip_last_doc_name"] = _uploaded_doc.name
                    st.success(f"✓ Loaded **{_uploaded_doc.name}** ({len(_doc_text):,} chars) → reference field auto-filled")
                except Exception as _de:
                    st.error(f"Could not parse doc: {_de}")

        ip_refs = st.text_area(
            "Reference Material — paste stats, or upload .docx above (auto-populated)",
            height=120,
            key="ip_refs_text",
        )

        with st.expander("⚙️ Custom Instructions (optional — focus areas, key hazards, regional requirements)"):
            ip_custom_inst = st.text_area(
                "",
                placeholder="e.g. Focus on methane gas detection and ATEX compliance. Include UAE Taqa project context. Emphasise LTI reduction metric.",
                height=90,
                key="ip_custom_instructions",
                label_visibility="collapsed",
            )

        st.write("")

        # ── HITL Review: show what will be used before generation ────────────
        _preview_industry = ip_industry if ip_industry else "—"
        _preview_comp = INDUSTRY_COMPETITOR_URLS.get(ip_industry, []) if ip_industry and ip_industry != "Custom (type below) →" else []
        _preview_refs = (ip_refs or "").strip()
        _preview_custom = (ip_custom_inst or "").strip() if "ip_custom_instructions" in st.session_state else ""
        with st.expander("🔍 Review Inputs Before Generating", expanded=False):
            st.markdown(_html(
                '<div style="font-size:0.82rem; color:#8b949e; line-height:1.8;">'
                f'<strong style="color:#c9d1d9;">Industry:</strong> {_t(_preview_industry)}<br>'
                f'<strong style="color:#c9d1d9;">Competitor pages to scrape:</strong> '
                + (_t(", ".join(_preview_comp)) if _preview_comp else '<span style="color:#f85149;">None (custom industry)</span>')
                + '<br>'
                f'<strong style="color:#c9d1d9;">Reference material:</strong> '
                + (f'<span style="color:#3fb950;">✓ {len(_preview_refs)} chars loaded</span>' if _preview_refs else '<span style="color:#e3b341;">⚠ None — using viAct verified stats only</span>')
                + '<br>'
                + (f'<strong style="color:#c9d1d9;">Custom instructions:</strong> {_t(_preview_custom[:120])}{"..." if len(_preview_custom) > 120 else ""}<br>' if _preview_custom else '')
                + '</div>'
            ), unsafe_allow_html=True)

        if st.button("🏭  Generate Industry Page", type="primary", key="ip_generate"):
            if not ip_industry or not ip_industry.strip():
                st.warning("Please enter an industry name.")
                st.stop()

            # Resolve slug and URLs
            if ip_industry in INDUSTRY_VIACT_URLS:
                _industry_slug = INDUSTRY_VIACT_URLS[ip_industry].rsplit("/", 1)[-1]
                _viact_url     = INDUSTRY_VIACT_URLS[ip_industry]
                _comp_urls     = INDUSTRY_COMPETITOR_URLS.get(ip_industry, [])
            else:
                _industry_slug = ip_industry.lower().replace(" ", "-").replace("&", "and").replace("/", "-")
                _viact_url     = ""
                _comp_urls     = []

            _refs_combined = ip_refs.strip()
            if ip_yt.strip():
                _refs_combined = f"Hero YouTube Video URL: {ip_yt.strip()}\n\n" + _refs_combined

            from agent2_data_extractor import extract_competitor_content
            from agent3_content_architect import generate_industry_page

            _prog = st.empty()

            if _viact_url:
                with _prog.container():
                    st.info("Step 1/3 — Scraping viAct industry page (tone reference)...")
                try:
                    _viact_scraped = extract_competitor_content([_viact_url])
                    _viact_md = _viact_scraped.get(_viact_url, {}).get("markdown", "")
                except Exception:
                    _viact_md = ""
                _comp_data_all = {_viact_url: {"success": bool(_viact_md), "markdown": _viact_md, "word_count": len(_viact_md.split())}}
            else:
                # Custom industry: scrape viact.ai case-studies page as reference
                with _prog.container():
                    st.info("Step 1/3 — Scraping viAct case studies (custom industry reference)...")
                try:
                    from agent3_content_architect import INDUSTRY_CASE_STUDY_URL
                    _cs_scraped = extract_competitor_content([INDUSTRY_CASE_STUDY_URL])
                    _viact_md = _cs_scraped.get(INDUSTRY_CASE_STUDY_URL, {}).get("markdown", "")
                except Exception:
                    _viact_md = ""
                _comp_data_all = {}

            with _prog.container():
                st.info(f"Step 2/3 — Scraping {len(_comp_urls)} competitor industry pages (Firecrawl)..." if _comp_urls else "Step 2/3 — No competitor URLs for this industry. Generating from viAct data...")
            try:
                _comp_data = extract_competitor_content(_comp_urls) if _comp_urls else {}
            except Exception:
                _comp_data = {}
            _comp_data_all.update(_comp_data)

            with _prog.container():
                st.info("Step 3/3 — Generating 8-section industry page (Llama 3.3 70B)...")
            try:
                _radar_viact_pages = st.session_state.get("r3_results", {}).get("viact_known_pages", [])
                _result = generate_industry_page(
                    industry_name=ip_industry,
                    industry_slug=_industry_slug,
                    viact_page_content=_viact_md,
                    competitor_content=_comp_data,
                    references=_refs_combined,
                    viact_pages=_radar_viact_pages,
                    custom_instructions=ip_custom_inst.strip() if ip_custom_inst else "",
                )
                st.session_state["ip_content"]          = _result
                st.session_state["ip_competitor_data"]  = _comp_data_all
                st.session_state["ip_industry_label"]   = ip_industry
                st.session_state["ip_industry_slug"]    = _industry_slug
                st.session_state["ip_viact_url"]        = _viact_url
                st.session_state["ip_comp_urls"]        = _comp_urls
                st.session_state["ip_refs_saved"]       = _refs_combined
                st.session_state["ip_custom_inst_saved"]= ip_custom_inst.strip() if ip_custom_inst else ""
                st.session_state["ip_step"]             = 1
                # Auto-save to Sheets immediately after generation
                try:
                    from push_to_sheets import push_industry_page_vertical as _push_ip_auto
                    _auto_ip_sid = os.getenv("INDUSTRY_SHEET_ID") or os.getenv("SHEET_ID", "")
                    _push_ip_auto(content=_result, industry_name=ip_industry, sheet_id=_auto_ip_sid)
                    st.toast(f"✓ Auto-saved to Sheets: {ip_industry}")
                except Exception as _ip_ae:
                    st.toast(f"⚠ Auto-save failed: {_ip_ae}")
                _prog.empty()
                st.rerun()
            except Exception as _e:
                _prog.empty()
                st.error(f"Generation failed: {_e}")

    # =========================================================================
    # INDUSTRY STEP 1 — Preview + Push to Sheets
    # =========================================================================
    elif ip_step == 1:
        _ip_content    = st.session_state["ip_content"]
        _ip_comp_data  = st.session_state.get("ip_competitor_data", {})
        _ip_label      = st.session_state.get("ip_industry_label", "")
        _ip_viact_url  = st.session_state.get("ip_viact_url", "")

        # ── Quality gate warning ─────────────────────────────────────────────
        _gate_errors = _ip_content.get("quality_gate_errors", [])
        if _gate_errors:
            st.warning(
                "⚠️ **Quality gate retry was triggered** — the following issues were detected in the first generation "
                "and a correction was requested:\n" + "\n".join(f"- {e}" for e in _gate_errors)
                + "\n\nReview the output below to confirm the corrections were applied."
            )

        st.markdown(_html(
            '<div class="glass-card" style="border-color:rgba(63,185,80,0.3); background:rgba(22,25,33,0.8);">'
            '<div style="display:flex; align-items:center; justify-content:space-between; flex-wrap:wrap; gap:12px;">'
            '<div style="display:flex; align-items:center; gap:12px;">'
            '<div style="background:rgba(63,185,80,0.15); border:1px solid #238636; border-radius:50%; width:40px; height:40px; display:flex; align-items:center; justify-content:center; font-size:1.2rem; flex-shrink:0;">&#127970;</div>'
            '<div>'
            '<div style="color:#3fb950; font-size:0.72rem; font-weight:700; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:3px;">Industry Landing Page Ready</div>'
            f'<h3 style="margin:0; color:#e6edf3; font-size:1.1rem;">{_t(_ip_label)}</h3>'
            '</div></div>'
            '<div style="display:flex; gap:8px; flex-wrap:wrap;">'
            '<span style="background:rgba(255,106,61,0.1); color:#ff6a3d; border:1px solid rgba(255,106,61,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#128293; Firecrawl scraped</span>'
            '<span style="background:rgba(63,185,80,0.08); color:#3fb950; border:1px solid rgba(63,185,80,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#129302; Llama 3.3 written</span>'
            '<span style="background:rgba(88,166,255,0.1); color:#58a6ff; border:1px solid rgba(88,166,255,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#127970; 8-section structure</span>'
            '</div></div></div>'
        ), unsafe_allow_html=True)

        _ip_push_col, _ip_info_col = st.columns([1, 3])
        with _ip_push_col:
            if st.button("📊  Save to Google Sheets", type="primary", key="ip_push", use_container_width=True):
                try:
                    from push_to_sheets import push_industry_page_vertical
                    _ip_sheet_id = os.getenv("INDUSTRY_SHEET_ID") or os.getenv("SHEET_ID", "")
                    push_industry_page_vertical(
                        content=_ip_content,
                        industry_name=_ip_label,
                        sheet_id=_ip_sheet_id,
                    )
                    _sheet_url = f"https://docs.google.com/spreadsheets/d/{_ip_sheet_id}/edit"
                    st.success(f"✅ Saved to **'{_ip_label}'** tab in your Sheet — [Open Sheet ↗]({_sheet_url})")
                except Exception as _e:
                    st.error(f"Sheets error: {_e}")
        with _ip_info_col:
            st.markdown(_html(
                f'<div style="background:rgba(22,25,33,0.5); border:1px solid #2d303a; border-radius:8px; padding:10px 14px; font-size:0.8rem; color:#8b949e; margin-top:4px;">'
                f'&#128221; Creates a <strong style="color:#c9d1d9;">{_ip_label}</strong> tab in your Sheet. Each CMS field gets its own column — copy any field directly.'
                f'</div>'
            ), unsafe_allow_html=True)

        st.markdown("<hr/>", unsafe_allow_html=True)
        st.markdown("<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:8px;'>PREVIEW ALL SECTIONS</p>", unsafe_allow_html=True)

        (
            _ip_tab_dl, _ip_tab_sources, _ip_tab_body, _ip_tab_cms,
            _ip_tab_seo, _ip_tab_faqs, _ip_tab_schema, _ip_tab_geo,
            _ip_tab_visual, _ip_tab_links, _ip_tab_raw,
        ) = st.tabs([
            "📋 Decision Logic",
            "🔍 Proof & Sources",
            "📄 Page Body",
            "🗂️ Wix CMS Fields",
            "🔎 SEO Tags",
            "❓ FAQs",
            "🏷️ Schema Markup",
            "🌐 AI Citations",
            "📷 Image Briefs",
            "🔗 Internal Links",
            "🔧 Raw JSON",
        ])

        with _ip_tab_dl:
            st.text_area("Decision Logic", _ip_content.get("decision_logic", ""), height=220, key="ip_dl_text")

        with _ip_tab_sources:
            st.markdown("<h4 style='color:#e6edf3;'>viAct Industry Page (Tone Reference)</h4>", unsafe_allow_html=True)
            _viact_res = _ip_comp_data.get(_ip_viact_url, {})
            if _viact_res.get("success"):
                st.markdown(f"<div style='color:#3fb950; font-size:0.84rem; margin-bottom:4px;'>&#10003; <code>{_t(_ip_viact_url[:70])}</code> &mdash; {_viact_res.get('word_count', 0)} words scraped</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='color:#f85149; font-size:0.84rem; margin-bottom:4px;'>&#x1F6AB; <code>{_t(_ip_viact_url[:70])}</code> &mdash; ACCESS DENIED (fresh content generated)</div>", unsafe_allow_html=True)
            st.markdown("<hr/><h4 style='color:#e6edf3;'>Competitor Industry Pages (Firecrawl)</h4>", unsafe_allow_html=True)
            for _url, _res in _ip_comp_data.items():
                if _url == _ip_viact_url:
                    continue
                if _res.get("success"):
                    st.markdown(f"<div style='color:#3fb950; font-size:0.84rem; margin-bottom:4px;'>&#10003; <code>{_t(_url[:65])}</code> &mdash; {_res.get('word_count', 0)} words</div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div style='color:#f85149; font-size:0.84rem; margin-bottom:4px;'>&#x1F6AB; <code>{_t(_url[:65])}</code> &mdash; ACCESS DENIED</div>", unsafe_allow_html=True)

        with _ip_tab_body:
            _ip_body = _ip_content.get("webpage_body", "")
            st.text_area("Webpage Body (Markdown — copy-paste into Wix CMS)", _ip_body, height=500, key="ip_body_text")
            with st.expander("👁️ Preview — see how the page renders"):
                st.markdown(_ip_body)

        with _ip_tab_cms:
            _ip_cms = _ip_content.get("industry_cms_fields", {})
            if not _ip_cms:
                st.info("No CMS fields in this output — regenerate to get individual Wix CMS copy-paste fields.")
            else:
                st.markdown(
                    "<div style='background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; font-size:0.83rem; color:#58a6ff;'>"
                    "<strong>Wix CMS Fields</strong> — Each field maps to one dynamic field in your Wix CMS dataset. Copy one field at a time directly into the matching CMS field."
                    "</div>",
                    unsafe_allow_html=True,
                )
                st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:10px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Hero Section</div>", unsafe_allow_html=True)
                st.text_input("Hero Subheadline", _ip_cms.get("hero_subheadline", ""), key="ip_cms_hero_sub")
                st.text_area("Hero Body Copy", _ip_cms.get("hero_body_copy", ""), height=90, key="ip_cms_hero_body")

                st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Impact Section</div>", unsafe_allow_html=True)
                st.text_input("Impact Section Title", _ip_cms.get("impact_section_title", ""), key="ip_cms_impact_title")
                st.text_input("Impact Subtitle", _ip_cms.get("impact_subtitle", ""), key="ip_cms_impact_sub")

                st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Metrics (3 Stats)</div>", unsafe_allow_html=True)
                for _mi, _met in enumerate(_ip_cms.get("metrics", []), 1):
                    _mc1, _mc2 = st.columns([1, 2])
                    with _mc1:
                        st.text_input(f"Metric {_mi} Label", _met.get("label", ""), key=f"ip_cms_met_label_{_mi}")
                    with _mc2:
                        st.text_input(f"Metric {_mi} Description", _met.get("description", ""), key=f"ip_cms_met_desc_{_mi}")

                st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Use Cases Section</div>", unsafe_allow_html=True)
                st.text_input("Use Cases Section Title", _ip_cms.get("use_cases_section_title", ""), key="ip_cms_uc_section_title")
                for _uci, _uc in enumerate(_ip_cms.get("use_cases", []), 1):
                    with st.expander(f"Use Case {_uci}: {_uc.get('title', '')}"):
                        st.text_input(f"UC{_uci} Title", _uc.get("title", ""), key=f"ip_cms_uc_title_{_uci}")
                        st.text_area(f"UC{_uci} Description", _uc.get("description", ""), height=80, key=f"ip_cms_uc_desc_{_uci}")

                st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Solutions & viGent</div>", unsafe_allow_html=True)
                st.text_area("Solutions Description", _ip_cms.get("solutions_description", ""), height=80, key="ip_cms_solutions")
                st.text_area("viGent Description", _ip_cms.get("vigent_description", ""), height=100, key="ip_cms_vigent")

                st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Testimonials (5)</div>", unsafe_allow_html=True)
                for _ti, _test in enumerate(_ip_cms.get("testimonials", []), 1):
                    with st.expander(f"Testimonial {_ti} — {_test.get('source', '')}"):
                        st.text_area(f"T{_ti} Quote", _test.get("quote", ""), height=90, key=f"ip_cms_test_quote_{_ti}")
                        st.text_input(f"T{_ti} Source", _test.get("source", ""), key=f"ip_cms_test_src_{_ti}")

                st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>CTA Section</div>", unsafe_allow_html=True)
                st.text_input("CTA Headline", _ip_cms.get("cta_headline", ""), key="ip_cms_cta_headline")
                st.text_area("CTA Description", _ip_cms.get("cta_description", ""), height=80, key="ip_cms_cta_desc")

        with _ip_tab_seo:
            _ip_seo = _ip_content.get("seo_suite", {})
            _c1, _c2 = st.columns(2)
            with _c1:
                st.text_input("Meta Title", _ip_seo.get("meta_title", ""), key="ip_meta_title")
                st.caption(f"{len(_ip_seo.get('meta_title', ''))} chars — target 50-60")
                st.text_area("Meta Description", _ip_seo.get("meta_description", ""), height=90, key="ip_meta_desc")
                st.caption(f"{len(_ip_seo.get('meta_description', ''))} chars — target 140-155")
                st.text_input("Canonical Slug", _ip_seo.get("canonical_url_slug", ""), key="ip_slug")
            with _c2:
                st.markdown(f"<div style='margin-bottom:8px;'><span style='color:#8b949e; font-size:0.82rem;'>Primary: </span><code>{_t(_ip_seo.get('primary_keyword', ''))}</code></div>", unsafe_allow_html=True)
                st.markdown("<div style='margin-bottom:8px; color:#8b949e; font-size:0.82rem;'>Secondary: " + " &nbsp;&middot;&nbsp; ".join(f"<code>{_t(k)}</code>" for k in _ip_seo.get("secondary_keywords", [])) + "</div>", unsafe_allow_html=True)
                st.markdown("<div style='margin-bottom:12px; color:#8b949e; font-size:0.82rem;'>LSI: " + " &nbsp;&middot;&nbsp; ".join(f"<code>{_t(k)}</code>" for k in _ip_seo.get("lsi_keywords", [])) + "</div>", unsafe_allow_html=True)
                for _h in _ip_seo.get("heading_map", []):
                    _depth = 1 if _h.startswith("H1") else (2 if _h.startswith("H2") else 3)
                    st.markdown(f"<div style='margin-left:{(_depth-1)*14}px; color:#c9d1d9; font-size:0.84rem; margin-bottom:4px;'>{_t(_h)}</div>", unsafe_allow_html=True)

        with _ip_tab_faqs:
            st.markdown("<div style='color:#e6edf3; font-weight:700; margin-bottom:12px;'>Schema FAQs <span style='color:#8b949e; font-size:0.82rem; font-weight:400;'>— 5 items · used in JSON-LD rich results</span></div>", unsafe_allow_html=True)
            for _i, _faq in enumerate(_ip_content.get("schema_faqs", []), 1):
                with st.expander(f"Q{_i}: {_faq.get('question', '')}"):
                    st.write(_faq.get("answer", ""))
                    st.caption(f"{len(_faq.get('answer', '').split())} words")
            st.markdown("<hr/><div style='color:#e6edf3; font-weight:700; margin:12px 0;'>Extended FAQs <span style='color:#8b949e; font-size:0.82rem; font-weight:400;'>— 2 items · 80-120 words · on-page only</span></div>", unsafe_allow_html=True)
            for _i, _faq in enumerate(_ip_content.get("extended_faqs", []), 1):
                with st.expander(f"Extended Q{_i}: {_faq.get('question', '')}"):
                    st.write(_faq.get("answer", ""))

        with _ip_tab_schema:
            st.markdown(
                "<div style='background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; font-size:0.83rem; color:#58a6ff;'>"
                "<strong>What is this?</strong> Add this inside a &lt;script type=\"application/ld+json\"&gt; tag in the page &lt;head&gt;. Makes Google show FAQs as rich results."
                "</div>",
                unsafe_allow_html=True,
            )
            st.code(_ip_content.get("schema_json_ld", ""), language="json")

        with _ip_tab_geo:
            _ip_geo = _ip_content.get("geo_package", {})
            st.markdown("<div style='color:#e6edf3; font-weight:600; margin-bottom:8px;'>Opening 200 Words — AI-citation optimized:</div>", unsafe_allow_html=True)
            st.text_area("Opening 200 words", _ip_geo.get("opening_200_words", ""), height=180, key="ip_geo")
            st.markdown("<div style='color:#e6edf3; font-weight:600; margin:14px 0 8px;'>Citation Framing Tips:</div>", unsafe_allow_html=True)
            for _tip in _ip_geo.get("citation_framing_tips", []):
                st.markdown(f"<div style='background:rgba(22,25,33,0.6); border:1px solid #2d303a; border-radius:6px; padding:8px 12px; margin-bottom:6px; color:#c9d1d9; font-size:0.85rem;'>&#8594; {_t(_tip)}</div>", unsafe_allow_html=True)

        with _ip_tab_visual:
            _ip_prompts = _ip_content.get("nano_banana_prompts", [])
            st.markdown("<div style='color:#8b949e; font-size:0.82rem; margin-bottom:12px;'>11 image prompts — copy each prompt and paste into Nano Banana or any image tool to create the images.</div>", unsafe_allow_html=True)
            for _i, _v in enumerate(_ip_prompts, 1):
                with st.expander(f"Image {_i} — {_v.get('placement', '')}"):
                    _ip_prompt_txt = _v.get("prompt", "")
                    st.text_area(f"Prompt {_i}", _ip_prompt_txt, height=140, key=f"ip_vis_{_i}")
                    st.markdown(f"<div style='color:#8b949e; font-size:0.82rem; margin-top:6px;'><strong style='color:#e6edf3;'>Alt text:</strong> {_t(_v.get('alt_text', ''))}</div>", unsafe_allow_html=True)

        with _ip_tab_links:
            st.markdown("<div style='color:#e6edf3; font-weight:600; margin-bottom:10px;'>Internal Links — verified viAct.ai URLs only:</div>", unsafe_allow_html=True)
            for _link in _ip_content.get("internal_links", []):
                _url_val = _link.get("url", "")
                st.markdown(
                    f"<div style='background:rgba(22,25,33,0.6); border:1px solid #2d303a; border-radius:6px; padding:10px 14px; margin-bottom:8px;'>"
                    f"<div style='color:#ff6a3d; font-weight:600; font-size:0.88rem;'>{_t(_link.get('anchor_text', ''))}</div>"
                    f"<div style='margin:4px 0;'><a href='{_t(_url_val)}' target='_blank' style='color:#58a6ff; font-size:0.82rem;'>{_t(_url_val)}</a></div>"
                    f"<div style='color:#8b949e; font-size:0.8rem;'>{_t(_link.get('context', ''))}</div>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

        with _ip_tab_raw:
            st.json(_ip_content)

        # ── Refine / Improve loop ─────────────────────────────────────────────
        st.markdown("<hr style='border-color:#2d303a; margin:24px 0 16px;'/>", unsafe_allow_html=True)
        with st.expander("✏️ Refine this output — give feedback and regenerate"):
            st.markdown(
                "<div style='color:#8b949e; font-size:0.82rem; margin-bottom:10px;'>"
                "Describe what to improve. The agent will regenerate the full page with your feedback as highest priority. No re-scraping needed."
                "</div>",
                unsafe_allow_html=True,
            )
            ip_feedback = st.text_area(
                "What to change?",
                placeholder="e.g. Make testimonials more specific to mining safety. Change metric 2 to focus on LTI reduction. Add methane detection to use case 3.",
                height=100,
                key="ip_refine_feedback",
            )
            if st.button("🔄 Regenerate with Feedback", key="ip_regenerate_btn"):
                if not ip_feedback.strip():
                    st.warning("Please describe what to improve.")
                else:
                    _combined_inst = (
                        st.session_state.get("ip_custom_inst_saved", "")
                        + "\n\nIMPROVEMENT FEEDBACK (highest priority — apply these changes):\n"
                        + ip_feedback.strip()
                    ).strip()
                    from agent2_data_extractor import extract_competitor_content as _ece
                    from agent3_content_architect import generate_industry_page as _gip
                    with st.spinner("Regenerating with feedback..."):
                        try:
                            _new_result = _gip(
                                industry_name=st.session_state.get("ip_industry_label", ""),
                                industry_slug=st.session_state.get("ip_industry_slug", ""),
                                viact_page_content=st.session_state.get("ip_competitor_data", {}).get(
                                    st.session_state.get("ip_viact_url", ""), {}
                                ).get("markdown", ""),
                                competitor_content={
                                    k: v for k, v in st.session_state.get("ip_competitor_data", {}).items()
                                    if k != st.session_state.get("ip_viact_url", "")
                                },
                                references=st.session_state.get("ip_refs_saved", ""),
                                viact_pages=[],
                                custom_instructions=_combined_inst,
                            )
                            st.session_state["ip_content"] = _new_result
                            st.session_state["ip_custom_inst_saved"] = _combined_inst
                            st.rerun()
                        except Exception as _re:
                            st.error(f"Regeneration failed: {_re}")



# =============================================================================
# TAB — CASE STUDIES (Agent 06)
# Session state prefix: cs_
# =============================================================================
elif _sel == "casestudy":
    if "cs_step" not in st.session_state:
        st.session_state["cs_step"] = 0

    cs_step = st.session_state["cs_step"]

    if cs_step == 0:
        st.markdown(_html("""
        <div style="display:flex;gap:8px;margin-bottom:20px;align-items:stretch;">
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #58a6ff;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">📋</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 1 — Client Details</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Fill in company type, industry, location, and viAct products used. Upload a .docx reference file with project data, metrics, or approved copy.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #58a6ff;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">⚙️</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 2 — Generate</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Click "Generate Case Study". The AI builds all 56 Wix CMS fields — Problem, Solution, Impact, Metrics, Testimonials, Image Briefs, and SEO.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #58a6ff;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">📊</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 3 — Review &amp; Save</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Review across tabs — CMS Fields, SEO, Alt Texts, Image Prompts. Click "Save to Google Sheets" — the tab will be named after the client.</div>
          </div>
        </div>
        """), unsafe_allow_html=True)

    if cs_step > 0:
        st.write("")
        if st.button("↩ Start Over", key="cs_reset"):
            for _k in [_k for _k in st.session_state if _k.startswith("cs_")]:
                del st.session_state[_k]
            st.rerun()

    st.write("")

    # ── Suggested topic banner (from Daily Intel Scan) ────────────────────────
    _sug_cs = st.session_state.get("suggested_cs_topic")
    if _sug_cs and cs_step == 0:
        _sc_col, _sc_btn = st.columns([5, 1])
        with _sc_col:
            st.info(
                f"💡 **Today's suggested case study:** {_sug_cs.get('company_type', '')} · "
                f"{_sug_cs.get('location', '')} · {_sug_cs.get('detection_focus', '')}\n\n"
                f"_{_sug_cs.get('why', '')}_"
            )
        with _sc_btn:
            st.write("")
            if st.button("Use This →", key="cs_use_suggested", use_container_width=True):
                st.session_state["cs_location"] = _sug_cs.get("location", "")
                st.session_state["cs_products"] = _sug_cs.get("detection_focus", "")
                st.rerun()

    # =========================================================================
    # CASE STUDY STEP 0 — Inputs
    # =========================================================================
    if cs_step == 0:
        st.markdown(
            "<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:12px;'>GENERATE A FULL viAct CASE STUDY PAGE</p>",
            unsafe_allow_html=True,
        )
        st.markdown(_html(
            '<div class="glass-card" style="margin-bottom:18px;">'
            '<div style="color:#8b949e; font-size:0.82rem; line-height:1.7;">'
            "Fill in company details. The system will:<br>"
            "&nbsp;1. Research the company via Tavily<br>"
            "&nbsp;2. Scrape viAct case-studies reference page (tone + structure)<br>"
            "&nbsp;3. Generate Hero → Challenge → Solution → Impact → Testimonials (Llama 3.3 70B)"
            "</div></div>"
        ), unsafe_allow_html=True)

        import sys as _sys
        _sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))
        from agent6_case_study_builder import VIACT_PRODUCTS

        _cs_col1, _cs_col2 = st.columns(2, gap="medium")
        with _cs_col1:
            cs_company  = st.text_input("Company Name *", placeholder="e.g. Samsung C&T, ADNOC, Gamuda", key="cs_company")
            cs_location = st.text_input("Location *", placeholder="e.g. Singapore, Dubai UAE, Kuala Lumpur", key="cs_location")
        with _cs_col2:
            _INDUSTRIES = [
                "Construction", "Oil & Gas", "Chemical", "Energy",
                "Infrastructure", "Mining", "Food & Beverage",
                "Pharmaceuticals", "Port & Logistics", "Manufacturing",
                "Marine", "Other",
            ]
            cs_industry = st.selectbox("Industry *", _INDUSTRIES, key="cs_industry_select")
            cs_products = st.multiselect(
                "viAct Products Used *",
                VIACT_PRODUCTS,
                default=["PPE Detection"],
                key="cs_products",
            )

        _cs_doc = st.file_uploader(
            "📄 Upload case study .docx (auto-fills reference field below)",
            type=["docx"],
            key="cs_doc_upload",
            help="Upload your Word doc — metrics, quotes, project details auto-extracted.",
        )
        if _cs_doc is not None:
            if st.session_state.get("cs_last_doc_name") != _cs_doc.name:
                try:
                    import io
                    import docx as _docx_lib
                    _cs_doc_obj = _docx_lib.Document(io.BytesIO(_cs_doc.read()))
                    _cs_doc_text = "\n".join(p.text for p in _cs_doc_obj.paragraphs if p.text.strip())
                    st.session_state["cs_refs"] = _cs_doc_text[:6000]
                    st.session_state["cs_last_doc_name"] = _cs_doc.name
                    st.success(f"✓ Loaded **{_cs_doc.name}** ({len(_cs_doc_text):,} chars) → reference field auto-filled")
                except Exception as _de:
                    st.error(f"Could not parse doc: {_de}")

        cs_refs = st.text_area(
            "Reference Material — paste real metrics, quotes, or upload .docx above (auto-populated)",
            height=110,
            placeholder="e.g. Reduced incidents by 75%, 5,000 worker-hours saved, Quote from EHS Director",
            key="cs_refs",
        )

        with st.expander("⚙️ Custom Instructions (optional)"):
            cs_custom = st.text_area(
                "",
                placeholder="e.g. Focus on e-PTW implementation. Mention MOM compliance. Include night-shift context.",
                height=80,
                key="cs_custom",
                label_visibility="collapsed",
            )

        cs_tavily = st.checkbox("🔍 Research company via Tavily", value=True, key="cs_tavily")

        st.write("")

        with st.expander("🔍 Review Before Generating", expanded=False):
            st.markdown(_html(
                '<div style="font-size:0.82rem; color:#8b949e; line-height:1.8;">'
                f'<strong style="color:#c9d1d9;">Company:</strong> {_t(cs_company or "—")}<br>'
                f'<strong style="color:#c9d1d9;">Industry:</strong> {_t(cs_industry)}<br>'
                f'<strong style="color:#c9d1d9;">Location:</strong> {_t(cs_location or "—")}<br>'
                f'<strong style="color:#c9d1d9;">Products:</strong> {_t(", ".join(cs_products) or "—")}<br>'
                + (f'<strong style="color:#c9d1d9;">References:</strong> <span style="color:#3fb950;">✓ {len((cs_refs or "").strip())} chars</span><br>' if (cs_refs or "").strip() else '<strong style="color:#c9d1d9;">References:</strong> <span style="color:#e3b341;">⚠ None — AI will infer from research</span><br>')
                + '</div>'
            ), unsafe_allow_html=True)

        if st.button("📋  Generate Case Study", type="primary", key="cs_generate"):
            if not cs_company.strip():
                st.warning("Please enter a company name.")
                st.stop()
            if not cs_location.strip():
                st.warning("Please enter a location.")
                st.stop()
            if not cs_products:
                st.warning("Please select at least one viAct product.")
                st.stop()

            from agent6_case_study_builder import generate_case_study

            _cs_prog = st.empty()
            _cs_steps = []

            def _cs_cb(msg):
                _cs_steps.append(msg)
                with _cs_prog.container():
                    st.info(_cs_steps[-1])

            try:
                _cs_result = generate_case_study(
                    company=cs_company.strip(),
                    industry=cs_industry,
                    location=cs_location.strip(),
                    products_used=cs_products,
                    references=cs_refs.strip() if cs_refs else "",
                    custom_instructions=cs_custom.strip() if cs_custom else "",
                    run_tavily=cs_tavily,
                    progress_callback=_cs_cb,
                )
                st.session_state["cs_result"]       = _cs_result
                st.session_state["cs_company_saved"] = cs_company.strip()
                st.session_state["cs_step"]          = 1
                # Auto-save to Sheets immediately after generation
                try:
                    from push_to_sheets import push_case_study as _push_cs_auto
                    _auto_cs_sid = os.getenv("INDUSTRY_SHEET_ID") or os.getenv("SHEET_ID", "")
                    _push_cs_auto(_cs_result, sheet_id=_auto_cs_sid)
                    st.toast(f"✓ Auto-saved to Sheets: {cs_company.strip()}")
                except Exception as _cs_ae:
                    st.toast(f"⚠ Auto-save failed: {_cs_ae}")
                _cs_prog.empty()
                st.rerun()
            except Exception as _ce:
                _cs_prog.empty()
                st.error(f"Generation failed: {_ce}")

    # =========================================================================
    # CASE STUDY STEP 1 — Results + Push to Sheets
    # =========================================================================
    elif cs_step == 1:
        _cs_result  = st.session_state["cs_result"]
        _cs_cms     = _cs_result.get("cms_fields", {})
        _cs_meta    = _cs_result.get("generation_meta", {})
        _cs_errors  = _cs_result.get("quality_gate_errors", [])
        _cs_company = st.session_state.get("cs_company_saved", "")

        if _cs_errors:
            st.warning(
                "⚠️ **Quality gate issues detected:**\n" +
                "\n".join(f"- {e}" for e in _cs_errors) +
                "\n\nReview output — placeholders in [brackets] need real data."
            )

        st.markdown(_html(
            '<div class="glass-card" style="border-color:rgba(88,166,255,0.3); background:rgba(22,25,33,0.8);">'
            '<div style="display:flex; align-items:center; gap:12px;">'
            '<div style="background:rgba(88,166,255,0.15); border:1px solid #388bfd; border-radius:50%; width:40px; height:40px; display:flex; align-items:center; justify-content:center; font-size:1.2rem; flex-shrink:0;">📋</div>'
            '<div>'
            '<div style="color:#58a6ff; font-size:0.72rem; font-weight:700; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:3px;">Case Study Ready</div>'
            f'<h3 style="margin:0; color:#e6edf3; font-size:1.1rem;">{_t(_cs_cms.get("hero_h1", _cs_company))}</h3>'
            '</div></div></div>'
        ), unsafe_allow_html=True)

        _cs_push_col, _cs_info_col = st.columns([1, 3])
        with _cs_push_col:
            if st.button("📊  Save to Google Sheets", type="primary", key="cs_push", use_container_width=True):
                try:
                    from push_to_sheets import push_case_study
                    _cs_sid = os.getenv("INDUSTRY_SHEET_ID") or os.getenv("SHEET_ID", "")
                    push_case_study(_cs_result, sheet_id=_cs_sid)
                    st.success("✓ Saved!")
                except Exception as _pe:
                    st.error(f"Sheets push failed: {_pe}")

        st.write("")

        _cs_t1, _cs_t2, _cs_t3, _cs_t4, _cs_t5 = st.tabs(["📝 CMS Fields", "🔍 SEO", "🖼 Alt Texts", "🎨 Image Prompts", "🔧 Raw JSON"])

        with _cs_t1:
            st.markdown("#### Hero")
            st.text_area("h1", value=_cs_cms.get("hero_h1", ""), height=60, key="cs_out_h1")
            st.text_area("h2", value=_cs_cms.get("h2", ""), height=60, key="cs_out_h2")
            st.text_area("h3 Intro", value=_cs_cms.get("h3", ""), height=80, key="cs_out_h3")
            st.text_area("Hero Image Brief", value=_cs_cms.get("hero_image_brief", ""), height=80, key="cs_out_img")

            st.markdown("#### Company Info")
            _ci1, _ci2 = st.columns(2)
            with _ci1:
                st.text_input("Company Name",  value=_cs_cms.get("company_name", ""),  key="cs_out_cname")
                st.text_input("Industry",      value=_cs_cms.get("industry", ""),      key="cs_out_ind")
                st.text_input("Location",      value=_cs_cms.get("location", ""),      key="cs_out_loc")
                st.text_input("Use Case",      value=_cs_cms.get("use_case", ""),      key="cs_out_uc")
            with _ci2:
                st.text_input("Company Size",  value=_cs_cms.get("company_size", ""),  key="cs_out_csize")
                st.text_input("Company Type",  value=_cs_cms.get("company_type", ""),  key="cs_out_ctype")
                _prods = _cs_cms.get("products_used", [])
                st.text_input("Products Used", value=", ".join(_prods) if isinstance(_prods, list) else str(_prods), key="cs_out_prods")
            st.text_area("Company Overview", value=_cs_cms.get("company_overview", ""), height=100, key="cs_out_cov")
            st.text_area("Story Snapshot",   value=_cs_cms.get("story_snapshot", ""),  height=80,  key="cs_out_snap")

            st.markdown("#### Key Metrics")
            for _i in (1, 2, 3):
                _mc, _ml, _md_col = st.columns([1, 2, 3])
                with _mc:
                    st.text_input(f"Metric {_i} Value", value=_cs_cms.get(f"metric_{_i}_value", ""), key=f"cs_out_mv{_i}")
                with _ml:
                    st.text_input(f"Metric {_i} Label", value=_cs_cms.get(f"metric_{_i}_label", ""), key=f"cs_out_ml{_i}")
                with _md_col:
                    st.text_input(f"Metric {_i} Description", value=_cs_cms.get(f"metric_{_i}_description", ""), key=f"cs_out_mdesc{_i}")

            st.markdown("#### The Challenge")
            st.text_input("Challenge Title", value=_cs_cms.get("challenge_title", ""), key="cs_out_chal_t")
            st.text_area("Challenge Body",   value=_cs_cms.get("challenge_body", ""),  height=180, key="cs_out_chal")

            st.markdown("#### The Solution")
            st.text_input("Solution Title", value=_cs_cms.get("solution_title", ""), key="cs_out_sol_t")
            st.text_area("Solution Body",   value=_cs_cms.get("solution_body", ""),  height=150, key="cs_out_sol")
            st.text_input("Subsection 1 Title", value=_cs_cms.get("solution_sub1_title", ""), key="cs_out_sub1t")
            st.text_area("Subsection 1 Body",   value=_cs_cms.get("solution_sub1_body", ""),  height=120, key="cs_out_sub1b")
            st.text_input("Subsection 2 Title", value=_cs_cms.get("solution_sub2_title", ""), key="cs_out_sub2t")
            st.text_area("Subsection 2 Body",   value=_cs_cms.get("solution_sub2_body", ""),  height=120, key="cs_out_sub2b")

            st.markdown("#### The Impact")
            st.text_input("Impact Title", value=_cs_cms.get("impact_title", ""), key="cs_out_imp_t")
            st.text_area("Impact Body",   value=_cs_cms.get("impact_body", ""),  height=150, key="cs_out_imp")

            st.markdown("#### Testimonials")
            for _i in (1, 2):
                _tq_col, _tr_col, _tc_col = st.columns([3, 2, 2])
                with _tq_col:
                    st.text_area(f"Quote {_i}",   value=_cs_cms.get(f"testimonial_{_i}_quote", ""),   height=80, key=f"cs_out_tq{_i}")
                with _tr_col:
                    st.text_input(f"Role {_i}",    value=_cs_cms.get(f"testimonial_{_i}_role", ""),    key=f"cs_out_tr{_i}")
                with _tc_col:
                    st.text_input(f"Company {_i}", value=_cs_cms.get(f"testimonial_{_i}_company", ""), key=f"cs_out_tc{_i}")

            st.markdown("#### CTA")
            st.text_input("CTA Headline", value=_cs_cms.get("cta_headline", ""), key="cs_out_cta")

        with _cs_t2:
            _mt = _cs_cms.get("meta_title", "")
            _md = _cs_cms.get("meta_description", "")
            _sl = _cs_cms.get("slug", "")
            _url = _cs_cms.get("url", "")
            _kw  = _cs_cms.get("keywords", "")
            _tags = _cs_cms.get("tags", [])
            st.text_input(f"Meta Title ({len(_mt)}/60 chars)", value=_mt, key="cs_out_mt")
            if len(_mt) > 60:
                st.error(f"Meta title is {len(_mt)} chars — must be ≤60")
            st.text_area(f"Meta Description ({len(_md)}/160 chars)", value=_md, height=80, key="cs_out_md")
            if not (140 <= len(_md) <= 165):
                st.warning(f"Meta description is {len(_md)} chars — aim for 140-160")
            st.text_input("URL Slug", value=_sl, key="cs_out_slug")
            st.text_input("URL (Full)", value=_url, key="cs_out_url")
            st.text_input("Tags / Filter Tag", value=", ".join(_tags) if isinstance(_tags, list) else str(_tags), key="cs_out_tags")
            st.text_area("Keywords", value=_kw, height=70, key="cs_out_kw")

        with _cs_t3:
            st.markdown("**Hero Section**")
            st.text_input("Hero Section image alt text",  value=_cs_cms.get("hero_alt_text", ""),            key="cs_out_alt_hero")
            st.markdown("**Metrics (2nd Section)**")
            st.text_input("1st alt text (Metric 1)",      value=_cs_cms.get("metric_1_alt_text", ""),        key="cs_out_alt_m1")
            st.text_input("2nd alt text (Metric 2)",      value=_cs_cms.get("metric_2_alt_text", ""),        key="cs_out_alt_m2")
            st.text_input("3rd alt text (Metric 3)",      value=_cs_cms.get("metric_3_alt_text", ""),        key="cs_out_alt_m3")
            st.markdown("**Company Overview (3rd Section)**")
            st.text_input("Company Overview alt text",    value=_cs_cms.get("section_alt_text", ""),         key="cs_out_alt_sec")
            st.markdown("**The Solution**")
            st.text_input("1 image alt text",             value=_cs_cms.get("solution_1_alt_text", ""),      key="cs_out_alt_s1")
            st.text_input("2 image alt text",             value=_cs_cms.get("solution_2_alt_text", ""),      key="cs_out_alt_s2")
            st.markdown("**Testimonial**")
            st.text_input("Company logo alt text",        value=_cs_cms.get("company_logo_alt_text", ""),    key="cs_out_alt_logo")
            st.text_input("Profile image alt text (1&2)", value=_cs_cms.get("profile_image_alt_text", ""),   key="cs_out_alt_profile")
            st.markdown("**Other**")
            st.text_input("Industry Alt Text",            value=_cs_cms.get("industry_alt_text", ""),        key="cs_out_alt_ind")
            st.text_input("Location Alt Text",            value=_cs_cms.get("location_alt_text", ""),        key="cs_out_alt_loc")
            st.text_input("Use Case Alt Text",            value=_cs_cms.get("use_case_alt_text", ""),        key="cs_out_alt_uc")

        with _cs_t4:
            st.markdown("<div style='color:#8b949e; font-size:0.82rem; margin-bottom:16px;'>5 image prompts — copy each into Nano Banana to generate images for the case study page.</div>", unsafe_allow_html=True)
            _img_fields = [
                ("1 — Hero Banner (426×423px — square)",              "hero_image_brief"),
                ("2 — Company Overview / Story Snapshot (342×414px — portrait)", "overview_image_brief"),
                ("3 — Solution Section 1 (1440×978px — landscape)",   "solution_1_image_brief"),
                ("4 — Solution Section 2 (2289×1400px — ultra-wide)", "solution_2_image_brief"),
                ("5 — Testimonial / Review Section (400×400px — square)", "testimonial_image_brief"),
            ]
            for _ilabel, _ikey in _img_fields:
                st.text_area(_ilabel, value=_cs_cms.get(_ikey, ""), height=110, key=f"cs_img_{_ikey}")

        with _cs_t5:
            st.json(_cs_result)


# =============================================================================
# TAB — VIDEO ANALYTICS ITEM PAGES (Agent 07)
# Session state prefix: va_
# One detection type → all ~45 Wix CMS text fields → Google Sheet tab
# =============================================================================
elif _sel == "va":
    if "va_result" not in st.session_state:
        st.session_state["va_result"] = None

    if st.session_state["va_result"] is None:
        st.markdown(_html("""
        <div style="display:flex;gap:8px;margin-bottom:20px;align-items:stretch;">
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #bc8cff;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">🎯</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 1 — Detection Type</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Enter a detection type — e.g. "Fall Detection", "PPE Detection", "Fire &amp; Smoke Detection". Or use a suggested topic from the Market Radar tab.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #bc8cff;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">⚙️</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 2 — Generate</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Click "Generate Page". The AI runs live research + viAct style scraping to build the full item page — Hero, Challenges, How It Works, Use Cases, Case Study, SEO.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #bc8cff;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">📊</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 3 — Review &amp; Save</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Review across 4 tabs — Webpage Content, SEO, Image Alt Texts, Raw JSON. Click "Save to Google Sheets" — saved as tab "VA — {Detection}".</div>
          </div>
        </div>
        """), unsafe_allow_html=True)

    st.markdown("### 🎯 Video Analytics Item Page Generator")
    st.caption("Generate all Wix CMS text fields for one detection-type item page. Output → Google Sheet tab 'VA — {Detection}'.")

    # ── Suggested topic banner (from Daily Intel Scan) ────────────────────────
    _sug_va = st.session_state.get("suggested_va_topic")
    if _sug_va and st.session_state.get("va_result") is None:
        _sv_col, _sv_btn = st.columns([5, 1])
        with _sv_col:
            st.info(f"💡 **Today's suggested detection:** **{_sug_va.get('detection_name', '')}**\n\n_{_sug_va.get('why', '')}_")
        with _sv_btn:
            st.write("")
            if st.button("Use This →", key="va_use_suggested", use_container_width=True):
                st.session_state["va_detection_input"] = _sug_va.get("detection_name", "")
                st.rerun()

    _va_detection = st.text_input(
        "Detection Type",
        placeholder="e.g. Hot Work Perimeter Violation Detection",
        key="va_detection_input",
    )

    _run_va = st.button(
        "🚀  Generate Page",
        type="primary",
        key="run_va",
        use_container_width=True,
        disabled=not _va_detection.strip(),
    )

    if _run_va:
        st.session_state["va_result"] = None
        with st.spinner(f"Generating '{_va_detection}' page content..."):
            from agent7_video_analytics_page import generate_va_page as _gen_va_page
            _va_res = _gen_va_page(_va_detection, progress_callback=lambda m: st.toast(m))
            st.session_state["va_result"] = _va_res
            # Auto-save to Sheets immediately after generation
            try:
                from push_to_sheets import push_video_analytics_page as _push_va_auto
                _push_va_auto(_va_res)
                st.toast(f"✓ Auto-saved to Sheets: VA — {_va_detection}")
            except Exception as _va_ae:
                st.toast(f"⚠ Auto-save failed: {_va_ae}")

        _va_errors = _va_res.get("quality_gate_errors", [])
        if _va_errors:
            st.warning("Quality gate warnings — review before publishing:")
            for _e in _va_errors:
                st.markdown(f"- ⚠️ {_e}")
        else:
            st.success(f"✅ '{_va_detection}' page generated with no quality issues.")

    _va_result = st.session_state.get("va_result")
    if _va_result:
        _va_cms = _va_result.get("cms_fields", {})
        _va_meta = _va_result.get("generation_meta", {})
        st.divider()

        _va_push_col, _va_info_col = st.columns([1, 3])
        with _va_push_col:
            if st.button("📊  Save to Google Sheets", type="primary", key="va_push", use_container_width=True):
                try:
                    from push_to_sheets import push_video_analytics_page as _push_va_fn
                    _push_va_fn(_va_result)
                    _det = _va_cms.get("title", "")
                    st.success(f"✓ Saved to 'VA — {_det}'")
                except Exception as _pe:
                    st.error(f"Sheets push failed: {_pe}")
        with _va_info_col:
            st.caption(f"Generated: {_va_meta.get('timestamp','')[:19]} UTC · Model: {_va_meta.get('model_used','')} · Retries: {_va_meta.get('retry_count',0)}")

        _va_tab_cms, _va_tab_seo, _va_tab_imgs, _va_tab_raw = st.tabs(["📝 Webpage Content", "🔎 SEO", "🖼 Image Alt Texts", "🔧 Raw JSON"])

        with _va_tab_cms:
            st.markdown("#### Hero")
            st.text_input("Title", value=_va_cms.get("title", ""), key="va_v_title")
            st.text_input("H1", value=_va_cms.get("h1", ""), key="va_v_h1")
            st.text_input("H2", value=_va_cms.get("h2", ""), key="va_v_h2")
            st.text_input("H3", value=_va_cms.get("h3", ""), key="va_v_h3")
            st.text_area("First Paragraph", value=_va_cms.get("first_paragraph", ""), height=160, key="va_v_fp")

            st.markdown("#### Challenges")
            st.text_input("Section Title [t1]", value=_va_cms.get("t1", ""), key="va_v_t1")
            st.text_area("Challenges Body [td]", value=_va_cms.get("td", ""), height=220, key="va_v_td")

            st.markdown("#### How Computer Vision Works")
            st.text_input("Section Title [t2]", value=_va_cms.get("t2", ""), key="va_v_t2")
            for _step, _tk, _dk in [
                ("Step 1 — Choose", "t2_ct2", "t2_cdesc2"),
                ("Step 2 — Connect", "t2_t1", "t2_1d"),
                ("Step 3 — Capture", "t3_1t", "t3_1d"),
                ("Step 4 — Control", "t4_t1", "t4_td"),
            ]:
                st.text_input(f"{_step} Title", value=_va_cms.get(_tk, ""), key=f"va_v_{_tk}")
                st.text_area(f"{_step} Description", value=_va_cms.get(_dk, ""), height=90, key=f"va_v_{_dk}")

            st.markdown("#### Where Needed Most")
            st.text_input("Section Title [s6_title]", value=_va_cms.get("s6_title", ""), key="va_v_s6t")
            st.text_area("Intro", value=_va_cms.get("s6_descriptions", ""), height=90, key="va_v_s6d")
            for _i in range(1, 6):
                st.text_input(f"Use Case {_i} Title", value=_va_cms.get(f"s6_t{_i}", ""), key=f"va_v_s6t{_i}")
                st.text_area(f"Use Case {_i} Description", value=_va_cms.get(f"s6_desc{_i}", ""), height=90, key=f"va_v_s6d{_i}")

            st.markdown("#### Case Study Snapshot")
            st.text_input("Headline [s7_title]", value=_va_cms.get("s7_title", ""), key="va_v_s7t")
            st.text_input("Industry Label", value=_va_cms.get("construction", ""), key="va_v_ind")
            st.text_input("Location Label", value=_va_cms.get("singapore", ""), key="va_v_loc")
            st.text_input("Module Label", value=_va_cms.get("open_edge_detection", ""), key="va_v_mod")
            st.text_area("The Problem", value=_va_cms.get("problem_description", ""), height=90, key="va_v_prob")
            st.text_area("The Solution", value=_va_cms.get("solution_description", ""), height=90, key="va_v_sol")
            st.text_area("The viAct impAct", value=_va_cms.get("viact_impact_descriptions", ""), height=90, key="va_v_imp")

            st.markdown("#### Why viAct")
            st.text_input("Section Title [s8_title]", value=_va_cms.get("s8_title", ""), key="va_v_s8t")
            st.text_input("Intro Line", value=_va_cms.get("s8_description", ""), key="va_v_s8d")
            for _i in range(1, 8):
                st.text_input(f"Bullet {_i}", value=_va_cms.get(f"s8_{_i}", ""), key=f"va_v_s8_{_i}")

        with _va_tab_seo:
            _mt = _va_cms.get("meta_title", "")
            st.text_input(f"Meta Title ({len(_mt)}/60 chars)", value=_mt, key="va_v_mt")
            _md = _va_cms.get("meta_descriptions", "")
            st.text_area(f"Meta Description ({len(_md)}/160 chars)", value=_md, height=90, key="va_v_md")
            st.text_area("Keywords", value=_va_cms.get("keywords", ""), height=70, key="va_v_kw")

        with _va_tab_imgs:
            st.text_input("Hero Image Alt Text", value=_va_cms.get("hero_image_alt_text", ""), key="va_v_alt0")
            for _i in range(1, 5):
                st.text_input(f"Image Alt Text {_i}", value=_va_cms.get(f"image_alt_text_{_i}", ""), key=f"va_v_alt{_i}")
            st.text_input("S7 Image Alt Text", value=_va_cms.get("s7_image_alt_text", ""), key="va_v_alts7")

        with _va_tab_raw:
            st.json(_va_result)


# =============================================================================
# TAB — PRODUCT PAGES (Agent 04)
# Session state prefix: pp_
# Gary's approach: viact.ai reference pages → LLM content → Claude HTML design
# =============================================================================
elif _sel == "product":
    if "pp_step" not in st.session_state:
        st.session_state["pp_step"] = 0

    pp_step = st.session_state["pp_step"]

    if pp_step == 0:
        st.markdown(_html("""
        <div style="display:flex;gap:8px;margin-bottom:20px;align-items:stretch;">
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #d6a126;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">🖥️</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 1 — Select Product</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Select a viAct product from the dropdown (viGent, viLID, viHUB…) or type a custom product name. Optionally upload .docx specs and add competitor URLs.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #d6a126;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">⚙️</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 2 — Generate</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Click "Generate Product Page". The AI builds Hero, Features, How It Works, Use Cases, Testimonials, FAQs, SEO, and 6 Image Briefs.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #d6a126;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">🎨</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 3 — HTML Design</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Review the content. Click "Generate HTML Page" — Claude AI designs a fully styled HTML page that can be downloaded directly.</div>
          </div>
        </div>
        """), unsafe_allow_html=True)

    if pp_step > 0:
        st.write("")
        if st.button("↩ Start Over", key="pp_reset"):
            for _k in [_k for _k in st.session_state if _k.startswith("pp_")]:
                del st.session_state[_k]
            st.rerun()

    st.write("")

    # ── Competitor launch prefill banner ──────────────────────────────────────
    _pp_prefill_url  = st.session_state.pop("product_prefill_url", None)
    _pp_prefill_name = st.session_state.pop("product_prefill_name", None)
    if _pp_prefill_url:
        st.info(f"💡 Responding to competitor launch: **{_pp_prefill_name}** — URL pre-filled in competitor URLs below.")
        st.session_state["pp_competitor_urls"] = _pp_prefill_url

    # =========================================================================
    # PRODUCT STEP 0 — Configure + Run
    # =========================================================================
    if pp_step == 0:
        st.markdown(
            "<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:12px;'>GENERATE A PRODUCT PAGE — CONTENT + DESIGNED HTML</p>",
            unsafe_allow_html=True,
        )
        st.markdown(_html(
            '<div class="glass-card" style="margin-bottom:18px;">'
            '<div style="color:#8b949e; font-size:0.82rem; line-height:1.7;">'
            "Select a viAct product. The system will:<br>"
            "&nbsp;1. Scrape the existing viact.ai product page (tone reference via Firecrawl)<br>"
            "&nbsp;2. Scrape competitor product pages (optional)<br>"
            "&nbsp;3. Generate all CMS fields + image prompts (Groq / Llama 3.3 70B)"
            "</div></div>"
        ), unsafe_allow_html=True)

        import sys as _sys
        _sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))
        from agent3_product_page import PRODUCT_VIACT_URLS

        _PRESET_PRODUCTS = list(PRODUCT_VIACT_URLS.keys()) + ["Custom (type below) →"]

        _pp_col1, _pp_col2 = st.columns([1, 1], gap="medium")
        with _pp_col1:
            pp_product_choice = st.selectbox("Product", _PRESET_PRODUCTS, key="pp_product_select")
            if pp_product_choice == "Custom (type below) →":
                pp_product = st.text_input(
                    "Custom product name",
                    placeholder="e.g. viSense Edge AI Module, AI Fatigue Detection...",
                    key="pp_custom_product_text",
                )
                pp_product_slug = st.text_input(
                    "URL slug (lowercase, hyphens)",
                    placeholder="e.g. visense-edge-ai",
                    key="pp_custom_slug_text",
                )
            else:
                pp_product = pp_product_choice
                pp_product_slug = pp_product_choice.lower().split(" ")[0].replace("(", "").replace(")", "").strip()

        with _pp_col2:
            pp_competitor_urls_raw = st.text_area(
                "Competitor product page URLs (one per line, optional)",
                placeholder="https://www.protex.ai/products/...\nhttps://visionify.ai/...",
                height=100,
                key="pp_competitor_urls",
            )

        _pp_doc = st.file_uploader(
            "📄 Upload product spec / brief .docx (auto-fills reference field)",
            type=["docx"],
            key="pp_doc_upload",
            help="Product specs, approved messaging, or case study data — used as ground-truth reference.",
        )
        if _pp_doc is not None:
            if st.session_state.get("pp_last_doc_name") != _pp_doc.name:
                try:
                    import io
                    import docx as _docx_lib
                    _doc_obj = _docx_lib.Document(io.BytesIO(_pp_doc.read()))
                    _doc_text = "\n".join(p.text for p in _doc_obj.paragraphs if p.text.strip())
                    st.session_state["pp_refs_text"] = _doc_text[:6000]
                    st.session_state["pp_last_doc_name"] = _pp_doc.name
                    st.success(f"✓ Loaded **{_pp_doc.name}** ({len(_doc_text):,} chars) → reference auto-filled")
                except Exception as _de:
                    st.error(f"Could not parse doc: {_de}")

        pp_refs = st.text_area(
            "Reference Material — paste product specs, metrics, or approved copy (or upload .docx above)",
            height=120,
            key="pp_refs_text",
        )

        with st.expander("⚙️ Custom Instructions (optional — key messaging, target audience, regional focus)"):
            pp_custom_inst = st.text_area(
                "",
                placeholder="e.g. Focus on oil & gas sector. Emphasise ATEX certification. Target UAE market.",
                height=100,
                key="pp_custom_inst",
            )

        st.write("")
        pp_run_disabled = not pp_product or not pp_product.strip()
        if st.button("🚀 Generate Product Page Content", disabled=pp_run_disabled, key="pp_run_btn", type="primary"):
            if not pp_product or not pp_product.strip():
                st.error("Please enter a product name.")
            else:
                _pp_viact_url = PRODUCT_VIACT_URLS.get(pp_product, "")
                _pp_competitor_urls = [
                    u.strip() for u in pp_competitor_urls_raw.strip().splitlines() if u.strip().startswith("http")
                ]

                _pp_progress = st.empty()

                with st.spinner("Scraping reference pages..."):
                    from agent2_data_extractor import extract_competitor_content as _ece
                    _pp_progress.info("🔍 Scraping viact.ai product page for tone reference...")

                    _pp_all_urls = ([_pp_viact_url] if _pp_viact_url else []) + _pp_competitor_urls
                    _pp_scraped = {}
                    if _pp_all_urls:
                        _pp_scraped = _ece(_pp_all_urls)

                    _pp_viact_md = (
                        _pp_scraped.get(_pp_viact_url, {}).get("markdown", "")
                        if _pp_viact_url else ""
                    )
                    _pp_competitor_data = {
                        k: v for k, v in _pp_scraped.items() if k != _pp_viact_url
                    }

                    _pp_progress.info("✍️ Generating product page content (Groq / Llama 3.3 70B)...")

                try:
                    from agent3_product_page import generate_product_page as _gpp
                    _pp_result = _gpp(
                        product_name=pp_product,
                        product_slug=pp_product_slug or pp_product.lower().replace(" ", "-"),
                        viact_page_content=_pp_viact_md,
                        competitor_content=_pp_competitor_data,
                        references=st.session_state.get("pp_refs_text", ""),
                        custom_instructions=st.session_state.get("pp_custom_inst", ""),
                    )
                    st.session_state["pp_content"] = _pp_result
                    st.session_state["pp_viact_md"] = _pp_viact_md
                    st.session_state["pp_product_label"] = pp_product
                    st.session_state["pp_step"] = 1
                    _pp_progress.empty()
                    st.rerun()
                except Exception as _pp_err:
                    _pp_progress.empty()
                    st.error(f"Generation failed: {_pp_err}")

    # =========================================================================
    # PRODUCT STEP 1 — Preview + Design + Export
    # =========================================================================
    elif pp_step == 1:
        _pp_result = st.session_state.get("pp_content", {})
        _pp_product_label = st.session_state.get("pp_product_label", "Product")

        st.markdown(
            f"<p style='color:#3fb950; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:12px;'>✓ {_pp_product_label.upper()} — CONTENT GENERATED</p>",
            unsafe_allow_html=True,
        )

        # ── Content Preview Tabs ───────────────────────────────────────────────
        _pp_t1, _pp_t2, _pp_t3, _pp_t4, _pp_t5, _pp_t6 = st.tabs([
            "📝 Hero & Problem",
            "⚡ Features & Steps",
            "🏭 Use Cases & Social Proof",
            "💬 Testimonials & FAQ",
            "🔍 SEO",
            "🎨 Image Prompts",
        ])

        with _pp_t1:
            _hero = _pp_result.get("hero_section", {})
            st.markdown("**Hero Section**")
            st.text_input("H1 Headline", value=_hero.get("h1", ""), key="pp_out_h1")
            st.text_area("Subheadline", value=_hero.get("subheadline", ""), height=70, key="pp_out_sub")
            st.text_area("Hero Body", value=_hero.get("hero_body", ""), height=90, key="pp_out_body")
            _pp_cta_c1, _pp_cta_c2 = st.columns(2)
            _pp_cta_c1.text_input("Primary CTA", value=_hero.get("primary_cta", ""), key="pp_out_cta1")
            _pp_cta_c2.text_input("Secondary CTA", value=_hero.get("secondary_cta", ""), key="pp_out_cta2")
            st.markdown("---")
            _prob = _pp_result.get("problem_statement", {})
            st.markdown("**Problem Statement**")
            st.text_input("Section Heading", value=_prob.get("heading", ""), key="pp_out_prob_h")
            st.text_area("Body", value=_prob.get("body", ""), height=100, key="pp_out_prob_b")
            for _pi, _pp_val in enumerate(_prob.get("pain_points", []), 1):
                st.text_input(f"Pain Point {_pi}", value=_pp_val, key=f"pp_out_pain_{_pi}")

        with _pp_t2:
            st.markdown("**Key Features (4)**")
            for _fi, _feat in enumerate(_pp_result.get("key_features", []), 1):
                with st.expander(f"Feature {_fi} — {_feat.get('title', '')}"):
                    st.text_input("Title", value=_feat.get("title", ""), key=f"pp_out_feat_title_{_fi}")
                    st.text_area("Description", value=_feat.get("description", ""), height=80, key=f"pp_out_feat_desc_{_fi}")
                    st.text_input("Icon Hint", value=_feat.get("icon_hint", ""), key=f"pp_out_feat_icon_{_fi}")
            st.markdown("---")
            st.markdown("**How It Works (3 Steps)**")
            for _si, _step in enumerate(_pp_result.get("how_it_works", []), 1):
                _sc1, _sc2 = st.columns([1, 3])
                _sc1.text_input(f"Step {_si} Title", value=_step.get("title", ""), key=f"pp_out_step_title_{_si}")
                _sc2.text_area(f"Step {_si} Description", value=_step.get("description", ""), height=70, key=f"pp_out_step_desc_{_si}")
            st.markdown("---")
            _specs = _pp_result.get("technical_specs", {})
            st.markdown("**Technical Specs**")
            st.text_area("Integrations", value="\n".join(_specs.get("integrations", [])), height=80, key="pp_out_specs_int")
            _sc1, _sc2 = st.columns(2)
            _sc1.text_input("Deployment Options", value=_specs.get("deployment_options", ""), key="pp_out_specs_deploy")
            _sc2.text_input("Compatibility", value=_specs.get("compatibility", ""), key="pp_out_specs_compat")

        with _pp_t3:
            st.markdown("**Use Cases (3)**")
            for _ui, _uc in enumerate(_pp_result.get("use_cases", []), 1):
                with st.expander(f"Use Case {_ui} — {_uc.get('industry', '')} · {_uc.get('title', '')}"):
                    st.text_input("Industry", value=_uc.get("industry", ""), key=f"pp_out_uc_ind_{_ui}")
                    st.text_input("Title", value=_uc.get("title", ""), key=f"pp_out_uc_title_{_ui}")
                    st.text_area("Description", value=_uc.get("description", ""), height=80, key=f"pp_out_uc_desc_{_ui}")
            st.markdown("---")
            st.markdown("**Social Proof Stats**")
            for _sti, _stat in enumerate(_pp_result.get("social_proof", {}).get("stats", []), 1):
                _stc1, _stc2 = st.columns([1, 2])
                _stc1.text_input(f"Metric {_sti}", value=_stat.get("metric", ""), key=f"pp_out_stat_m_{_sti}")
                _stc2.text_input(f"Label {_sti}", value=_stat.get("label", ""), key=f"pp_out_stat_l_{_sti}")

        with _pp_t4:
            st.markdown("**Testimonials (3)**")
            for _ti, _test in enumerate(_pp_result.get("testimonials", []), 1):
                with st.expander(f"Testimonial {_ti} — {_test.get('role', '')}"):
                    st.text_area("Quote", value=_test.get("quote", ""), height=100, key=f"pp_out_test_q_{_ti}")
                    st.text_input("Role", value=_test.get("role", ""), key=f"pp_out_test_r_{_ti}")
            st.markdown("---")
            st.markdown("**FAQs (5)**")
            for _qi, _faq in enumerate(_pp_result.get("faqs", []), 1):
                with st.expander(f"FAQ {_qi} — {_faq.get('question', '')[:60]}..."):
                    st.text_area("Question", value=_faq.get("question", ""), height=60, key=f"pp_out_faq_q_{_qi}")
                    st.text_area("Answer", value=_faq.get("answer", ""), height=100, key=f"pp_out_faq_a_{_qi}")
            st.markdown("---")
            _pp_cta_sec = _pp_result.get("cta_section", {})
            st.markdown("**CTA Section**")
            st.text_input("Heading", value=_pp_cta_sec.get("heading", ""), key="pp_out_cta_h")
            st.text_area("Description", value=_pp_cta_sec.get("description", ""), height=80, key="pp_out_cta_d")

        with _pp_t5:
            _pp_seo = _pp_result.get("seo_suite", {})
            _mt = _pp_seo.get("meta_title", "")
            _md = _pp_seo.get("meta_description", "")
            st.text_input(f"Meta Title ({len(_mt)}/60 chars)", value=_mt, key="pp_out_seo_title")
            if len(_mt) > 60:
                st.error(f"Meta title is {len(_mt)} chars — must be ≤60")
            st.text_area(f"Meta Description ({len(_md)}/160 chars)", value=_md, height=80, key="pp_out_seo_desc")
            if not (150 <= len(_md) <= 165):
                st.warning(f"Meta description is {len(_md)} chars — aim for 150-160")
            st.text_input("Primary Keyword", value=_pp_seo.get("primary_keyword", ""), key="pp_out_seo_kw")
            st.text_input("Keywords", value=", ".join(_pp_seo.get("keywords", [])), key="pp_out_seo_kws")
            st.text_input("Canonical Slug", value=_pp_seo.get("canonical_slug", ""), key="pp_out_seo_slug")

        with _pp_t6:
            st.markdown(
                "<div style='color:#8b949e; font-size:0.82rem; margin-bottom:16px;'>"
                "6 image prompts — copy each into Gemini / Imagen to generate visuals for the product page."
                "</div>",
                unsafe_allow_html=True,
            )
            for _img in _pp_result.get("image_prompts", []):
                _img_id = _img.get("id", "")
                _img_dim = _img.get("dimensions", "")
                _img_label = f"{_img_id.replace('_', ' ').title()} ({_img_dim})"
                st.text_area(_img_label, value=_img.get("prompt", ""), height=110, key=f"pp_img_{_img_id}")

        # ── Raw JSON ───────────────────────────────────────────────────────────
        with st.expander("📄 Raw JSON Output"):
            st.json(_pp_result)

# =============================================================================
# TAB — SOLUTIONS PAGES (Agent 08)
# Session state prefix: sol_
# =============================================================================
elif _sel == "solutions":
    import sys as _sys
    _sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))
    from agent8_solutions_page import SOLUTIONS_LIST as _SOL_LIST

    if "sol_step" not in st.session_state:
        st.session_state["sol_step"] = 0

    _sol_step = st.session_state["sol_step"]

    if _sol_step == 0:
        st.markdown(_html("""
        <div style="display:flex;gap:8px;margin-bottom:20px;align-items:stretch;">
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #f9c74f;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">🔧</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 1 — Solution Name</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Type the solution name — e.g. "Job Hazard Analysis Software", "Working at Height Safety Software". Or use a suggested topic from Market Radar.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #f9c74f;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">🤖</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 2 — Generate</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Click "Generate Solutions Page". The AI builds all Wix CMS fields in one shot — tagline, features (14 bullets), metrics, UVPs, 10 FAQs, and image alt texts.</div>
          </div>
          <div style="color:#1e1e1e;font-size:1.1rem;align-self:center;padding:0 2px;">→</div>
          <div style="flex:1;background:#0e0e0e;border:1px solid #1a1a1a;border-top:2px solid #f9c74f;border-radius:8px;padding:14px 16px;">
            <div style="font-size:1.1rem;margin-bottom:6px;">📊</div>
            <div style="color:#e6edf3;font-weight:700;font-size:0.8rem;margin-bottom:4px;">Step 3 — Review &amp; Save</div>
            <div style="color:#444;font-size:0.74rem;line-height:1.5;">Review across 3 tabs — Webpage Content, SEO, Image Alts. Click "Save to Google Sheets" — saved as tab "Sol — {Solution}".</div>
          </div>
        </div>
        """), unsafe_allow_html=True)

    if _sol_step > 0:
        st.write("")
        if st.button("↩ Start Over", key="sol_reset"):
            for _k in [_k for _k in st.session_state if _k.startswith("sol_")]:
                del st.session_state[_k]
            st.rerun()

    st.write("")

    # ── Suggested topic banner (from Daily Intel Scan) ────────────────────────
    _sug_sol = st.session_state.get("suggested_solutions_topic")
    if _sug_sol and _sol_step == 0:
        _ss_col, _ss_btn = st.columns([5, 1])
        with _ss_col:
            st.info(f"💡 **Today's suggested topic:** {_sug_sol.get('solution_name', '')} — _{_sug_sol.get('why', '')}_")
        with _ss_btn:
            st.write("")
            if st.button("Use This →", key="sol_use_suggested", use_container_width=True):
                st.session_state["sol_input_val"] = _sug_sol.get("solution_name", "")
                st.rerun()

    # =========================================================================
    # SOLUTIONS STEP 0 — Select Solution & Generate
    # =========================================================================
    if _sol_step == 0:
        st.markdown(
            "<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:12px;'>GENERATE A FULL SOLUTIONS LANDING PAGE</p>",
            unsafe_allow_html=True,
        )
        st.markdown(_html(
            '<div class="glass-card" style="margin-bottom:18px;">'
            '<div style="color:#8b949e; font-size:0.82rem; line-height:1.7;">'
            "Enter the solution name. System will:<br>"
            "&nbsp;1. Live research — Tavily / Google News RSS fallback<br>"
            "&nbsp;2. Generate complete Wix CMS page: Hero, Features (14 bullets in 5 tabs), Metrics, UVPs, FAQs, SEO, Image Alt Texts"
            "</div></div>"
        ), unsafe_allow_html=True)

        # Prefill from Market Radar
        _sol_prefill = st.session_state.pop("sol_prefill", "") or st.session_state.pop("sol_input_val", "")

        _sol_final = st.text_input(
            "Solution Name",
            value=_sol_prefill or "",
            placeholder="e.g. Working at Height Safety Software",
            key="sol_input",
        ).strip()

        with st.expander("Quick pick from standard solutions list"):
            _sol_pick = st.selectbox("Standard solutions", options=[""] + _SOL_LIST, key="sol_select")
            if _sol_pick:
                st.caption(f"Selected: **{_sol_pick}** — copy it into the field above")

        _sol_col1, _sol_col2 = st.columns([1, 3])
        with _sol_col1:
            _sol_run_tavily = st.checkbox("Live research (Tavily/RSS)", value=True, key="sol_tavily")

        st.write("")

        if st.button("🔧  Generate Solutions Page", type="primary", key="sol_gen_btn", disabled=not _sol_final):
            if not _sol_final.strip():
                st.warning("Please enter a solution name.")
                st.stop()

            _prog = st.empty()
            with _prog.container():
                st.info("Step 1/2 — Researching solution via Tavily / Google News RSS...")

            try:
                from agent8_solutions_page import generate_solutions_page as _gen_sol

                with _prog.container():
                    st.info("Step 2/2 — Generating all Wix CMS fields (Llama 3.3 70B)...")

                _sol_result = _gen_sol(
                    solution_name=_sol_final,
                    run_tavily=_sol_run_tavily,
                    progress_callback=lambda msg: None,
                )
                st.session_state["sol_result"]        = _sol_result
                st.session_state["sol_solution_label"] = _sol_final
                st.session_state["sol_step"]           = 1
                # Auto-save to Sheets
                try:
                    from push_to_sheets import push_solutions_page as _push_sol_auto
                    _push_sol_auto(result=_sol_result)
                    st.toast(f"✓ Auto-saved to Sheets: Sol — {_sol_final}")
                except Exception as _sol_ae:
                    st.toast(f"⚠ Auto-save failed: {_sol_ae}")
                _prog.empty()
                st.rerun()
            except Exception as _sol_err:
                _prog.empty()
                st.error(f"Generation failed: {_sol_err}")

    # =========================================================================
    # SOLUTIONS STEP 1 — Preview + Push to Sheets
    # =========================================================================
    elif _sol_step == 1:
        _sol_result = st.session_state["sol_result"]
        _sol_cms    = _sol_result.get("cms_fields", {})
        _sol_errors = _sol_result.get("quality_gate_errors", [])
        _sol_label  = st.session_state.get("sol_solution_label", _sol_cms.get("solution_name", ""))

        # ── Quality gate warning ─────────────────────────────────────────────
        if _sol_errors:
            st.warning(
                "⚠️ **Quality gate retry was triggered** — issues detected in first generation:\n"
                + "\n".join(f"- {e}" for e in _sol_errors)
                + "\n\nReview the output below to confirm corrections were applied."
            )

        st.markdown(_html(
            '<div class="glass-card" style="border-color:rgba(249,199,79,0.3); background:rgba(22,25,33,0.8);">'
            '<div style="display:flex; align-items:center; justify-content:space-between; flex-wrap:wrap; gap:12px;">'
            '<div style="display:flex; align-items:center; gap:12px;">'
            '<div style="background:rgba(249,199,79,0.15); border:1px solid rgba(249,199,79,0.4); border-radius:50%; width:40px; height:40px; display:flex; align-items:center; justify-content:center; font-size:1.2rem; flex-shrink:0;">🔧</div>'
            '<div>'
            '<div style="color:#f9c74f; font-size:0.72rem; font-weight:700; text-transform:uppercase; letter-spacing:1.5px; margin-bottom:3px;">Solutions Page Ready</div>'
            f'<h3 style="margin:0; color:#e6edf3; font-size:1.1rem;">{_t(_sol_label)}</h3>'
            '</div></div>'
            '<div style="display:flex; gap:8px; flex-wrap:wrap;">'
            '<span style="background:rgba(249,199,79,0.1); color:#f9c74f; border:1px solid rgba(249,199,79,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#129302; Llama 3.3 written</span>'
            '<span style="background:rgba(63,185,80,0.08); color:#3fb950; border:1px solid rgba(63,185,80,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#127979; Wix CMS ready</span>'
            '<span style="background:rgba(88,166,255,0.1); color:#58a6ff; border:1px solid rgba(88,166,255,0.25); border-radius:6px; padding:4px 10px; font-size:0.75rem;">&#128203; 50+ fields</span>'
            '</div></div></div>'
        ), unsafe_allow_html=True)

        _sol_push_col, _sol_info_col = st.columns([1, 3])
        with _sol_push_col:
            if st.button("📊  Save to Google Sheets", type="primary", key="sol_push_btn", use_container_width=True):
                try:
                    from push_to_sheets import push_solutions_page as _push_sol
                    _push_sol(result=_sol_result)
                    _sheet_id = os.getenv("SHEET_ID", "")
                    _sheet_url = f"https://docs.google.com/spreadsheets/d/{_sheet_id}/edit"
                    st.success(f"✅ Saved to **'Sol — {_sol_label}'** tab in your Sheet — [Open Sheet ↗]({_sheet_url})")
                except Exception as _push_err:
                    st.error(f"Sheets error: {_push_err}")
        with _sol_info_col:
            st.markdown(_html(
                f'<div style="background:rgba(22,25,33,0.5); border:1px solid #2d303a; border-radius:8px; padding:10px 14px; font-size:0.8rem; color:#8b949e; margin-top:4px;">'
                f'&#128221; Creates a <strong style="color:#c9d1d9;">Sol — {_t(_sol_label)}</strong> tab in your Sheet. Fields in col A, values in col B — copy directly into Wix CMS.'
                f'</div>'
            ), unsafe_allow_html=True)

        st.markdown("<hr/>", unsafe_allow_html=True)
        st.markdown("<p style='color:#8b949e; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:1.8px; margin-bottom:8px;'>PREVIEW ALL SECTIONS</p>", unsafe_allow_html=True)

        _sol_t1, _sol_t2, _sol_t3, _sol_t4 = st.tabs([
            "📝 Webpage Content",
            "🔍 SEO",
            "🖼 Image Alts",
            "🎨 Image Prompts",
        ])

        with _sol_t1:
            st.markdown(
                "<div style='background:rgba(249,199,79,0.06); border:1px solid rgba(249,199,79,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; font-size:0.83rem; color:#f9c74f;'>"
                "<strong>Wix CMS Fields</strong> — Each field maps directly to a Wix CMS dynamic field. Copy one field at a time into the matching Wix field."
                "</div>",
                unsafe_allow_html=True,
            )

            # HERO
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:10px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Hero Section</div>", unsafe_allow_html=True)
            st.text_input("Title (Text 1)", value=_sol_cms.get("solution_name", ""), key="sol_title")
            st.text_input("Tagline (Text 2)", value=_sol_cms.get("tagline", ""), key="sol_tagline")
            st.text_area("Short Description (Text 3)", value=_sol_cms.get("short_description", ""), height=80, key="sol_short_desc")

            # TESTIMONIAL
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Testimonial</div>", unsafe_allow_html=True)
            st.text_area("Quote (Text 5)", value=_sol_cms.get("testimonial_quote", ""), height=70, key="sol_testimonial")
            st.text_input("Attribution (review subtitle)", value=_sol_cms.get("testimonial_attribution", ""), key="sol_attrib")

            # DIFFERENCE SECTION
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Difference Section</div>", unsafe_allow_html=True)
            st.text_input("Diff Section Heading (Text 7)", value=_sol_cms.get("diff_section_title", ""), key="sol_diff_h")
            _d1, _d2, _d3 = st.columns(3)
            with _d1:
                st.caption("TRENDS")
                st.text_input("Title (Text 8)", value=_sol_cms.get("trend_title", ""), key="sol_trend_t")
                st.text_area("Description (Text 9)", value=_sol_cms.get("trend_description", ""), height=120, key="sol_trend_d")
            with _d2:
                st.caption("STATISTICS")
                st.text_input("Title (Text 10)", value=_sol_cms.get("stats_title", ""), key="sol_stats_t")
                st.text_area("Description (Text 11)", value=_sol_cms.get("stats_description", ""), height=120, key="sol_stats_d")
            with _d3:
                st.caption("OUTCOME")
                st.text_input("Title (Text 12)", value=_sol_cms.get("outcome_title", ""), key="sol_outcome_t")
                st.text_area("Description (Text 13)", value=_sol_cms.get("outcome_description", ""), height=120, key="sol_outcome_d")

            # CTA
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>CTA Block</div>", unsafe_allow_html=True)
            st.text_area("CTA Block (Text 14)", value=_sol_cms.get("cta_text", ""), height=90, key="sol_cta_text")
            st.text_input("CTA Button (Text 15)", value=_sol_cms.get("cta_button", ""), key="sol_cta_btn")

            # FEATURES
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Key Features (5 Tabs × Bullets)</div>", unsafe_allow_html=True)
            st.text_input("Features Section Title (Text 16)", value=_sol_cms.get("features_title", ""), key="sol_feat_title")
            _bullet_groups = [(1, 3), (4, 6), (7, 9), (10, 11), (12, 14)]
            for _fi in range(1, 6):
                _tab_name = _sol_cms.get(f"feature_tab_{_fi}", f"Feature Tab {_fi}")
                with st.expander(f"Tab {_fi}: {_tab_name}"):
                    st.text_input(f"Tab Name (Text {16+_fi})", value=_tab_name, key=f"sol_tab_{_fi}")
                    _b_start, _b_end = _bullet_groups[_fi - 1]
                    for _bi in range(_b_start, _b_end + 1):
                        st.text_area(f"Bullet {_bi}", value=_sol_cms.get(f"bullet_{_bi}", ""), height=70, key=f"sol_bullet_{_bi}")

            # METRICS
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Post Deployment Metrics (3 Stats)</div>", unsafe_allow_html=True)
            _mc1, _mc2, _mc3 = st.columns(3)
            for _mi, _mc in enumerate([_mc1, _mc2, _mc3], start=1):
                with _mc:
                    st.text_input(f"Metric {_mi} Value (Text {36+(_mi-1)*2})", value=_sol_cms.get(f"metric_{_mi}_value", ""), key=f"sol_m{_mi}v")
                    st.text_area(f"Metric {_mi} Desc (Text {37+(_mi-1)*2})", value=_sol_cms.get(f"metric_{_mi}_desc", ""), height=90, key=f"sol_m{_mi}d")

            # UVPs
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Unique Value Propositions (5)</div>", unsafe_allow_html=True)
            for _ui in range(1, 6):
                _uc1, _uc2 = st.columns([1, 2])
                with _uc1:
                    st.text_input(f"UVP {_ui} Title (Text {43+(_ui-1)*2})", value=_sol_cms.get(f"uvp_{_ui}_title", ""), key=f"sol_uvp{_ui}t")
                with _uc2:
                    st.text_area(f"UVP {_ui} Desc (Text {44+(_ui-1)*2})", value=_sol_cms.get(f"uvp_{_ui}_desc", ""), height=70, key=f"sol_uvp{_ui}d")

            # BOTTOM CTA
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Bottom CTA</div>", unsafe_allow_html=True)
            st.text_input("Bottom CTA (new cta text)", value=_sol_cms.get("new_cta_text", ""), key="sol_new_cta")

            # FAQs
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>FAQs — Ask Our Expert (5)</div>", unsafe_allow_html=True)
            for _qi in range(1, 6):
                _fq = _sol_cms.get(f"faq_{_qi}_q", "")
                _fa = _sol_cms.get(f"faq_{_qi}_a", "")
                if _fq or _fa:
                    with st.expander(f"FAQ {_qi}: {_fq[:60]}{'...' if len(_fq) > 60 else ''}"):
                        st.text_input("Q", value=_fq, key=f"sol_faq{_qi}q")
                        st.text_area("A", value=_fa, height=120, key=f"sol_faq{_qi}a")

        with _sol_t2:
            st.markdown(
                "<div style='background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; font-size:0.83rem; color:#58a6ff;'>"
                "<strong>SEO Fields</strong> — Paste meta title &amp; description into Wix SEO settings. Use slug for the page URL. Keywords for on-page copy."
                "</div>",
                unsafe_allow_html=True,
            )
            _c1, _c2 = st.columns(2)
            with _c1:
                st.text_input("Slug (Solutions URL)", value=_sol_cms.get("slug", ""), key="sol_slug")
                _stc = len(_sol_cms.get("seo_meta_title", ""))
                st.text_input(f"Meta Title", value=_sol_cms.get("seo_meta_title", ""), key="sol_seo_title")
                st.caption(f"{_stc} chars — target 50-60")
                _sdc = len(_sol_cms.get("seo_meta_description", ""))
                st.text_area("Meta Description", value=_sol_cms.get("seo_meta_description", ""), height=90, key="sol_seo_desc")
                st.caption(f"{_sdc} chars — target 140-165")
            with _c2:
                st.text_area("Meta Keywords", value=_sol_cms.get("seo_keywords", ""), height=140, key="sol_keywords")

        with _sol_t3:
            st.markdown(
                "<div style='background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; font-size:0.83rem; color:#58a6ff;'>"
                "<strong>Image Alt Texts</strong> — Copy each alt text into the matching image field in Wix CMS. Field names match exactly."
                "</div>",
                unsafe_allow_html=True,
            )
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:10px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Hero &amp; Sections</div>", unsafe_allow_html=True)
            st.text_input("hero image", value=_sol_cms.get("hero_image_alt", ""), key="sol_alt_hero")
            st.text_input("Trends img", value=_sol_cms.get("trends_image_alt", ""), key="sol_alt_trends")
            st.text_input("STATISTICS img", value=_sol_cms.get("stats_image_alt", ""), key="sol_alt_stats")
            st.text_input("OUTCOME img", value=_sol_cms.get("outcome_image_alt", ""), key="sol_alt_outcome")
            st.text_input("dashboard img", value=_sol_cms.get("dashboard_image_alt", ""), key="sol_alt_dash")
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Key Features Images (5)</div>", unsafe_allow_html=True)
            for _fi in range(1, 6):
                st.text_input(f"{_fi} Key features", value=_sol_cms.get(f"feature_{_fi}_img_alt", ""), key=f"sol_alt_feat{_fi}")
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Unique Value Images (5)</div>", unsafe_allow_html=True)
            for _ui in range(1, 6):
                st.text_input(f"{_ui} Unique Value", value=_sol_cms.get(f"uvp_{_ui}_img_alt", ""), key=f"sol_alt_uvp{_ui}")
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin:14px 0 6px; text-transform:uppercase; letter-spacing:1.2px;'>Other</div>", unsafe_allow_html=True)
            st.text_input("1 check", value=_sol_cms.get("check_img_alt", ""), key="sol_alt_check")
            st.text_input("Solution list image", value=_sol_cms.get("list_image_alt", ""), key="sol_alt_list")

        with _sol_t4:
            _sol_imgs = _sol_result.get("image_prompts", [])
            if not _sol_imgs:
                st.info("No image prompts in this result — regenerate to get Nano Banana prompts.")
            else:
                st.markdown(
                    "<div style='background:rgba(88,166,255,0.06); border:1px solid rgba(88,166,255,0.2); border-radius:8px; padding:12px 16px; margin-bottom:14px; font-size:0.83rem; color:#58a6ff;'>"
                    "<strong>Nano Banana Prompts</strong> — Copy each prompt and paste into Nano Banana (or any image tool). "
                    "7 prompts with exact Wix dimensions: Hero (1155×764), Dashboard (1620×705), "
                    "Key Features 1–5 (800×672 / 755×561 / 699×498 / 794×504 / 851×572)."
                    "</div>",
                    unsafe_allow_html=True,
                )
                for _i, _img in enumerate(_sol_imgs, 1):
                    _placement = _img.get("placement", f"Image {_i}")
                    _prompt    = _img.get("prompt", "")
                    _alt       = _img.get("alt_text", "")
                    with st.expander(f"Image {_i} — {_placement}"):
                        st.text_area("Prompt", value=_prompt, height=140, key=f"sol_img_prompt_{_i}")
                        st.markdown(
                            f"<div style='color:#8b949e; font-size:0.82rem; margin-top:6px;'>"
                            f"<strong style='color:#e6edf3;'>Alt text:</strong> {_t(_alt)}</div>",
                            unsafe_allow_html=True,
                        )

# =============================================================================
# TAB — BLOG WRITER (Agent 09)
# Session state prefix: blog_
# =============================================================================
elif _sel == "blog":
    import sys as _sys
    _sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))

    st.markdown(_html(
        '<div style="margin-bottom:18px;">'
        '<div style="font-family:\'Oxanium\',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#e879f9; margin-bottom:4px;">AGENT 09</div>'
        '<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.3rem; color:#fff; margin-bottom:4px;">Blog / SEO Article Writer</div>'
        '<div style="font-family:\'Jost\',sans-serif; font-weight:300; font-size:0.88rem; color:#8b949e;">Generate a full, CMS-ready SEO blog post — introduction, body sections, conclusion, CTA, meta tags, and image prompts. All from a topic + keyword.</div>'
        '</div>'
    ), unsafe_allow_html=True)

    st.markdown(_html("""
    <div style="display:flex;gap:8px;margin-bottom:20px;align-items:stretch;">
      <div style="flex:1;background:#12121c;border:1px solid rgba(255,255,255,0.05);border-top:2px solid #e879f9;border-radius:8px;padding:14px 16px;">
        <div style="font-family:'Oxanium',sans-serif;font-size:0.55rem;font-weight:600;letter-spacing:2px;text-transform:uppercase;color:#e879f9;margin-bottom:8px;">STEP 01</div>
        <div style="color:#E9ECF1;font-family:'Jost',sans-serif;font-weight:700;font-size:0.82rem;margin-bottom:6px;">Enter Topic &amp; Keyword</div>
        <div style="color:#8b949e;font-size:0.74rem;line-height:1.55;">Type your blog topic and primary SEO keyword. Optionally use a suggested topic from the Market Radar scan.</div>
      </div>
      <div style="color:#3a3a4a;font-size:1.1rem;align-self:center;padding:0 4px;">&#8594;</div>
      <div style="flex:1;background:#12121c;border:1px solid rgba(255,255,255,0.05);border-top:2px solid #e879f9;border-radius:8px;padding:14px 16px;">
        <div style="font-family:'Oxanium',sans-serif;font-size:0.55rem;font-weight:600;letter-spacing:2px;text-transform:uppercase;color:#e879f9;margin-bottom:8px;">STEP 02</div>
        <div style="color:#E9ECF1;font-family:'Jost',sans-serif;font-weight:700;font-size:0.82rem;margin-bottom:6px;">Generate</div>
        <div style="color:#8b949e;font-size:0.74rem;line-height:1.55;">Click "Generate Blog Post". AI researches the topic via Tavily/RSS and writes a full CMS-ready blog with all sections, meta tags, and image prompts.</div>
      </div>
      <div style="color:#3a3a4a;font-size:1.1rem;align-self:center;padding:0 4px;">&#8594;</div>
      <div style="flex:1;background:#12121c;border:1px solid rgba(255,255,255,0.05);border-top:2px solid #e879f9;border-radius:8px;padding:14px 16px;">
        <div style="font-family:'Oxanium',sans-serif;font-size:0.55rem;font-weight:600;letter-spacing:2px;text-transform:uppercase;color:#e879f9;margin-bottom:8px;">STEP 03</div>
        <div style="color:#E9ECF1;font-family:'Jost',sans-serif;font-weight:700;font-size:0.82rem;margin-bottom:6px;">Review &amp; Save</div>
        <div style="color:#8b949e;font-size:0.74rem;line-height:1.55;">Review across 4 tabs — Blog Content, SEO &amp; Meta, Image Prompts, Raw JSON. Click "Save to Google Sheets" — saved as tab "Blog — {Title}".</div>
      </div>
    </div>
    """), unsafe_allow_html=True)

    # ── "Use This →" banner from Market Radar suggested topics ───────────────
    _sug_blog = (
        st.session_state.get("suggested_va_topic")
        or st.session_state.get("suggested_solutions_topic")
        or st.session_state.get("suggested_industry_topic")
    )
    if _sug_blog and "blog_result" not in st.session_state:
        _sug_name = (
            _sug_blog.get("detection_type")
            or _sug_blog.get("solution_name")
            or _sug_blog.get("topic", "")
        )
        _sug_why  = _sug_blog.get("why", "")
        _sb_col, _sb_btn = st.columns([5, 1])
        with _sb_col:
            st.info(f"💡 **Suggested blog topic:** {_sug_name} — _{_sug_why}_")
        with _sb_btn:
            st.write("")
            if st.button("Use This →", key="blog_use_suggested", use_container_width=True):
                st.session_state["blog_topic_in"] = _sug_name
                st.rerun()

    _bc1, _bc2 = st.columns([3, 2], gap="medium")
    with _bc1:
        _blog_topic   = st.text_input("Blog Topic", placeholder="e.g. PPE Detection in Oil & Gas Industry", key="blog_topic_in")
        _blog_keyword = st.text_input("Target Keyword", placeholder="e.g. PPE detection software", key="blog_kw_in")
        _blog_industry = st.text_input("Industry (optional)", placeholder="e.g. Oil & Gas, Construction, Manufacturing", key="blog_industry_in", value="Industrial")
    with _bc2:
        _blog_words = st.selectbox("Target Word Count", [800, 1200, 1500, 2000, 2500, 3000], index=1, key="blog_words_in")
        st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)
        _blog_go = st.button("Generate Blog Post", type="primary", use_container_width=True, key="blog_go_btn")

    if _blog_go:
        if not _blog_topic or not _blog_keyword:
            st.warning("Please enter a topic and target keyword.")
        else:
            _blog_bar = st.progress(0, text="Starting...")
            _blog_status = st.empty()
            def _blog_prog(msg):
                _blog_status.markdown(f"<div style='color:#e879f9; font-size:0.85rem;'>{msg}</div>", unsafe_allow_html=True)

            try:
                from agent9_blog_writer import generate_blog_post as _gen_blog
                _blog_bar.progress(10, text="Researching...")
                _blog_result = _gen_blog(
                    topic=_blog_topic,
                    target_keyword=_blog_keyword,
                    industry=_blog_industry or "Industrial",
                    word_count=_blog_words,
                    progress_callback=_blog_prog,
                )
                _blog_bar.progress(100, text="Done!")
                _blog_status.empty()
                st.session_state["blog_result"] = _blog_result
                st.success("Blog post generated!")
            except Exception as _blog_err:
                _blog_bar.progress(0)
                st.error(f"Error: {_blog_err}")

    if "blog_result" in st.session_state:
        _br = st.session_state["blog_result"]
        _bc = _br.get("cms_fields", {})
        _bs = _br.get("seo", {})

        st.markdown("<hr/>", unsafe_allow_html=True)

        _bt1, _bt2, _bt3, _bt4 = st.tabs(["📝 Blog Content", "🔍 SEO & Meta", "🎨 Image Prompts", "🔧 Raw JSON"])

        with _bt1:
            st.markdown("<div style='color:#ff6a3d; font-weight:700; font-size:0.9rem; margin-bottom:8px; text-transform:uppercase; letter-spacing:1.2px;'>CMS Fields</div>", unsafe_allow_html=True)
            st.text_input("Blog Title (H1)", value=_bc.get("blog_title", ""), key="blog_out_title")
            st.text_input("Slug", value=_bc.get("slug", ""), key="blog_out_slug")
            _bcat, _btag = st.columns(2)
            with _bcat:
                st.text_input("Category", value=_bc.get("category", ""), key="blog_out_cat")
            with _btag:
                st.text_input("Tags", value=", ".join(_bc.get("tags", [])), key="blog_out_tags")
            st.text_input("Reading Time", value=_bc.get("reading_time", ""), key="blog_out_rt")
            st.text_area("Excerpt (blog card summary)", value=_bc.get("excerpt", ""), height=70, key="blog_out_excerpt")
            st.markdown("<div style='color:#e879f9; font-weight:600; font-size:0.82rem; margin:12px 0 4px; text-transform:uppercase; letter-spacing:1px;'>Body</div>", unsafe_allow_html=True)
            st.text_area("Introduction", value=_bc.get("introduction", ""), height=120, key="blog_out_intro")
            if _bc.get("tldr"):
                st.markdown("<div style='color:#e879f9; font-weight:600; font-size:0.82rem; margin:12px 0 4px; text-transform:uppercase; letter-spacing:1px;'>TLDR</div>", unsafe_allow_html=True)
                _tldr_text = "\n".join([f"{i}. {p}" for i, p in enumerate(_bc["tldr"], 1)])
                st.text_area("TLDR (top summary block)", value=_tldr_text, height=100, key="blog_out_tldr")
            if _bc.get("table_of_contents"):
                st.markdown("<div style='color:#e879f9; font-weight:600; font-size:0.82rem; margin:12px 0 4px; text-transform:uppercase; letter-spacing:1px;'>Table of Contents</div>", unsafe_allow_html=True)
                _toc_text = "\n".join([f"{i}. {h}" for i, h in enumerate(_bc["table_of_contents"], 1)])
                st.text_area("TOC (paste into Wix TOC block)", value=_toc_text, height=100, key="blog_out_toc")
            for _bi, _sec in enumerate(_bc.get("body_sections", []), 1):
                with st.expander(f"Section {_bi}: {_sec.get('heading', '')}"):
                    st.text_input(f"H2 Heading {_bi}", value=_sec.get("heading", ""), key=f"blog_sec_h_{_bi}")
                    if _sec.get("content"):
                        st.text_area(f"Intro Paragraph {_bi}", value=_sec.get("content", ""), height=100, key=f"blog_sec_c_{_bi}")
                    for _sj, _sub in enumerate(_sec.get("subsections", []), 1):
                        st.text_input(f"  H3 Subheading {_bi}.{_sj}", value=_sub.get("subheading", ""), key=f"blog_sub_h_{_bi}_{_sj}")
                        st.text_area(f"  Content {_bi}.{_sj}", value=_sub.get("content", ""), height=120, key=f"blog_sub_c_{_bi}_{_sj}")
            if _bc.get("case_study"):
                with st.expander("📌 Case Study Section"):
                    _cs = _bc["case_study"]
                    st.text_input("Section Heading", value=_cs.get("heading", "Seeing It in Practice"), key="blog_cs_heading")
                    st.text_input("Client", value=_cs.get("client", ""), key="blog_cs_client")
                    st.text_area("Challenge", value=_cs.get("challenge", ""), height=70, key="blog_cs_challenge")
                    st.text_area("Solution", value=_cs.get("solution", ""), height=70, key="blog_cs_solution")
                    st.text_area("Outcome", value=_cs.get("outcome", ""), height=70, key="blog_cs_outcome")
            st.text_area("Conclusion", value=_bc.get("conclusion", ""), height=100, key="blog_out_conclusion")
            if _bc.get("key_takeaways"):
                st.markdown("<div style='color:#e879f9; font-weight:600; font-size:0.82rem; margin:12px 0 4px; text-transform:uppercase; letter-spacing:1px;'>Key Takeaways</div>", unsafe_allow_html=True)
                _kt_text = "\n".join([f"• {t}" for t in _bc["key_takeaways"]])
                st.text_area("Key Takeaways (bullet list)", value=_kt_text, height=120, key="blog_out_kt")
            if _bc.get("faqs"):
                with st.expander(f"❓ FAQs ({len(_bc['faqs'])})"):
                    for _fi, _faq in enumerate(_bc["faqs"], 1):
                        st.text_input(f"Q{_fi}", value=_faq.get("question", ""), key=f"blog_faq_q_{_fi}")
                        st.text_area(f"A{_fi}", value=_faq.get("answer", ""), height=70, key=f"blog_faq_a_{_fi}")
            st.markdown("<div style='color:#e879f9; font-weight:600; font-size:0.82rem; margin:12px 0 4px; text-transform:uppercase; letter-spacing:1px;'>CTA Block</div>", unsafe_allow_html=True)
            _bca, _bcb, _bcc = st.columns(3)
            with _bca:
                st.text_input("CTA Heading", value=_bc.get("cta_heading", ""), key="blog_cta_h")
            with _bcb:
                st.text_input("CTA Button Text", value=_bc.get("cta_button_text", ""), key="blog_cta_btn")
            with _bcc:
                st.text_input("CTA URL", value=_bc.get("cta_url", ""), key="blog_cta_url")
            st.text_area("CTA Body", value=_bc.get("cta_body", ""), height=70, key="blog_cta_body")

            if _bc.get("internal_links"):
                st.markdown("<div style='color:#e879f9; font-weight:600; font-size:0.82rem; margin:12px 0 4px; text-transform:uppercase; letter-spacing:1px;'>Suggested Internal Links</div>", unsafe_allow_html=True)
                for _lnk in _bc.get("internal_links", []):
                    st.markdown(f"<div style='color:#8b949e; font-size:0.82rem;'>→ <strong style='color:#e6edf3;'>{_t(_lnk.get('anchor',''))}</strong> → {_t(_lnk.get('suggested_page',''))}</div>", unsafe_allow_html=True)

        with _bt2:
            _mt_v = _bc.get("meta_title", "")
            _md_v = _bc.get("meta_description", "")
            st.text_input(f"Meta Title ({len(_mt_v)}/60 chars)", value=_mt_v, key="blog_out_mt")
            if len(_mt_v) > 60:
                st.error(f"Meta title {len(_mt_v)} chars — must be ≤60")
            st.text_area(f"Meta Description ({len(_md_v)}/155 chars)", value=_md_v, height=80, key="blog_out_md")
            if len(_md_v) > 155:
                st.warning(f"Meta description {len(_md_v)} chars — aim for ≤155")
            st.text_input("Focus Keyword", value=_bs.get("focus_keyword", ""), key="blog_out_fk")
            st.text_input("Secondary Keywords", value=", ".join(_bs.get("secondary_keywords", [])), key="blog_out_sk")
            _sco = _bs.get("keyword_in_title", False)
            _scd = _bs.get("keyword_in_desc", False)
            st.markdown(
                f"<div style='display:flex; gap:16px; margin-top:10px;'>"
                f"<span style='color:{'#3fb950' if _sco else '#f85149'}; font-size:0.82rem;'>{'✓' if _sco else '✗'} Keyword in title</span>"
                f"<span style='color:{'#3fb950' if _scd else '#f85149'}; font-size:0.82rem;'>{'✓' if _scd else '✗'} Keyword in description</span>"
                f"<span style='color:#8b949e; font-size:0.82rem;'>~{_bs.get('estimated_word_count', 0)} words</span>"
                f"</div>",
                unsafe_allow_html=True,
            )

        with _bt3:
            st.text_area("Hero Image Alt Text", value=_bc.get("image_alt_main", ""), height=60, key="blog_out_img_alt")
            st.text_area("Hero Image Prompt (for AI generation)", value=_bc.get("image_prompt_main", ""), height=120, key="blog_out_img_prompt")

        with _bt4:
            st.json(_br)

        st.write("")
        _blog_push_col, _ = st.columns([1, 2])
        with _blog_push_col:
            if st.button("📊  Save to Google Sheets", type="primary", key="blog_push_btn", use_container_width=True):
                try:
                    from push_to_sheets import push_blog_post as _push_blog
                    _push_blog(_br)
                    _saved_title = _br.get("cms_fields", {}).get("blog_title", "Blog")[:40]
                    st.success(f"✅ Saved — tab 'Blog — {_saved_title}'")
                except Exception as _bp_err:
                    st.error(f"Save failed: {_bp_err}")

# =============================================================================
# TAB — META / SEO TAG GENERATOR (Agent 10)
# Session state prefix: meta_
# =============================================================================
elif _sel == "meta_seo":
    import sys as _sys
    _sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))

    st.markdown(_html(
        '<div style="margin-bottom:18px;">'
        '<div style="font-family:\'Oxanium\',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#38bdf8; margin-bottom:4px;">AGENT 10</div>'
        '<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.3rem; color:#fff; margin-bottom:4px;">Meta / SEO Tag Generator</div>'
        '<div style="font-family:\'Jost\',sans-serif; font-weight:300; font-size:0.88rem; color:#8b949e;">Generate all SEO meta tags, Open Graph, Twitter Card, and JSON-LD schema for any viAct page. Paste directly into Wix SEO settings.</div>'
        '</div>'
    ), unsafe_allow_html=True)

    _PAGE_TYPE_OPTIONS = {
        "Industry Page":       "industry",
        "Product Page":        "product",
        "Solution Page":       "solution",
        "Blog Post":           "blog",
        "Case Study":          "case_study",
        "Video Analytics Page": "video_analytics",
        "Homepage":            "homepage",
    }

    _mc1, _mc2 = st.columns([3, 2], gap="medium")
    with _mc1:
        _meta_page_label = st.selectbox("Page Type", list(_PAGE_TYPE_OPTIONS.keys()), key="meta_page_type_sel")
        _meta_page_type  = _PAGE_TYPE_OPTIONS[_meta_page_label]
        _meta_title      = st.text_input("Page Title", placeholder="e.g. PPE Detection Software for Construction", key="meta_title_in")
        _meta_desc_in    = st.text_area("Page Description (brief)", placeholder="Describe what this page is about in 1-2 sentences...", height=80, key="meta_desc_in")
    with _mc2:
        _meta_keyword    = st.text_input("Primary Keyword", placeholder="e.g. PPE detection software", key="meta_kw_in")
        _meta_slug       = st.text_input("URL Slug", placeholder="e.g. /solutions/ppe-detection", key="meta_slug_in")
        _meta_sec_kws    = st.text_input("Secondary Keywords (optional, comma-separated)", key="meta_sec_kws_in")
        _meta_go = st.button("Generate SEO Tags", type="primary", use_container_width=True, key="meta_go_btn")

    if _meta_go:
        if not _meta_title or not _meta_keyword or not _meta_slug:
            st.warning("Please fill in Page Title, Primary Keyword, and URL Slug.")
        else:
            _meta_bar = st.progress(0, text="Generating...")
            _meta_status = st.empty()
            def _meta_prog(msg):
                _meta_status.markdown(f"<div style='color:#38bdf8; font-size:0.85rem;'>{msg}</div>", unsafe_allow_html=True)

            _sec_kw_list = [k.strip() for k in _meta_sec_kws.split(",") if k.strip()] if _meta_sec_kws else []

            try:
                from agent10_meta_seo import generate_meta_tags as _gen_meta
                _meta_bar.progress(20, text="Generating meta tags...")
                _meta_result = _gen_meta(
                    page_type=_meta_page_type,
                    page_title=_meta_title,
                    page_description=_meta_desc_in,
                    target_keyword=_meta_keyword,
                    url_slug=_meta_slug,
                    secondary_keywords=_sec_kw_list,
                    progress_callback=_meta_prog,
                )
                _meta_bar.progress(100, text="Done!")
                _meta_status.empty()
                st.session_state["meta_result"] = _meta_result
                st.success("SEO tags generated!")
            except Exception as _meta_err:
                _meta_bar.progress(0)
                st.error(f"Error: {_meta_err}")

    if "meta_result" in st.session_state:
        _mr = st.session_state["meta_result"]
        _mt = _mr.get("meta_tags", {})
        _ms = _mr.get("seo_scores", {})

        st.markdown("<hr/>", unsafe_allow_html=True)

        # SEO Score bar
        _sc_val = _ms.get("overall_score", "0/9")
        _sc_num = int(_sc_val.split("/")[0]) if "/" in _sc_val else 0
        _sc_color = "#3fb950" if _sc_num >= 7 else ("#f9c74f" if _sc_num >= 4 else "#f85149")
        st.markdown(
            f"<div style='background:#12121c; border:1px solid rgba(255,255,255,0.06); border-radius:8px; padding:12px 16px; margin-bottom:16px; display:flex; align-items:center; gap:16px;'>"
            f"<div style='font-family:\"Oxanium\",sans-serif; font-size:1.6rem; font-weight:700; color:{_sc_color};'>{_sc_val}</div>"
            f"<div style='font-family:\"Jost\",sans-serif; font-size:0.85rem; color:#8b949e;'>SEO Score — {'Excellent' if _sc_num >= 7 else ('Good' if _sc_num >= 4 else 'Needs Work')}</div>"
            f"</div>",
            unsafe_allow_html=True,
        )

        _mtt1, _mtt2, _mtt3, _mtt4 = st.tabs(["🏷 Meta Tags", "📣 Open Graph & Twitter", "🔧 Schema JSON-LD", "📋 HTML Snippet"])

        with _mtt1:
            _m_title_v = _mt.get("meta_title", "")
            _m_desc_v  = _mt.get("meta_description", "")
            st.text_input(f"Meta Title ({len(_m_title_v)}/60 chars)", value=_m_title_v, key="meta_out_title")
            if len(_m_title_v) > 60:
                st.error(f"Title is {len(_m_title_v)} chars — must be ≤60")
            st.text_area(f"Meta Description ({len(_m_desc_v)}/155 chars)", value=_m_desc_v, height=80, key="meta_out_desc")
            if len(_m_desc_v) > 155:
                st.warning(f"Description {len(_m_desc_v)} chars — aim for ≤155")
            st.text_input("Canonical URL", value=_mt.get("canonical", ""), key="meta_out_canonical")
            st.text_input("Robots", value=_mt.get("robots", "index, follow"), key="meta_out_robots")
            st.text_input("Focus Keyword", value=_mt.get("focus_keyword", ""), key="meta_out_fk")
            st.text_input("Secondary Keywords", value=", ".join(_mt.get("secondary_keywords", [])), key="meta_out_sk")
            _check_cols = st.columns(4)
            checks = [
                ("Keyword in title", _ms.get("keyword_in_title")),
                ("Keyword in desc", _ms.get("keyword_in_description")),
                ("OG tags", _ms.get("has_og_tags")),
                ("FAQ schema", _ms.get("has_faq_schema")),
            ]
            for (_lbl, _val), _col in zip(checks, _check_cols):
                _col.markdown(f"<div style='color:{'#3fb950' if _val else '#f85149'}; font-size:0.82rem; text-align:center;'>{'✓' if _val else '✗'} {_lbl}</div>", unsafe_allow_html=True)

        with _mtt2:
            st.text_input("OG Title", value=_mt.get("og_title", ""), key="meta_out_og_title")
            st.text_area("OG Description", value=_mt.get("og_description", ""), height=80, key="meta_out_og_desc")
            st.text_input("OG Image URL", value=_mt.get("og_image", ""), key="meta_out_og_img")
            st.text_input("OG Image Alt", value=_mt.get("og_image_alt", ""), key="meta_out_og_alt")
            st.text_input("OG Type", value=_mt.get("og_type", "website"), key="meta_out_og_type")
            st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)
            st.text_input("Twitter Title", value=_mt.get("twitter_title", ""), key="meta_out_tw_title")
            st.text_area("Twitter Description", value=_mt.get("twitter_description", ""), height=70, key="meta_out_tw_desc")
            st.text_input("Twitter Card Type", value=_mt.get("twitter_card", "summary_large_image"), key="meta_out_tw_card")
            st.text_input("Twitter Site", value=_mt.get("twitter_site", "@viactai"), key="meta_out_tw_site")

        with _mtt3:
            _schema = _mr.get("schema_json_ld", {})
            if _schema.get("webpage"):
                st.markdown("<div style='color:#38bdf8; font-weight:600; font-size:0.82rem; margin-bottom:6px;'>WebPage / Product Schema</div>", unsafe_allow_html=True)
                st.text_area("JSON-LD (paste into Wix SEO > Schema)", value=_schema.get("webpage", ""), height=220, key="meta_out_schema_wp")
            if _schema.get("faq"):
                st.markdown("<div style='color:#38bdf8; font-weight:600; font-size:0.82rem; margin:12px 0 6px;'>FAQ Schema</div>", unsafe_allow_html=True)
                st.text_area("FAQ JSON-LD", value=_schema.get("faq", ""), height=220, key="meta_out_schema_faq")
                if _mr.get("faq_pairs"):
                    st.markdown("<div style='color:#8b949e; font-size:0.82rem; margin:10px 0 4px;'>Generated FAQ pairs:</div>", unsafe_allow_html=True)
                    for _fq in _mr.get("faq_pairs", []):
                        with st.expander(_fq.get("question", "")):
                            st.write(_fq.get("answer", ""))

        with _mtt4:
            st.markdown("<div style='color:#8b949e; font-size:0.82rem; margin-bottom:8px;'>Copy this entire block and paste into your Wix page's &lt;head&gt; custom code section.</div>", unsafe_allow_html=True)
            st.text_area("Ready-to-paste HTML", value=_mr.get("html_snippet", ""), height=420, key="meta_out_html")


# =============================================================================
# TAB — PARTNER OUTREACH (Agent 11)
# Session state prefix: po_
# =============================================================================
elif _sel == "partner":
    import sys as _sys
    _sys.path.insert(0, os.path.join(os.path.dirname(__file__), "tools"))

    st.markdown(_html(
        '<div style="margin-bottom:18px;">'
        '<div style="font-family:\'Oxanium\',sans-serif; font-size:0.52rem; font-weight:600; letter-spacing:3px; text-transform:uppercase; color:#22d3ee; margin-bottom:4px;">AGENT 11</div>'
        '<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.3rem; color:#fff; margin-bottom:4px;">Partner Outreach — Competitor Partner Discovery</div>'
        '<div style="font-family:\'Jost\',sans-serif; font-weight:300; font-size:0.88rem; color:#8b949e;">Auto-discovers competitors, extracts their partners, enriches with emails / websites / descriptions, and pushes BD-ready leads to the Partnership Leads Google Sheet.</div>'
        '</div>'
    ), unsafe_allow_html=True)

    # Sheet link — big CTA at top
    _po_sheet_id = os.getenv("PARTNER_SHEET_ID", "1Q2XJZ2STaCN94DK4JEnS1mHkrgILfFljjNCc1dy_5qw")
    _po_sheet_url = f"https://docs.google.com/spreadsheets/d/{_po_sheet_id}/edit"
    st.markdown(_html(
        f'<a href="{_po_sheet_url}" target="_blank" style="display:block; text-decoration:none;">'
        f'<div style="background:linear-gradient(135deg, #22d3ee 0%, #0891b2 100%); border-radius:10px; padding:14px 20px; margin-bottom:20px;">'
        f'<div style="font-family:\'Oxanium\',sans-serif; font-size:0.55rem; font-weight:600; letter-spacing:2px; text-transform:uppercase; color:rgba(0,0,0,0.65);">Google Sheet</div>'
        f'<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.1rem; color:#fff; margin-top:3px;">Open Partnership Leads Sheet →</div>'
        f'</div></a>'
    ), unsafe_allow_html=True)

    # Load dashboard data from the sheet
    _po_load_ok = False
    _po_error = ""
    try:
        from push_to_sheets import get_sheets_service
        _po_svc = get_sheets_service()

        # Get all tabs
        _po_meta = _po_svc.spreadsheets().get(spreadsheetId=_po_sheet_id).execute()
        _po_all_tabs = [s["properties"]["title"] for s in _po_meta.get("sheets", [])]

        # Get Competitors tab
        _po_comp_resp = _po_svc.spreadsheets().values().get(
            spreadsheetId=_po_sheet_id, range="'Competitors'!A2:G",
        ).execute()
        _po_comp_rows = _po_comp_resp.get("values", [])

        # Categorize competitors by Status
        _po_tracked, _po_blank, _po_skip = [], [], []
        for _r in _po_comp_rows:
            while len(_r) < 7:
                _r.append("")
            _po_name = _r[0]
            _po_website = _r[1]
            _po_status = _r[6].strip()
            _po_item = {"name": _po_name, "website": _po_website}
            if _po_status == "Track":
                _po_tracked.append(_po_item)
            elif _po_status.lower() in ("skip", "done"):
                _po_skip.append(_po_item)
            else:
                _po_blank.append(_po_item)

        # Count partners per tracked tab
        _po_partner_tabs = []
        for _t in _po_tracked:
            if _t["name"] in _po_all_tabs:
                _resp = _po_svc.spreadsheets().values().get(
                    spreadsheetId=_po_sheet_id, range=f"'{_t['name']}'!A2:A",
                ).execute()
                _n = len(_resp.get("values", []))
                if _n:
                    _po_partner_tabs.append({"name": _t["name"], "count": _n})

        _po_total_partners = sum(_p["count"] for _p in _po_partner_tabs)
        _po_load_ok = True
    except Exception as _e:
        _po_error = str(_e)

    if not _po_load_ok:
        st.error(f"Could not load sheet data: {_po_error[:200]}")
    else:
        # Top metrics
        _pm1, _pm2, _pm3, _pm4 = st.columns(4, gap="small")
        with _pm1:
            st.markdown(_html(
                f'<div style="background:#12121c; border:1px solid rgba(34,211,238,0.15); border-radius:8px; padding:14px;">'
                f'<div style="font-family:\'Oxanium\',sans-serif; font-size:0.52rem; letter-spacing:2px; color:#22d3ee; text-transform:uppercase;">Tracked Competitors</div>'
                f'<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.8rem; color:#fff; margin-top:3px;">{len(_po_tracked)}</div>'
                f'</div>'
            ), unsafe_allow_html=True)
        with _pm2:
            st.markdown(_html(
                f'<div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:8px; padding:14px;">'
                f'<div style="font-family:\'Oxanium\',sans-serif; font-size:0.52rem; letter-spacing:2px; color:#8b949e; text-transform:uppercase;">Pending Review</div>'
                f'<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.8rem; color:#fff; margin-top:3px;">{len(_po_blank)}</div>'
                f'</div>'
            ), unsafe_allow_html=True)
        with _pm3:
            st.markdown(_html(
                f'<div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:8px; padding:14px;">'
                f'<div style="font-family:\'Oxanium\',sans-serif; font-size:0.52rem; letter-spacing:2px; color:#8b949e; text-transform:uppercase;">Partners Discovered</div>'
                f'<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.8rem; color:#fff; margin-top:3px;">{_po_total_partners}</div>'
                f'</div>'
            ), unsafe_allow_html=True)
        with _pm4:
            st.markdown(_html(
                f'<div style="background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:8px; padding:14px;">'
                f'<div style="font-family:\'Oxanium\',sans-serif; font-size:0.52rem; letter-spacing:2px; color:#8b949e; text-transform:uppercase;">Skipped (0 Partners)</div>'
                f'<div style="font-family:\'Jost\',sans-serif; font-weight:700; font-size:1.8rem; color:#fff; margin-top:3px;">{len(_po_skip)}</div>'
                f'</div>'
            ), unsafe_allow_html=True)

        st.markdown("<div style='height:16px;'></div>", unsafe_allow_html=True)

        _pt1, _pt2, _pt3 = st.tabs(["📊 Partner Tabs", "🎯 Competitors", "⚙️ How It Works"])

        # ── Tab 1: Partner tabs with counts ──────────────────────────────────
        with _pt1:
            if not _po_partner_tabs:
                st.info("No competitor tabs with partners yet. The daily cron runs Mon-Fri at 6:30 AM IST.")
            else:
                st.markdown(f"<div style='color:#8b949e; font-size:0.82rem; margin-bottom:10px;'>{len(_po_partner_tabs)} tabs · sorted by partner count</div>", unsafe_allow_html=True)
                _sorted = sorted(_po_partner_tabs, key=lambda x: -x["count"])
                for _ptb in _sorted:
                    _tab_url = f"{_po_sheet_url}#gid=0"
                    st.markdown(_html(
                        f'<div style="display:flex; justify-content:space-between; align-items:center; background:#12121c; border:1px solid rgba(255,255,255,0.05); border-radius:6px; padding:10px 14px; margin-bottom:6px;">'
                        f'<div style="font-family:\'Jost\',sans-serif; font-weight:600; font-size:0.9rem; color:#e9ecf1;">{_t(_ptb["name"])}</div>'
                        f'<div style="font-family:\'Oxanium\',sans-serif; font-weight:700; font-size:1rem; color:#22d3ee;">{_ptb["count"]}</div>'
                        f'</div>'
                    ), unsafe_allow_html=True)

        # ── Tab 2: Competitors (Track / Pending / Skip) ──────────────────────
        with _pt2:
            _sub1, _sub2, _sub3 = st.tabs([f"Track ({len(_po_tracked)})", f"Pending ({len(_po_blank)})", f"Skip ({len(_po_skip)})"])
            with _sub1:
                for _c in _po_tracked:
                    st.markdown(f"• **{_c['name']}** — {_c['website']}")
            with _sub2:
                if not _po_blank:
                    st.info("No competitors awaiting Track/Skip decision.")
                else:
                    st.markdown(f"<div style='color:#8b949e; font-size:0.8rem; margin-bottom:8px;'>Mark these Track or Skip in the sheet.</div>", unsafe_allow_html=True)
                    for _c in _po_blank[:50]:
                        st.markdown(f"• {_c['name']} — {_c['website']}")
                    if len(_po_blank) > 50:
                        st.markdown(f"<div style='color:#8b949e; font-size:0.75rem;'>...and {len(_po_blank)-50} more (open sheet to see all)</div>", unsafe_allow_html=True)
            with _sub3:
                for _c in _po_skip:
                    st.markdown(f"• ~~{_c['name']}~~ — {_c['website']}")

        # ── Tab 3: How it works ──────────────────────────────────────────────
        with _pt3:
            st.markdown(_html(
                '<div style="color:#c9d1d9; font-size:0.9rem; line-height:1.6;">'
                '<b style="color:#22d3ee;">Daily Pipeline (6:30 AM IST, Mon-Fri):</b><br>'
                '1. Rotates through Track-status competitors (1 per day)<br>'
                '2. Discovers partners across 6 sources — sitemap, homepage, /partners, /integrations, news, case-studies<br>'
                '3. Uses LLM (Groq) to classify + filter for viAct-relevant industrial-safety BD leads<br>'
                '4. Discovers each partner\'s real website (canonical brand map + DDG + LLM verify)<br>'
                '5. Scrapes emails (5-tier: website → footer → social → DDG → WHOIS)<br>'
                '6. Fills descriptions / country / address via LLM<br>'
                '7. Every Monday also runs Agent 1 (finds NEW competitors)<br>'
                '<br>'
                '<b style="color:#22d3ee;">Sheet Schema (per competitor tab):</b><br>'
                'Company Name · Description · Website · Phone · Email · Address · Country · Status · Email Source · Discovered Via · Discovered At · Relationship<br>'
                '<br>'
                '<b style="color:#22d3ee;">Manual Actions in Sheet:</b><br>'
                '• Mark competitor <b>Track</b>/<b>Skip</b> in Competitors tab (Status column)<br>'
                '• Mark partner <b>Shortlist</b>/<b>Done</b> in individual competitor tabs<br>'
                '• Add missing emails/phones manually — pipeline won\'t overwrite<br>'
                '</div>'
            ), unsafe_allow_html=True)
